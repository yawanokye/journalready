from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import re
import secrets
import sqlite3
import uuid
from contextlib import contextmanager
from datetime import date, datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable, Iterator

from docx import Document
from docx.shared import Pt


PAYMENT_DB_PATH = Path(os.environ.get("ARTICLEREADY_SQLITE_DB_PATH", "articleready_payments.db"))
DEFAULT_REVIEW_DB = PAYMENT_DB_PATH.parent / "articleready_review_workspace.db"
REVIEW_DB_PATH = Path(os.environ.get("ARTICLEREADY_REVIEW_DB_PATH", str(DEFAULT_REVIEW_DB)))

SCHEMA = """
CREATE TABLE IF NOT EXISTS review_projects (
    id TEXT PRIMARY KEY,
    token_hash TEXT NOT NULL,
    title TEXT NOT NULL,
    article_type TEXT NOT NULL DEFAULT 'Systematic review',
    review_question TEXT NOT NULL DEFAULT '',
    protocol_positioning TEXT NOT NULL DEFAULT 'Auto',
    eligibility_criteria TEXT NOT NULL DEFAULT '',
    screening_process TEXT NOT NULL DEFAULT '',
    quality_appraisal TEXT NOT NULL DEFAULT '',
    synthesis_method TEXT NOT NULL DEFAULT '',
    software TEXT NOT NULL DEFAULT '',
    notes TEXT NOT NULL DEFAULT '',
    status TEXT NOT NULL DEFAULT 'active',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS review_search_runs (
    id TEXT PRIMARY KEY,
    project_id TEXT NOT NULL,
    database_name TEXT NOT NULL,
    platform TEXT NOT NULL DEFAULT '',
    source_route TEXT NOT NULL DEFAULT 'database',
    search_string TEXT NOT NULL DEFAULT '',
    search_date TEXT NOT NULL DEFAULT '',
    date_limits TEXT NOT NULL DEFAULT '',
    language_limits TEXT NOT NULL DEFAULT '',
    document_types TEXT NOT NULL DEFAULT '',
    reported_result_count INTEGER,
    imported_record_count INTEGER NOT NULL DEFAULT 0,
    notes TEXT NOT NULL DEFAULT '',
    created_at TEXT NOT NULL,
    FOREIGN KEY(project_id) REFERENCES review_projects(id) ON DELETE CASCADE
);
CREATE TABLE IF NOT EXISTS review_records (
    id TEXT PRIMARY KEY,
    project_id TEXT NOT NULL,
    search_run_id TEXT,
    source_database TEXT NOT NULL DEFAULT '',
    source_route TEXT NOT NULL DEFAULT 'database',
    external_id TEXT NOT NULL DEFAULT '',
    title TEXT NOT NULL,
    normalized_title TEXT NOT NULL,
    abstract TEXT NOT NULL DEFAULT '',
    authors TEXT NOT NULL DEFAULT '',
    publication_year TEXT NOT NULL DEFAULT '',
    doi TEXT NOT NULL DEFAULT '',
    normalized_doi TEXT NOT NULL DEFAULT '',
    url TEXT NOT NULL DEFAULT '',
    keywords TEXT NOT NULL DEFAULT '',
    document_type TEXT NOT NULL DEFAULT '',
    source_title TEXT NOT NULL DEFAULT '',
    raw_json TEXT NOT NULL DEFAULT '{}',
    duplicate_of TEXT,
    duplicate_candidate_of TEXT,
    duplicate_confidence REAL,
    title_abstract_decision TEXT NOT NULL DEFAULT 'not_screened',
    title_abstract_reason TEXT NOT NULL DEFAULT '',
    full_text_decision TEXT NOT NULL DEFAULT 'not_assessed',
    full_text_reason TEXT NOT NULL DEFAULT '',
    reviewer_notes TEXT NOT NULL DEFAULT '',
    full_text_filename TEXT NOT NULL DEFAULT '',
    full_text_text TEXT NOT NULL DEFAULT '',
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL,
    FOREIGN KEY(project_id) REFERENCES review_projects(id) ON DELETE CASCADE,
    FOREIGN KEY(search_run_id) REFERENCES review_search_runs(id) ON DELETE SET NULL,
    FOREIGN KEY(duplicate_of) REFERENCES review_records(id) ON DELETE SET NULL,
    FOREIGN KEY(duplicate_candidate_of) REFERENCES review_records(id) ON DELETE SET NULL
);
CREATE TABLE IF NOT EXISTS review_audit_events (
    id TEXT PRIMARY KEY,
    project_id TEXT NOT NULL,
    record_id TEXT,
    event_type TEXT NOT NULL,
    event_json TEXT NOT NULL DEFAULT '{}',
    created_at TEXT NOT NULL,
    FOREIGN KEY(project_id) REFERENCES review_projects(id) ON DELETE CASCADE,
    FOREIGN KEY(record_id) REFERENCES review_records(id) ON DELETE SET NULL
);
CREATE INDEX IF NOT EXISTS idx_review_records_project ON review_records(project_id);
CREATE INDEX IF NOT EXISTS idx_review_records_doi ON review_records(project_id, normalized_doi);
CREATE INDEX IF NOT EXISTS idx_review_records_title ON review_records(project_id, normalized_title);
CREATE INDEX IF NOT EXISTS idx_review_records_ta ON review_records(project_id, title_abstract_decision);
CREATE INDEX IF NOT EXISTS idx_review_records_ft ON review_records(project_id, full_text_decision);
CREATE INDEX IF NOT EXISTS idx_review_records_duplicate ON review_records(project_id, duplicate_of);
CREATE INDEX IF NOT EXISTS idx_review_search_project ON review_search_runs(project_id);
CREATE INDEX IF NOT EXISTS idx_review_audit_project ON review_audit_events(project_id);
"""


def _utc_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _hash_token(token: str) -> str:
    return hashlib.sha256(str(token or "").encode("utf-8")).hexdigest()


def _json(value: Any) -> str:
    return json.dumps(value if value is not None else {}, ensure_ascii=False, default=str)


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalise_doi(value: Any) -> str:
    text = _clean(value).lower()
    text = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", text)
    text = re.sub(r"^doi:\s*", "", text)
    return text.strip().rstrip(".,;)")


def normalise_title(value: Any) -> str:
    text = _clean(value).lower()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _year(value: Any) -> str:
    match = re.search(r"(?:19|20)\d{2}", str(value or ""))
    return match.group(0) if match else ""


@contextmanager
def connection() -> Iterator[sqlite3.Connection]:
    REVIEW_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(REVIEW_DB_PATH, timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("PRAGMA journal_mode = WAL")
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def init_review_workspace_tables() -> None:
    with connection() as conn:
        conn.executescript(SCHEMA)


def _row(row: sqlite3.Row | None) -> dict[str, Any] | None:
    if row is None:
        return None
    data = dict(row)
    for key in ("raw_json", "event_json"):
        if key in data and isinstance(data[key], str):
            try:
                data[key] = json.loads(data[key] or "{}")
            except Exception:
                data[key] = {}
    return data


def _audit(conn: sqlite3.Connection, project_id: str, event_type: str, payload: Any, record_id: str | None = None) -> None:
    conn.execute(
        "INSERT INTO review_audit_events (id, project_id, record_id, event_type, event_json, created_at) VALUES (?, ?, ?, ?, ?, ?)",
        (str(uuid.uuid4()), project_id, record_id, event_type, _json(payload), _utc_iso()),
    )


def create_project(data: dict[str, Any]) -> dict[str, Any]:
    init_review_workspace_tables()
    title = _clean(data.get("title"))
    if len(title) < 3:
        raise ValueError("Enter a review workspace title.")
    project_id = str(uuid.uuid4())
    access_token = secrets.token_urlsafe(32)
    now = _utc_iso()
    with connection() as conn:
        conn.execute(
            """
            INSERT INTO review_projects (
                id, token_hash, title, article_type, review_question, protocol_positioning,
                eligibility_criteria, screening_process, quality_appraisal, synthesis_method,
                software, notes, status, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'active', ?, ?)
            """,
            (
                project_id,
                _hash_token(access_token),
                title,
                _clean(data.get("article_type")) or "Systematic review",
                _clean(data.get("review_question")),
                _clean(data.get("protocol_positioning")) or "Auto",
                _clean(data.get("eligibility_criteria")),
                _clean(data.get("screening_process")),
                _clean(data.get("quality_appraisal")),
                _clean(data.get("synthesis_method")),
                _clean(data.get("software")),
                _clean(data.get("notes")),
                now,
                now,
            ),
        )
        _audit(conn, project_id, "project_created", {"title": title})
    project = get_project(project_id, access_token)
    assert project is not None
    project["access_token"] = access_token
    return project


def verify_project(project_id: str, token: str) -> dict[str, Any]:
    if not project_id or not token:
        raise PermissionError("Review workspace access is missing.")
    with connection() as conn:
        row = conn.execute("SELECT * FROM review_projects WHERE id=?", (project_id,)).fetchone()
    project = _row(row)
    if not project or not secrets.compare_digest(str(project.get("token_hash") or ""), _hash_token(token)):
        raise PermissionError("Review workspace access is invalid or no longer available.")
    project.pop("token_hash", None)
    return project


def get_project(project_id: str, token: str) -> dict[str, Any]:
    project = verify_project(project_id, token)
    project["summary"] = calculate_summary(project_id)
    project["search_runs"] = list_search_runs(project_id)
    return project


def update_project(project_id: str, token: str, data: dict[str, Any]) -> dict[str, Any]:
    verify_project(project_id, token)
    allowed = {
        "title", "article_type", "review_question", "protocol_positioning",
        "eligibility_criteria", "screening_process", "quality_appraisal",
        "synthesis_method", "software", "notes", "status",
    }
    updates = {key: _clean(value) for key, value in data.items() if key in allowed and value is not None}
    if "title" in updates and len(updates["title"]) < 3:
        raise ValueError("Enter a review workspace title.")
    if updates:
        columns = ", ".join(f"{key}=?" for key in updates)
        params = list(updates.values()) + [_utc_iso(), project_id]
        with connection() as conn:
            conn.execute(f"UPDATE review_projects SET {columns}, updated_at=? WHERE id=?", params)
            _audit(conn, project_id, "project_updated", updates)
    return get_project(project_id, token)


def delete_project(project_id: str, token: str) -> None:
    verify_project(project_id, token)
    with connection() as conn:
        conn.execute("DELETE FROM review_projects WHERE id=?", (project_id,))


def create_search_run(project_id: str, data: dict[str, Any]) -> dict[str, Any]:
    database_name = _clean(data.get("database_name"))
    if not database_name:
        raise ValueError("Enter the database, platform or source used for this search.")
    run_id = str(uuid.uuid4())
    source_route = _clean(data.get("source_route")) or "database"
    if source_route not in {"database", "backward_citation", "forward_citation", "manual_journal", "expert_recommendation", "grey_literature", "other"}:
        source_route = "other"
    reported_count = data.get("reported_result_count")
    try:
        reported_count = None if reported_count in (None, "") else max(0, int(reported_count))
    except Exception:
        reported_count = None
    with connection() as conn:
        conn.execute(
            """
            INSERT INTO review_search_runs (
                id, project_id, database_name, platform, source_route, search_string,
                search_date, date_limits, language_limits, document_types,
                reported_result_count, imported_record_count, notes, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, ?, ?)
            """,
            (
                run_id, project_id, database_name, _clean(data.get("platform")), source_route,
                str(data.get("search_string") or "").strip(), _clean(data.get("search_date")),
                _clean(data.get("date_limits")), _clean(data.get("language_limits")),
                _clean(data.get("document_types")), reported_count, _clean(data.get("notes")), _utc_iso(),
            ),
        )
        _audit(conn, project_id, "search_run_created", {"search_run_id": run_id, "database_name": database_name, "source_route": source_route})
        row = conn.execute("SELECT * FROM review_search_runs WHERE id=?", (run_id,)).fetchone()
    return _row(row) or {}


def list_search_runs(project_id: str) -> list[dict[str, Any]]:
    with connection() as conn:
        rows = conn.execute(
            "SELECT * FROM review_search_runs WHERE project_id=? ORDER BY created_at DESC",
            (project_id,),
        ).fetchall()
    return [_row(row) or {} for row in rows]


def _find_value(row: dict[str, Any], candidates: Iterable[str]) -> Any:
    lowered = {str(key).strip().lower(): value for key, value in row.items()}
    for name in candidates:
        if name.lower() in lowered and lowered[name.lower()] not in (None, ""):
            return lowered[name.lower()]
    return ""


def _canonical_record(row: dict[str, Any]) -> dict[str, Any] | None:
    title = _clean(_find_value(row, ["title", "article title", "document title", "ti", "t1"]))
    if not title:
        return None
    authors = _find_value(row, ["authors", "author", "au", "af", "creator"])
    if isinstance(authors, list):
        authors = "; ".join(_clean(item) for item in authors if _clean(item))
    abstract = _find_value(row, ["abstract", "ab", "n2", "summary", "description"])
    year = _year(_find_value(row, ["year", "publication year", "py", "y1", "date", "published"] ))
    doi = normalise_doi(_find_value(row, ["doi", "digital object identifier", "di", "do"]))
    url = _clean(_find_value(row, ["url", "link", "ur", "web of science link", "scopus link"]))
    external_id = _clean(_find_value(row, ["external_id", "id", "eid", "ut", "accession number", "an"]))
    keywords = _find_value(row, ["keywords", "author keywords", "de", "kw"])
    if isinstance(keywords, list):
        keywords = "; ".join(_clean(item) for item in keywords if _clean(item))
    return {
        "title": title,
        "abstract": str(abstract or "").strip(),
        "authors": _clean(authors),
        "publication_year": year,
        "doi": doi,
        "url": url,
        "external_id": external_id,
        "keywords": _clean(keywords),
        "document_type": _clean(_find_value(row, ["document type", "type", "dt", "ty"])),
        "source_title": _clean(_find_value(row, ["source title", "journal", "publication name", "jo", "jf", "t2"])),
        "raw": row,
    }


def parse_csv_records(content: bytes, delimiter: str | None = None) -> list[dict[str, Any]]:
    text = content.decode("utf-8-sig", errors="replace")
    if delimiter is None:
        sample = text[:8192]
        try:
            delimiter = csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
        except Exception:
            delimiter = ","
    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
    records = []
    for raw in reader:
        item = _canonical_record(dict(raw or {}))
        if item:
            records.append(item)
    return records


def parse_ris_records(content: bytes) -> list[dict[str, Any]]:
    text = content.decode("utf-8-sig", errors="replace")
    rows: list[dict[str, Any]] = []
    current: dict[str, Any] = {}
    for line in text.splitlines():
        match = re.match(r"^([A-Z0-9]{2})\s*-\s*(.*)$", line.rstrip())
        if not match:
            continue
        tag, value = match.groups()
        if tag == "TY" and current:
            item = _canonical_record(current)
            if item:
                rows.append(item)
            current = {}
        if tag in {"AU", "A1", "A2", "KW"}:
            current.setdefault(tag, []).append(value.strip())
        else:
            previous = current.get(tag)
            current[tag] = f"{previous} {value}".strip() if previous and tag in {"AB", "N2"} else value.strip()
        if tag == "ER":
            item = _canonical_record(current)
            if item:
                rows.append(item)
            current = {}
    if current:
        item = _canonical_record(current)
        if item:
            rows.append(item)
    return rows


def parse_bibtex_records(content: bytes) -> list[dict[str, Any]]:
    text = content.decode("utf-8-sig", errors="replace")
    records: list[dict[str, Any]] = []
    starts = [m.start() for m in re.finditer(r"(?m)^\s*@\w+\s*\{", text)]
    starts.append(len(text))
    for idx in range(len(starts) - 1):
        block = text[starts[idx]:starts[idx + 1]]
        entry_type = re.search(r"@([A-Za-z]+)", block)
        row: dict[str, Any] = {"type": entry_type.group(1) if entry_type else ""}
        key_match = re.search(r"@\w+\s*\{\s*([^,]+)", block)
        if key_match:
            row["id"] = key_match.group(1).strip()
        field_pattern = re.compile(
            r"(?ms)^\s*([A-Za-z][A-Za-z0-9_-]*)\s*=\s*(?:\{((?:[^{}]|\{[^{}]*\})*)\}|\"(.*?)\")\s*,?"
        )
        for match in field_pattern.finditer(block):
            key = match.group(1).lower()
            value = match.group(2) if match.group(2) is not None else match.group(3)
            value = re.sub(r"[{}]", "", value or "")
            row[key] = re.sub(r"\s+", " ", value).strip()
        item = _canonical_record(row)
        if item:
            records.append(item)
    return records


def parse_json_records(content: bytes) -> list[dict[str, Any]]:
    value = json.loads(content.decode("utf-8-sig", errors="replace"))
    if isinstance(value, dict):
        value = value.get("records") or value.get("items") or value.get("results") or [value]
    records = []
    for raw in value if isinstance(value, list) else []:
        if not isinstance(raw, dict):
            continue
        item = _canonical_record(raw)
        if item:
            records.append(item)
    return records


def parse_xlsx_records(content: bytes) -> list[dict[str, Any]]:
    from openpyxl import load_workbook
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheet = workbook.active
    rows = sheet.iter_rows(values_only=True)
    headers = [str(value or "").strip() for value in next(rows, [])]
    records = []
    for values in rows:
        raw = {headers[idx]: values[idx] for idx in range(min(len(headers), len(values))) if headers[idx]}
        item = _canonical_record(raw)
        if item:
            records.append(item)
    return records


def parse_records(filename: str, content: bytes) -> list[dict[str, Any]]:
    suffix = Path(filename or "").suffix.lower()
    if suffix in {".ris"}:
        return parse_ris_records(content)
    if suffix in {".bib", ".bibtex"}:
        return parse_bibtex_records(content)
    if suffix in {".json"}:
        return parse_json_records(content)
    if suffix in {".xlsx", ".xlsm"}:
        return parse_xlsx_records(content)
    if suffix in {".tsv"}:
        return parse_csv_records(content, delimiter="\t")
    if suffix in {".csv", ".txt"}:
        return parse_csv_records(content)
    raise ValueError("Upload a RIS, BibTeX, CSV, TSV, XLSX or JSON database export.")


def import_records(project_id: str, search_run_id: str, filename: str, content: bytes) -> dict[str, Any]:
    parsed = parse_records(filename, content)
    if not parsed:
        raise ValueError("No records with usable titles were found in the uploaded export.")
    with connection() as conn:
        run = conn.execute(
            "SELECT * FROM review_search_runs WHERE id=? AND project_id=?",
            (search_run_id, project_id),
        ).fetchone()
        if not run:
            raise ValueError("The search run could not be found.")
        run_data = dict(run)
        existing_rows = conn.execute(
            "SELECT id, normalized_title, publication_year, normalized_doi, duplicate_of FROM review_records WHERE project_id=?",
            (project_id,),
        ).fetchall()
        existing = [dict(row) for row in existing_rows]
        doi_index = {row["normalized_doi"]: row["id"] for row in existing if row.get("normalized_doi") and not row.get("duplicate_of")}
        title_index: dict[tuple[str, str], str] = {}
        title_blocks: dict[str, list[dict[str, Any]]] = {}
        for row in existing:
            if row.get("duplicate_of"):
                continue
            norm = row.get("normalized_title") or ""
            if norm:
                title_index[(norm, row.get("publication_year") or "")] = row["id"]
                title_blocks.setdefault(norm[:18], []).append(row)

        inserted = 0
        duplicates = 0
        candidates = 0
        for item in parsed:
            norm_title = normalise_title(item["title"])
            norm_doi = normalise_doi(item.get("doi"))
            year = _year(item.get("publication_year"))
            duplicate_of = doi_index.get(norm_doi) if norm_doi else None
            if not duplicate_of:
                duplicate_of = title_index.get((norm_title, year)) or (title_index.get((norm_title, "")) if not year else None)
            candidate_of = None
            confidence = None
            if not duplicate_of and len(norm_title) >= 24:
                pool = title_blocks.get(norm_title[:18], [])
                best = (0.0, None)
                for row in pool[:80]:
                    if year and row.get("publication_year") and year != row.get("publication_year"):
                        continue
                    ratio = SequenceMatcher(None, norm_title, row.get("normalized_title") or "").ratio()
                    if ratio > best[0]:
                        best = (ratio, row.get("id"))
                if best[0] >= 0.94:
                    candidate_of = best[1]
                    confidence = round(best[0], 4)
            record_id = str(uuid.uuid4())
            now = _utc_iso()
            conn.execute(
                """
                INSERT INTO review_records (
                    id, project_id, search_run_id, source_database, source_route,
                    external_id, title, normalized_title, abstract, authors,
                    publication_year, doi, normalized_doi, url, keywords,
                    document_type, source_title, raw_json, duplicate_of,
                    duplicate_candidate_of, duplicate_confidence, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    record_id, project_id, search_run_id, run_data["database_name"], run_data["source_route"],
                    item.get("external_id", ""), item["title"], norm_title, item.get("abstract", ""),
                    item.get("authors", ""), year, norm_doi, norm_doi, item.get("url", ""),
                    item.get("keywords", ""), item.get("document_type", ""), item.get("source_title", ""),
                    _json(item.get("raw", {})), duplicate_of, candidate_of, confidence, now, now,
                ),
            )
            inserted += 1
            if duplicate_of:
                duplicates += 1
            elif candidate_of:
                candidates += 1
            else:
                if norm_doi:
                    doi_index[norm_doi] = record_id
                title_index[(norm_title, year)] = record_id
                title_blocks.setdefault(norm_title[:18], []).append({
                    "id": record_id, "normalized_title": norm_title, "publication_year": year, "normalized_doi": norm_doi, "duplicate_of": None,
                })
        conn.execute(
            "UPDATE review_search_runs SET imported_record_count=imported_record_count+? WHERE id=?",
            (inserted, search_run_id),
        )
        conn.execute("UPDATE review_projects SET updated_at=? WHERE id=?", (_utc_iso(), project_id))
        _audit(
            conn,
            project_id,
            "records_imported",
            {"search_run_id": search_run_id, "filename": filename, "inserted": inserted, "exact_duplicates": duplicates, "possible_duplicates": candidates},
        )
    return {
        "filename": filename,
        "records_imported": inserted,
        "exact_duplicates": duplicates,
        "possible_duplicates": candidates,
        "summary": calculate_summary(project_id),
    }


def _record_where(stage: str, decision: str, search: str) -> tuple[str, list[Any]]:
    clauses = ["project_id=?"]
    params: list[Any] = []
    if stage == "possible_duplicates":
        clauses.append("duplicate_of IS NULL AND duplicate_candidate_of IS NOT NULL")
    elif stage == "title_abstract":
        clauses.append("duplicate_of IS NULL AND title_abstract_decision IN ('not_screened','uncertain')")
    elif stage == "full_text":
        clauses.append("duplicate_of IS NULL AND title_abstract_decision='include' AND full_text_decision IN ('not_assessed','not_retrieved')")
    elif stage == "included":
        clauses.append("duplicate_of IS NULL AND full_text_decision='include'")
    elif stage == "excluded":
        clauses.append("duplicate_of IS NULL AND (title_abstract_decision='exclude' OR full_text_decision='exclude')")
    elif stage == "duplicates":
        clauses.append("duplicate_of IS NOT NULL")
    if decision:
        clauses.append("(title_abstract_decision=? OR full_text_decision=?)")
        params.extend([decision, decision])
    if search:
        like = f"%{search.lower()}%"
        clauses.append("(LOWER(title) LIKE ? OR LOWER(abstract) LIKE ? OR LOWER(authors) LIKE ? OR LOWER(doi) LIKE ?)")
        params.extend([like, like, like, like])
    return " AND ".join(clauses), params


def list_records(
    project_id: str,
    *,
    stage: str = "all",
    decision: str = "",
    search: str = "",
    limit: int = 100,
    offset: int = 0,
) -> dict[str, Any]:
    limit = max(1, min(int(limit), 500))
    offset = max(0, int(offset))
    where, params = _record_where(stage, decision, search)
    params = [project_id] + params
    with connection() as conn:
        total = conn.execute(f"SELECT COUNT(*) FROM review_records WHERE {where}", params).fetchone()[0]
        rows = conn.execute(
            f"""
            SELECT id, project_id, search_run_id, source_database, source_route, external_id,
                   title, abstract, authors, publication_year, doi, url, keywords,
                   document_type, source_title, duplicate_of, duplicate_candidate_of,
                   duplicate_confidence, title_abstract_decision, title_abstract_reason,
                   full_text_decision, full_text_reason, reviewer_notes, full_text_filename,
                   CASE WHEN full_text_text <> '' THEN 1 ELSE 0 END AS has_full_text,
                   created_at, updated_at
            FROM review_records WHERE {where}
            ORDER BY created_at DESC LIMIT ? OFFSET ?
            """,
            params + [limit, offset],
        ).fetchall()
    return {"records": [_row(row) or {} for row in rows], "total": total, "limit": limit, "offset": offset}


def get_record(project_id: str, record_id: str) -> dict[str, Any]:
    with connection() as conn:
        row = conn.execute(
            "SELECT * FROM review_records WHERE id=? AND project_id=?",
            (record_id, project_id),
        ).fetchone()
    record = _row(row)
    if not record:
        raise ValueError("The review record could not be found.")
    text = str(record.get("full_text_text") or "")
    record["full_text_excerpt"] = text[:15000]
    record.pop("full_text_text", None)
    return record


def update_record(project_id: str, record_id: str, data: dict[str, Any]) -> dict[str, Any]:
    allowed = {
        "title_abstract_decision", "title_abstract_reason", "full_text_decision",
        "full_text_reason", "reviewer_notes", "url",
    }
    updates = {key: _clean(value) for key, value in data.items() if key in allowed and value is not None}
    ta_allowed = {"not_screened", "include", "exclude", "uncertain"}
    ft_allowed = {"not_assessed", "include", "exclude", "not_retrieved"}
    if "title_abstract_decision" in updates and updates["title_abstract_decision"] not in ta_allowed:
        raise ValueError("Select a valid title and abstract screening decision.")
    if "full_text_decision" in updates and updates["full_text_decision"] not in ft_allowed:
        raise ValueError("Select a valid full-text screening decision.")
    if updates.get("title_abstract_decision") == "exclude" and not updates.get("title_abstract_reason"):
        raise ValueError("Enter the title and abstract exclusion reason.")
    if updates.get("full_text_decision") in {"exclude", "not_retrieved"} and not updates.get("full_text_reason"):
        raise ValueError("Enter the full-text exclusion or retrieval reason.")
    if not updates:
        return get_record(project_id, record_id)
    with connection() as conn:
        exists = conn.execute("SELECT id FROM review_records WHERE id=? AND project_id=?", (record_id, project_id)).fetchone()
        if not exists:
            raise ValueError("The review record could not be found.")
        columns = ", ".join(f"{key}=?" for key in updates)
        conn.execute(f"UPDATE review_records SET {columns}, updated_at=? WHERE id=? AND project_id=?", list(updates.values()) + [_utc_iso(), record_id, project_id])
        conn.execute("UPDATE review_projects SET updated_at=? WHERE id=?", (_utc_iso(), project_id))
        _audit(conn, project_id, "record_screened", updates, record_id)
    return get_record(project_id, record_id)


def bulk_update_records(project_id: str, record_ids: list[str], stage: str, decision: str, reason: str = "", notes: str = "") -> dict[str, Any]:
    ids = [str(item) for item in record_ids if str(item).strip()][:500]
    if not ids:
        raise ValueError("Select at least one record.")
    if stage == "title_abstract":
        if decision not in {"include", "exclude", "uncertain", "not_screened"}:
            raise ValueError("Select a valid title and abstract decision.")
        if decision == "exclude" and not _clean(reason):
            raise ValueError("Enter the exclusion reason.")
        fields = {"title_abstract_decision": decision, "title_abstract_reason": _clean(reason)}
    elif stage == "full_text":
        if decision not in {"include", "exclude", "not_retrieved", "not_assessed"}:
            raise ValueError("Select a valid full-text decision.")
        if decision in {"exclude", "not_retrieved"} and not _clean(reason):
            raise ValueError("Enter the exclusion or retrieval reason.")
        fields = {"full_text_decision": decision, "full_text_reason": _clean(reason)}
    else:
        raise ValueError("Select title/abstract or full-text screening.")
    if notes:
        fields["reviewer_notes"] = _clean(notes)
    placeholders = ",".join("?" for _ in ids)
    with connection() as conn:
        columns = ", ".join(f"{key}=?" for key in fields)
        params = list(fields.values()) + [_utc_iso(), project_id] + ids
        cursor = conn.execute(
            f"UPDATE review_records SET {columns}, updated_at=? WHERE project_id=? AND duplicate_of IS NULL AND id IN ({placeholders})",
            params,
        )
        _audit(conn, project_id, "bulk_screening", {"record_ids": ids, "stage": stage, "decision": decision, "reason": reason, "updated": cursor.rowcount})
    return {"updated": cursor.rowcount, "summary": calculate_summary(project_id)}


def resolve_duplicate(project_id: str, record_id: str, action: str, duplicate_of: str = "") -> dict[str, Any]:
    if action not in {"confirm", "keep_unique", "clear"}:
        raise ValueError("Select a valid duplicate decision.")
    with connection() as conn:
        record = conn.execute("SELECT * FROM review_records WHERE id=? AND project_id=?", (record_id, project_id)).fetchone()
        if not record:
            raise ValueError("The review record could not be found.")
        record_data = dict(record)
        if action == "confirm":
            target = duplicate_of or record_data.get("duplicate_candidate_of") or ""
            if not target:
                raise ValueError("Select the retained record for this duplicate.")
            target_row = conn.execute("SELECT id FROM review_records WHERE id=? AND project_id=?", (target, project_id)).fetchone()
            if not target_row or target == record_id:
                raise ValueError("The retained record is invalid.")
            conn.execute(
                "UPDATE review_records SET duplicate_of=?, duplicate_candidate_of=NULL, duplicate_confidence=NULL, updated_at=? WHERE id=?",
                (target, _utc_iso(), record_id),
            )
        elif action == "keep_unique":
            conn.execute(
                "UPDATE review_records SET duplicate_of=NULL, duplicate_candidate_of=NULL, duplicate_confidence=NULL, updated_at=? WHERE id=?",
                (_utc_iso(), record_id),
            )
        else:
            conn.execute(
                "UPDATE review_records SET duplicate_of=NULL, updated_at=? WHERE id=?",
                (_utc_iso(), record_id),
            )
        _audit(conn, project_id, "duplicate_resolved", {"record_id": record_id, "action": action, "duplicate_of": duplicate_of}, record_id)
    return get_record(project_id, record_id)


def attach_full_text(project_id: str, record_id: str, filename: str, extracted_text: str) -> dict[str, Any]:
    text = str(extracted_text or "").strip()
    if not text:
        raise ValueError("No readable text was extracted from the full-text file.")
    text = text[:600_000]
    with connection() as conn:
        exists = conn.execute("SELECT id FROM review_records WHERE id=? AND project_id=?", (record_id, project_id)).fetchone()
        if not exists:
            raise ValueError("The review record could not be found.")
        conn.execute(
            "UPDATE review_records SET full_text_filename=?, full_text_text=?, updated_at=? WHERE id=?",
            (_clean(filename), text, _utc_iso(), record_id),
        )
        _audit(conn, project_id, "full_text_attached", {"filename": filename, "characters": len(text)}, record_id)
    return get_record(project_id, record_id)


def calculate_summary(project_id: str) -> dict[str, Any]:
    with connection() as conn:
        row = conn.execute(
            """
            SELECT
              COUNT(*) AS records_identified,
              SUM(CASE WHEN source_route='database' THEN 1 ELSE 0 END) AS database_records_identified,
              SUM(CASE WHEN source_route<>'database' THEN 1 ELSE 0 END) AS other_records_identified,
              SUM(CASE WHEN duplicate_of IS NOT NULL THEN 1 ELSE 0 END) AS duplicates_removed,
              SUM(CASE WHEN duplicate_of IS NULL AND duplicate_candidate_of IS NOT NULL THEN 1 ELSE 0 END) AS possible_duplicates,
              SUM(CASE WHEN duplicate_of IS NULL THEN 1 ELSE 0 END) AS unique_records,
              SUM(CASE WHEN duplicate_of IS NULL AND title_abstract_decision IN ('include','exclude') THEN 1 ELSE 0 END) AS records_screened,
              SUM(CASE WHEN duplicate_of IS NULL AND title_abstract_decision='exclude' THEN 1 ELSE 0 END) AS records_excluded,
              SUM(CASE WHEN duplicate_of IS NULL AND title_abstract_decision='include' THEN 1 ELSE 0 END) AS title_abstract_included,
              SUM(CASE WHEN duplicate_of IS NULL AND title_abstract_decision='include' AND full_text_decision IN ('include','exclude') THEN 1 ELSE 0 END) AS full_text_assessed,
              SUM(CASE WHEN duplicate_of IS NULL AND full_text_decision='not_retrieved' THEN 1 ELSE 0 END) AS reports_not_retrieved,
              SUM(CASE WHEN duplicate_of IS NULL AND full_text_decision='exclude' THEN 1 ELSE 0 END) AS full_text_excluded,
              SUM(CASE WHEN duplicate_of IS NULL AND full_text_decision='include' THEN 1 ELSE 0 END) AS final_corpus,
              SUM(CASE WHEN duplicate_of IS NULL AND full_text_decision='include' AND source_route IN ('backward_citation','forward_citation') THEN 1 ELSE 0 END) AS included_citation_tracking,
              SUM(CASE WHEN duplicate_of IS NULL AND title_abstract_decision IN ('not_screened','uncertain') THEN 1 ELSE 0 END) AS awaiting_title_abstract,
              SUM(CASE WHEN duplicate_of IS NULL AND title_abstract_decision='include' AND full_text_decision='not_assessed' THEN 1 ELSE 0 END) AS awaiting_full_text
            FROM review_records WHERE project_id=?
            """,
            (project_id,),
        ).fetchone()
        search_count = conn.execute("SELECT COUNT(*) FROM review_search_runs WHERE project_id=?", (project_id,)).fetchone()[0]
    data = {key: int(value or 0) for key, value in dict(row).items()}
    data["search_runs"] = int(search_count)
    data["screening_complete"] = data["awaiting_title_abstract"] == 0 and data["awaiting_full_text"] == 0 and data["possible_duplicates"] == 0 and data["records_identified"] > 0
    warnings: list[str] = []
    for run in list_search_runs(project_id):
        reported = run.get("reported_result_count")
        imported = int(run.get("imported_record_count") or 0)
        if reported is not None and int(reported) != imported:
            warnings.append(
                f"{run.get('database_name') or 'A search run'} reports {int(reported)} result(s), but {imported} record(s) were imported. Confirm whether the export was partial."
            )
        if run.get("source_route") == "database" and not str(run.get("search_string") or "").strip():
            warnings.append(f"The complete search string is missing for {run.get('database_name') or 'a database search' }.")
    if data["possible_duplicates"]:
        warnings.append(f"{data['possible_duplicates']} possible duplicate(s) require confirmation.")
    if data["awaiting_title_abstract"]:
        warnings.append(f"{data['awaiting_title_abstract']} unique record(s) await title and abstract screening.")
    if data["awaiting_full_text"]:
        warnings.append(f"{data['awaiting_full_text']} record(s) retained at initial screening await full-text assessment.")
    data["warnings"] = warnings
    return data


def _record_rows(project_id: str, scope: str = "all") -> list[dict[str, Any]]:
    clause = "project_id=?"
    if scope == "included":
        clause += " AND duplicate_of IS NULL AND full_text_decision='include'"
    elif scope == "excluded":
        clause += " AND duplicate_of IS NULL AND (title_abstract_decision='exclude' OR full_text_decision IN ('exclude','not_retrieved'))"
    elif scope == "duplicates":
        clause += " AND duplicate_of IS NOT NULL"
    with connection() as conn:
        rows = conn.execute(
            f"SELECT * FROM review_records WHERE {clause} ORDER BY source_database, title",
            (project_id,),
        ).fetchall()
    output = []
    for row in rows:
        item = dict(row)
        item.pop("raw_json", None)
        item.pop("full_text_text", None)
        output.append(item)
    return output


def export_records_csv(project_id: str, scope: str = "all") -> tuple[io.BytesIO, str]:
    rows = _record_rows(project_id, scope)
    fields = [
        "id", "source_database", "source_route", "external_id", "title", "authors",
        "publication_year", "doi", "url", "keywords", "document_type", "source_title",
        "duplicate_of", "duplicate_candidate_of", "duplicate_confidence",
        "title_abstract_decision", "title_abstract_reason", "full_text_decision",
        "full_text_reason", "reviewer_notes", "full_text_filename",
    ]
    buffer = io.StringIO()
    writer = csv.DictWriter(buffer, fieldnames=fields, extrasaction="ignore")
    writer.writeheader()
    writer.writerows(rows)
    data = io.BytesIO(buffer.getvalue().encode("utf-8-sig"))
    data.seek(0)
    return data, f"review_evidence_{scope}.csv"


def export_audit_json(project_id: str) -> tuple[io.BytesIO, str]:
    with connection() as conn:
        project = _row(conn.execute("SELECT * FROM review_projects WHERE id=?", (project_id,)).fetchone()) or {}
        project.pop("token_hash", None)
        events = [_row(row) or {} for row in conn.execute("SELECT * FROM review_audit_events WHERE project_id=? ORDER BY created_at", (project_id,)).fetchall()]
    payload = {
        "project": project,
        "summary": calculate_summary(project_id),
        "search_runs": list_search_runs(project_id),
        "records": _record_rows(project_id, "all"),
        "audit_events": events,
        "exported_at": _utc_iso(),
    }
    stream = io.BytesIO(json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8"))
    stream.seek(0)
    return stream, "review_evidence_audit.json"


def _methodology_lines(project: dict[str, Any], search_runs: list[dict[str, Any]], summary: dict[str, Any]) -> list[str]:
    databases = []
    search_strings = []
    for run in reversed(search_runs):
        label = run.get("database_name", "")
        platform = run.get("platform", "")
        databases.append(f"{label} ({platform})" if platform else label)
        if run.get("search_string"):
            search_strings.append(f"{label}: {run['search_string']}")
    lines = [
        f"Review positioning: {project.get('protocol_positioning') or project.get('article_type') or 'Not specified'}",
        f"Review question or objective: {project.get('review_question') or '[Author action: Confirm the review question or objective.]'}",
        f"Databases and routes documented: {'; '.join(databases) or '[Author action: Add formal database search runs.]'}",
        "Search strings:\n" + ("\n".join(search_strings) if search_strings else "[Author action: Add the complete database-specific search strings.]"),
        f"Eligibility criteria: {project.get('eligibility_criteria') or '[Author action: Confirm inclusion and exclusion criteria.]'}",
        f"Screening process: {project.get('screening_process') or '[Author action: Confirm reviewer roles, screening stages and disagreement resolution.]'}",
        f"Quality appraisal: {project.get('quality_appraisal') or '[Author action: Confirm the appraisal or critical-evaluation approach.]'}",
        f"Synthesis method: {project.get('synthesis_method') or '[Author action: Confirm extraction, coding and synthesis procedures.]'}",
        f"Software: {project.get('software') or '[Author action: Confirm software, versions and settings.]'}",
        f"Records identified: {summary['records_identified']}",
        f"Duplicates removed: {summary['duplicates_removed']}",
        f"Records screened: {summary['records_screened']}",
        f"Records excluded at title and abstract stage: {summary['records_excluded']}",
        f"Full texts assessed: {summary['full_text_assessed']}",
        f"Full texts excluded: {summary['full_text_excluded']}",
        f"Reports not retrieved: {summary['reports_not_retrieved']}",
        f"Final included corpus: {summary['final_corpus']}",
    ]
    return lines


def writer_payload(project_id: str) -> dict[str, Any]:
    with connection() as conn:
        project = _row(conn.execute("SELECT * FROM review_projects WHERE id=?", (project_id,)).fetchone()) or {}
    project.pop("token_hash", None)
    runs = list_search_runs(project_id)
    summary = calculate_summary(project_id)
    database_lines = []
    query_lines = []
    dates = []
    date_limits = []
    languages = []
    doc_types = []
    citation_routes = []
    duplicate_process = (
        "Records were imported into the ArticleReady Review Evidence Workspace. Exact DOI matches and exact normalised title-year matches were removed automatically. "
        "High-similarity title matches were flagged for manual confirmation, and all duplicate decisions were retained in the audit ledger."
    )
    for run in reversed(runs):
        label = run.get("database_name", "")
        platform = run.get("platform", "")
        database_lines.append(f"{label} ({platform})" if platform else label)
        if run.get("search_string"):
            query_lines.append(f"{label}: {run['search_string']}")
        if run.get("search_date"):
            dates.append(str(run["search_date"]))
        if run.get("date_limits"):
            date_limits.append(str(run["date_limits"]))
        if run.get("language_limits"):
            languages.append(str(run["language_limits"]))
        if run.get("document_types"):
            doc_types.append(str(run["document_types"]))
        if run.get("source_route") in {"backward_citation", "forward_citation"}:
            citation_routes.append(f"{run.get('source_route', '').replace('_', ' ')}: {label}")
    notes = _clean(project.get("notes"))
    if summary["included_citation_tracking"]:
        notes = f"{notes} {summary['included_citation_tracking']} included record(s) entered through backward or forward citation tracking.".strip()
    return {
        "workspace_project_id": project_id,
        "workspace_title": project.get("title", ""),
        "article_type": project.get("article_type", "Systematic review"),
        "review_protocol_positioning": project.get("protocol_positioning", "Auto"),
        "review_databases": "; ".join(item for item in database_lines if item),
        "review_search_strings": "\n".join(query_lines),
        "review_search_date": max(dates) if dates else "",
        "review_date_limits": "; ".join(dict.fromkeys(date_limits)),
        "review_language_limits": "; ".join(dict.fromkeys(languages)),
        "review_document_types": "; ".join(dict.fromkeys(doc_types)),
        "review_eligibility_criteria": project.get("eligibility_criteria", ""),
        "review_screening_process": project.get("screening_process", ""),
        "review_quality_appraisal": project.get("quality_appraisal", ""),
        "review_citation_tracking": "; ".join(citation_routes) or ("No separate citation-tracking run has been documented." if runs else ""),
        "review_duplicate_removal": duplicate_process,
        "review_synthesis_method": project.get("synthesis_method", ""),
        "review_software": project.get("software", ""),
        "review_protocol_notes": notes,
        "review_records_identified": summary["records_identified"],
        "review_duplicates_removed": summary["duplicates_removed"],
        "review_records_screened": summary["records_screened"],
        "review_records_excluded": summary["records_excluded"],
        "review_full_text_assessed": summary["full_text_assessed"],
        "review_full_text_excluded": summary["full_text_excluded"],
        "review_citation_tracking_additions": 0,
        "review_final_corpus_size": summary["final_corpus"],
        "research_problem": project.get("review_question", ""),
        "methodology": "\n".join(_methodology_lines(project, runs, summary)),
        "workspace_summary": summary,
    }


def export_protocol_docx(project_id: str) -> tuple[io.BytesIO, str]:
    with connection() as conn:
        project = _row(conn.execute("SELECT * FROM review_projects WHERE id=?", (project_id,)).fetchone()) or {}
    project.pop("token_hash", None)
    runs = list_search_runs(project_id)
    summary = calculate_summary(project_id)
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    document.add_heading(project.get("title") or "Review Evidence Workspace", 0)
    document.add_paragraph("Review protocol, record-flow and evidence audit generated from the stored workspace ledger.")

    document.add_heading("1. Review design", level=1)
    for line in _methodology_lines(project, runs, summary)[:9]:
        document.add_paragraph(line)

    document.add_heading("2. Search-run register", level=1)
    if runs:
        table = document.add_table(rows=1, cols=7)
        headers = ["Database/source", "Route", "Search date", "Reported", "Imported", "Limits", "Search string"]
        for idx, value in enumerate(headers):
            table.rows[0].cells[idx].text = value
        for run in reversed(runs):
            cells = table.add_row().cells
            cells[0].text = str(run.get("database_name") or "")
            cells[1].text = str(run.get("source_route") or "").replace("_", " ")
            cells[2].text = str(run.get("search_date") or "")
            cells[3].text = "" if run.get("reported_result_count") is None else str(run.get("reported_result_count"))
            cells[4].text = str(run.get("imported_record_count") or 0)
            cells[5].text = "; ".join(filter(None, [run.get("date_limits"), run.get("language_limits"), run.get("document_types")]))
            cells[6].text = str(run.get("search_string") or "")
    else:
        document.add_paragraph("[Author action: Add at least one formal search or citation-tracking run.]")

    document.add_heading("3. Record-flow summary", level=1)
    flow_table = document.add_table(rows=1, cols=2)
    flow_table.rows[0].cells[0].text = "Stage"
    flow_table.rows[0].cells[1].text = "Verified count"
    flow_items = [
        ("Records identified", "records_identified"),
        ("Database records identified", "database_records_identified"),
        ("Records identified through other methods", "other_records_identified"),
        ("Duplicates removed", "duplicates_removed"),
        ("Possible duplicates unresolved", "possible_duplicates"),
        ("Records screened", "records_screened"),
        ("Records excluded at title/abstract", "records_excluded"),
        ("Full texts assessed", "full_text_assessed"),
        ("Reports not retrieved", "reports_not_retrieved"),
        ("Full texts excluded", "full_text_excluded"),
        ("Final included corpus", "final_corpus"),
    ]
    for label, key in flow_items:
        cells = flow_table.add_row().cells
        cells[0].text = label
        cells[1].text = str(summary.get(key, 0))

    document.add_heading("4. Screening audit", level=1)
    if summary["warnings"]:
        for warning in summary["warnings"]:
            document.add_paragraph(warning, style="List Bullet")
    else:
        document.add_paragraph("No unresolved duplicate or screening-stage warnings were detected in the current ledger.")

    document.add_heading("5. Included-study corpus", level=1)
    included = _record_rows(project_id, "included")
    if included:
        corpus = document.add_table(rows=1, cols=5)
        for idx, value in enumerate(["Author(s)", "Year", "Title", "Source", "DOI"]):
            corpus.rows[0].cells[idx].text = value
        for record in included:
            cells = corpus.add_row().cells
            cells[0].text = str(record.get("authors") or "")
            cells[1].text = str(record.get("publication_year") or "")
            cells[2].text = str(record.get("title") or "")
            cells[3].text = str(record.get("source_title") or record.get("source_database") or "")
            cells[4].text = str(record.get("doi") or "")
    else:
        document.add_paragraph("[Author action: Complete full-text assessment before presenting a final included-study corpus.]")

    document.add_heading("6. Evidence-integrity statement", level=1)
    document.add_paragraph(
        "The counts in this document are calculated from imported records and confirmed screening decisions in the Review Evidence Workspace. "
        "ArticleReady metadata discovery is not treated as a formal database search unless the corresponding export and search-run details are recorded in this workspace."
    )
    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    safe = re.sub(r"[^A-Za-z0-9_-]+", "_", str(project.get("title") or "review_evidence"))[:80].strip("_") or "review_evidence"
    return stream, f"{safe}_protocol_and_evidence_audit.docx"
