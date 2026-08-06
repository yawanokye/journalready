from __future__ import annotations

import difflib
import io
import json
import logging
import os
import re
from datetime import datetime
from typing import Any

from app.article_service import (
    _article_prompt_quality_pack,
    _article_reference_expectations,
    _call_openai_response_with_fallback,
    _citation_density_report,
    _citation_density_requirements,
    _finalise_article_text,
    _humanisation_strength,
    _humanize_article_with_model,
    _humanizer_mode,
    _normalise_article_model,
    _normalise_reasoning_effort,
    _parse_inline_segments,
    _plain_inline_text,
    _safe_get_openai_client,
    _strong_humanisation_enabled,
    _strong_humanisation_requirements,
    _search_sources,
    _source_context,
)
from app.scholarly_humanizer import analyse_scholarly_style, humanize_scholarly_text

LOGGER = logging.getLogger("articleready.revision")

_REVISION_HEADING_RE = re.compile(
    r"^(?:"
    r"#{1,6}\s+.+|"
    r"\d+(?:\.\d+){0,4}\.?\s+[A-Z][^\n]{1,180}|"
    r"(?:Abstract|Introduction|Background|Literature Review|Theoretical Framework|Conceptual Framework|"
    r"Methods?|Methodology|Materials and Methods|Results?|Findings?|Discussion|Conclusion|Conclusions|"
    r"Recommendations?|Implications?|Limitations?|Declarations?|Acknowledgements?|Funding|"
    r"Competing Interests?|Conflict of Interest|Author Contributions?|Data Availability|Ethics Approval|"
    r"References|Bibliography|Source Use Audit|Appendix|Appendices)"
    r")$",
    flags=re.IGNORECASE,
)
_REVISION_PROTECTED_HEADING_RE = re.compile(
    r"^(?:#{1,6}\s*)?(?:References|Bibliography|Source Use Audit|Appendix|Appendices)\b",
    flags=re.IGNORECASE,
)

_REVISION_BLUE = (0, 112, 192)
_ACTION_RED = (192, 0, 0)


class RevisionServiceUnavailable(RuntimeError):
    """Raised when no substantive paid revision was produced.

    Propagating this exception causes the payment entitlement claim to roll back,
    so a temporary provider failure does not consume the customer's revision.
    """

    def __init__(self, message: str, *, provider_notes: list[Any] | None = None) -> None:
        super().__init__(message)
        self.provider_notes = list(provider_notes or [])


def _revision_model() -> str:
    """Model used for final scholarly rewriting."""
    return _normalise_article_model(
        os.getenv("OPENAI_ARTICLE_REVISION_MODEL")
        or os.getenv("OPENAI_ARTICLE_WRITING_MODEL")
        or os.getenv("OPENAI_ARTICLE_ADVANCED_MODEL")
        or os.getenv("OPENAI_ARTICLE_RESEARCH_MODEL")
        or "",
        "gpt-5.5",
    ) or "gpt-5.5"


def _revision_recovery_model() -> str:
    """Lower-cost analytical recovery model used after a GPT-5.5 failure."""
    return _normalise_article_model(
        os.getenv("OPENAI_ARTICLE_REVISION_RECOVERY_MODEL")
        or os.getenv("OPENAI_ARTICLE_ANALYSIS_MODEL")
        or os.getenv("OPENAI_ARTICLE_TERRA_MODEL")
        or "",
        "gpt-5.6-terra",
    ) or "gpt-5.6-terra"


def _analysis_model() -> str:
    return _normalise_article_model(
        os.getenv("OPENAI_ARTICLE_ANALYSIS_MODEL")
        or os.getenv("OPENAI_ARTICLE_TERRA_MODEL")
        or "",
        "gpt-5.6-terra",
    ) or "gpt-5.6-terra"


def _extract_response_text(response: Any) -> str:
    return str(getattr(response, "output_text", "") or "").strip()


def _strip_code_fences(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"^```(?:markdown|md)?\s*|\s*```$", "", text.strip(), flags=re.IGNORECASE | re.MULTILINE).strip()


def _split_revision_package(text: str) -> tuple[str, str, str]:
    raw = _strip_code_fences(text)
    revised_marker = "===REVISED_ARTICLE==="
    report_marker = "===REVISION_REPORT==="
    matrix_marker = "===REVIEWER_RESPONSE_MATRIX==="

    if revised_marker in raw:
        raw = raw.split(revised_marker, 1)[1]
    matrix = ""
    if matrix_marker in raw:
        raw, matrix = raw.split(matrix_marker, 1)
    report = ""
    if report_marker in raw:
        revised, report = raw.split(report_marker, 1)
    else:
        revised = raw
    return revised.strip(), report.strip(), matrix.strip()


def _revision_focus(payload: dict[str, Any]) -> list[str]:
    focus_map = [
        ("strengthen_conceptualisation", "conceptualisation, theoretical framing and the logic connecting the constructs or phenomena"),
        ("strengthen_contribution", "theoretical, empirical, methodological, contextual and practical contribution"),
        ("assess_method_fit", "alignment among the problem, objectives, design, sampling, measurement, identification and validity strategy"),
        ("assess_analysis", "appropriateness, assumptions, robustness and reporting of the analysis"),
        ("deepen_discussion", "interpretive depth, comparison with literature, boundary conditions, alternative explanations and limitations"),
        ("strengthen_recommendations", "evidence-linked implications and feasible recommendations for named actors"),
    ]
    return [description for key, description in focus_map if bool(payload.get(key, True))]


def _revision_human_writing_layer(payload: dict[str, Any]) -> dict[str, Any]:
    """Extend the article writer's human-supervised academic style controls for revision work.

    The shared quality pack keeps article drafting and article revision aligned.
    These additional rules are revision-specific and preserve the author's evidence,
    citations and substantive voice while improving scholarly flow.
    """
    base_pack = _article_prompt_quality_pack(payload)
    return {
        "base_article_quality_pack": base_pack,
        "strong_humanisation_requirements": _strong_humanisation_requirements(payload),
        "revision_specific_rules": [
            "Preserve the author's substantive position and disciplinary voice. Refine expression and argument without replacing the study with a generic manuscript.",
            "Use controlled high burstiness and strong lexical variation according to the complexity of the argument. Do not force every paragraph into the same rhythm or size.",
            "Combine concise emphasis with developed analytical sentences, but do not create fragments or artificial irregularities.",
            "Build substantive paragraphs around claim, evidence, interpretation, limitation and contribution where appropriate, rather than using repetitive topic-sentence templates.",
            "Vary paragraph openings and transitions. Avoid repeated constructions such as 'This study', 'The findings show', 'Furthermore' and 'Moreover' when a more direct transition is available.",
            "Prefer precise, grounded verbs such as indicates, qualifies, complicates, supports, constrains, extends, contradicts, explains and suggests.",
            "Remove generic filler, inflated novelty claims, mechanical summaries and author-by-author literature listing.",
            "Keep conceptual definitions, mechanisms, boundary conditions and contribution claims specific to the supplied study and evidence.",
            "Preserve valid in-text citations exactly where possible. Do not invent citations, alter author names or years, or detach a citation from the claim it supports.",
            "Do not introduce artificial mistakes, awkward grammar, random informality or deliberate imperfections.",
            "Do not add commentary about AI detection, humanisation, detector scores or making the article appear human. Produce normal scholarly prose only.",
            "Use formal British English, minimise long dashes and keep headings, tables, equations and citations clean.",
        ],
        "post_processing": (
            "After the model revision, the shared deterministic strong humanisation pass from article_service.py "
            "varies sentence rhythm, paragraph openings and repetitive wording while protecting evidence, citations, "
            "tables, equations, placeholders, references and document structure."
        ),
    }


def _fallback_revision_report(payload: dict[str, Any], source_records: list[dict[str, Any]], provider_errors: list[str]) -> str:
    comments = str(payload.get("review_comments") or "").strip()
    source_note = (
        f"{len(source_records)} scholarly record(s) were available for relevance screening, but no sources were inserted automatically in fallback mode."
        if source_records
        else "No scholarly records were available for source-supported revision."
    )
    errors = "\n".join(f"- {item}" for item in provider_errors) or "- The AI revision service was unavailable."
    reviewer_note = (
        "Reviewer comments were supplied, but they were not substantively resolved because the AI revision service was unavailable."
        if comments
        else "No reviewer comments were supplied."
    )
    return f"""# Revision and Publishability Report

## Revision Status

The original article has been returned without substantive rewriting because the revision model was unavailable. Do not treat this fallback as a completed publication-readiness revision.

## Provider or Processing Notes

{errors}

## Priority Review Areas

1. Clarify the article-level problem and the precise conceptual or theoretical gap.
2. State the central contribution in a way that distinguishes it from contextual novelty alone.
3. Confirm that the research design, sampling, measures and analysis directly answer the stated objectives.
4. Check whether robustness, diagnostic, endogeneity, sensitivity, validity or trustworthiness procedures are needed.
5. Rebuild the Discussion around the main findings, mechanisms, competing explanations, boundaries and contribution.
6. Ensure recommendations follow directly from confirmed findings and name the responsible actors.

## Additional Analysis

[review the supplied method and results to determine whether additional analysis is essential, strongly recommended or optional. Do not report any analysis as completed until actual output is supplied]

## Reviewer Comments

{reviewer_note}

## Source Review

{source_note}
""".strip()


def _fallback_reviewer_matrix(review_comments: str) -> str:
    comments = [line.strip(" -\t") for line in str(review_comments or "").splitlines() if line.strip()]
    if not comments:
        return ""
    rows = ["| Reviewer comment | Revision made | Location | Remaining action |", "|---|---|---|---|"]
    for comment in comments[:30]:
        safe = comment.replace("|", "/")
        rows.append(f"| {safe} | [substantive revision not completed in fallback mode] | [identify section] | [action required] |")
    return "\n".join(rows)


def _allow_non_substantive_fallback() -> bool:
    return os.getenv("ARTICLEREADY_ALLOW_REVISION_FALLBACK", "0").strip().lower() in {"1", "true", "yes", "on"}


def _revision_changed(original: str, revised: str) -> bool:
    original_clean = re.sub(r"\s+", " ", str(original or "")).strip()
    revised_clean = re.sub(r"\s+", " ", str(revised or "")).strip()
    if not revised_clean or revised_clean == original_clean:
        return False
    if len(revised_clean) < max(100, int(len(original_clean) * 0.45)):
        return False
    return True


def _public_revision_failure(provider_errors: list[Any]) -> RevisionServiceUnavailable:
    processing_notes: list[str] = []
    source_warnings: list[str] = []
    for item in provider_errors[-12:]:
        if isinstance(item, dict):
            provider = str(item.get("provider") or "provider")
            error = str(item.get("error") or "temporary error")
            source_warnings.append(f"Source warning, non-blocking: {provider}: {error[:170]}")
        else:
            processing_notes.append(str(item)[:240])
    notes = processing_notes[-6:] + source_warnings[-3:]
    return RevisionServiceUnavailable(
        "The revision service could not complete and validate a substantive full-manuscript revision. "
        "No revision entitlement has been consumed. Retry the request and review the processing notes or Render logs for the exact failed stage.",
        provider_notes=notes,
    )



def _env_int(name: str, default: int, *, minimum: int, maximum: int) -> int:
    try:
        value = int(str(os.getenv(name, str(default)) or default).strip())
    except (TypeError, ValueError):
        value = default
    return max(minimum, min(value, maximum))


def _word_count(text: str) -> int:
    return len(re.findall(r"\b[\w’'-]+\b", str(text or ""), flags=re.UNICODE))


def _clip_text(text: str, limit: int) -> str:
    value = str(text or "")
    if len(value) <= limit:
        return value
    return value[:limit].rstrip() + "\n\n[Content clipped for this support pass. The substantive revision still uses the complete section batches.]"


def _revision_timeout_seconds() -> int:
    return _env_int("ARTICLEREADY_REVISION_TIMEOUT_SECONDS", 600, minimum=120, maximum=1800)


def _revision_reasoning() -> str:
    return _normalise_reasoning_effort(os.getenv("OPENAI_ARTICLE_REVISION_REASONING", "high"), "high")


def _revision_recovery_reasoning() -> str:
    return _normalise_reasoning_effort(os.getenv("OPENAI_ARTICLE_REVISION_RECOVERY_REASONING", "high"), "high")


def _revision_plan_model() -> str:
    return _normalise_article_model(
        os.getenv("OPENAI_ARTICLE_REVISION_PLAN_MODEL")
        or os.getenv("OPENAI_ARTICLE_FAST_MODEL")
        or os.getenv("OPENAI_ARTICLE_LUNA_MODEL")
        or "",
        "gpt-5.6-luna",
    ) or "gpt-5.6-luna"


def _revision_escalation_model() -> str:
    return _normalise_article_model(
        os.getenv("OPENAI_ARTICLE_ESCALATION_MODEL")
        or os.getenv("OPENAI_ARTICLE_SOL_MODEL")
        or "",
        "gpt-5.6-sol",
    ) or "gpt-5.6-sol"


def _build_author_voice_profile(article_text: str, payload: dict[str, Any]) -> dict[str, Any]:
    """Build one global voice profile and reuse it across every revision batch."""
    report = analyse_scholarly_style(article_text)
    return {
        "language": "formal British English",
        "disciplinary_area": str(payload.get("research_area") or "").strip(),
        "target_journal": str(payload.get("target_journal") or "").strip(),
        "original_naturalness_score": report.get("naturalness_score"),
        "original_sentence_length_cv": report.get("sentence_length_cv"),
        "original_paragraph_length_cv": report.get("paragraph_length_cv"),
        "original_lexical_diversity": report.get("lexical_diversity_msttr"),
        "original_repeated_sentence_openings": report.get("repeated_sentence_openings"),
        "original_repeated_paragraph_openings": report.get("repeated_paragraph_openings"),
        "voice_directives": [
            "Preserve the author's disciplinary register, caution level and field-specific terminology.",
            "Prefer analytical, evidence-led prose to promotional or formulaic prose.",
            "Vary sentence and paragraph length only where the argument naturally requires it.",
            "Do not force every paragraph into the same claim-evidence-conclusion pattern.",
            "Avoid generic transitions, repeated summary sentences and three-part lists used only for rhythm.",
            "Keep causal, correlational and interpretive claims at the strength supported by the evidence.",
        ],
    }


def _context_head(text: str, limit: int = 1200) -> str:
    value = re.sub(r"\s+", " ", str(text or "")).strip()
    return value[:limit]


def _context_tail(text: str, limit: int = 1200) -> str:
    value = re.sub(r"\s+", " ", str(text or "")).strip()
    return value[-limit:]


def _section_heading(text: str, fallback: str) -> str:
    for line in str(text or "").splitlines():
        stripped = line.strip()
        if stripped:
            return stripped[:180]
    return fallback


def _split_revision_sections(article_text: str) -> list[dict[str, Any]]:
    """Split plain extracted manuscripts using common journal headings.

    DOCX extraction removes Word heading styles, so this parser recognises both
    numbered and conventional unnumbered article headings. Reference and
    appendix sections are marked as protected.
    """
    lines = str(article_text or "").splitlines()
    sections: list[dict[str, Any]] = []
    current: list[str] = []
    heading = ""

    def flush() -> None:
        nonlocal current, heading
        text = "\n".join(current).strip()
        if text:
            sections.append({
                "heading": heading,
                "text": text,
                "protected": bool(_REVISION_PROTECTED_HEADING_RE.match(heading.strip())) if heading else False,
                "word_count": _word_count(text),
            })
        current = []

    for line in lines:
        stripped = line.strip()
        if stripped and _REVISION_HEADING_RE.match(stripped):
            flush()
            heading = stripped
            current = [line]
        else:
            current.append(line)
    flush()
    return sections or [{
        "heading": "",
        "text": str(article_text or "").strip(),
        "protected": False,
        "word_count": _word_count(article_text),
    }]


def _split_long_text_unit(text: str, max_words: int) -> list[str]:
    """Split a long paragraph without cutting ordinary sentences where possible."""
    value = str(text or "").strip()
    if not value:
        return []
    if _word_count(value) <= max_words:
        return [value]

    sentences = [
        item.strip()
        for item in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\[])|\n+", value)
        if item.strip()
    ]
    if len(sentences) <= 1:
        words = value.split()
        return [" ".join(words[i:i + max_words]).strip() for i in range(0, len(words), max_words)]

    chunks: list[str] = []
    current: list[str] = []
    current_words = 0
    for sentence in sentences:
        sentence_words = _word_count(sentence)
        if sentence_words > max_words:
            if current:
                chunks.append(" ".join(current).strip())
                current = []
                current_words = 0
            words = sentence.split()
            chunks.extend(" ".join(words[i:i + max_words]).strip() for i in range(0, len(words), max_words))
            continue
        if current and current_words + sentence_words > max_words:
            chunks.append(" ".join(current).strip())
            current = []
            current_words = 0
        current.append(sentence)
        current_words += sentence_words
    if current:
        chunks.append(" ".join(current).strip())
    return [item for item in chunks if item]


def _split_oversized_section(section: dict[str, Any], max_words: int) -> list[dict[str, Any]]:
    text = str(section.get("text") or "").strip()
    heading = str(section.get("heading") or "").strip()
    protected = bool(section.get("protected"))
    if _word_count(text) <= max_words or protected:
        return [{"heading": heading, "text": text, "protected": protected, "continuation": False}]

    lines = text.splitlines()
    heading_line = ""
    if lines and heading and lines[0].strip() == heading:
        heading_line = lines.pop(0)

    units: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        paragraph = "\n".join(paragraph_buffer).strip()
        if paragraph:
            units.extend(_split_long_text_unit(paragraph, max_words))
        paragraph_buffer = []

    for line in lines:
        if not line.strip():
            flush_paragraph()
        else:
            paragraph_buffer.append(line)
    flush_paragraph()

    chunks: list[dict[str, Any]] = []
    current: list[str] = []
    current_words = 0
    part = 0

    def flush() -> None:
        nonlocal current, current_words, part
        if not current:
            return
        part += 1
        body = "\n\n".join(current).strip()
        if part == 1 and heading_line:
            body = f"{heading_line}\n{body}".strip()
        chunks.append({
            "heading": heading,
            "text": body,
            "protected": False,
            "continuation": part > 1,
        })
        current = []
        current_words = 0

    for unit in units:
        words = _word_count(unit)
        if current and current_words + words > max_words:
            flush()
        current.append(unit)
        current_words += words
    flush()
    return chunks or [{"heading": heading, "text": text, "protected": protected, "continuation": False}]


def _heading_family(heading: str) -> str:
    """Return a stable family so Results and Discussion are never packed together."""
    value = re.sub(r"^#{1,6}\s*", "", str(heading or "").strip())
    numbered = re.match(r"^(\d+)(?:\.\d+)*\.?\s+", value)
    if numbered:
        return f"numbered:{numbered.group(1)}"
    lower = value.lower()
    if lower in {"", "title", "keywords", "key words"} or lower.startswith("abstract"):
        return "front_matter"
    families = [
        ("introduction", ["introduction", "background"]),
        ("literature", ["literature review", "theoretical framework", "conceptual framework"]),
        ("methods", ["method", "methodology", "materials and methods"]),
        ("results", ["result", "findings", "analysis"]),
        ("discussion", ["discussion"]),
        ("conclusion", ["conclusion", "recommendation", "implication", "limitation"]),
        ("declarations", ["declaration", "funding", "competing interest", "conflict of interest", "author contribution", "data availability", "ethics approval", "acknowledgement"]),
        ("protected", ["reference", "bibliography", "source use audit", "appendix"]),
    ]
    for family, prefixes in families:
        if any(lower.startswith(prefix) for prefix in prefixes):
            return family
    return re.sub(r"[^a-z0-9]+", "_", lower)[:40] or "other"


def _split_revision_batches(article_text: str) -> list[dict[str, Any]]:
    max_words = _env_int("ARTICLEREADY_REVISION_SECTION_MAX_WORDS", 1400, minimum=500, maximum=4000)
    pack_limit = _env_int(
        "ARTICLEREADY_REVISION_PACK_MAX_WORDS",
        min(max_words, 1400),
        minimum=400,
        maximum=max_words,
    )
    pack_short = str(os.getenv("ARTICLEREADY_REVISION_PACK_SHORT_SECTIONS", "1") or "1").strip().lower() in {
        "1", "true", "yes", "on"
    }
    sections = _split_revision_sections(article_text)
    raw_batches: list[dict[str, Any]] = []
    for section in sections:
        raw_batches.extend(_split_oversized_section(section, max_words))
    if not raw_batches:
        raw_batches = [{"heading": "Complete manuscript", "text": article_text, "protected": False, "continuation": False}]

    batches: list[dict[str, Any]] = []
    current: list[dict[str, Any]] = []
    current_words = 0
    current_family = ""

    def flush() -> None:
        nonlocal current, current_words, current_family
        if not current:
            return
        labels = [str(item.get("heading") or "").strip() for item in current if str(item.get("heading") or "").strip()]
        text = "\n\n".join(str(item.get("text") or "").strip() for item in current if str(item.get("text") or "").strip()).strip()
        batches.append({
            "heading": labels[0] if labels else "",
            "label": " + ".join(dict.fromkeys(labels))[:220] if labels else f"Manuscript part {len(batches) + 1}",
            "text": text,
            "protected": False,
            "continuation": bool(current[0].get("continuation")) if len(current) == 1 else False,
        })
        current = []
        current_words = 0
        current_family = ""

    for item in raw_batches:
        item_text = str(item.get("text") or "").strip()
        item_words = _word_count(item_text)
        item_family = _heading_family(str(item.get("heading") or ""))
        if bool(item.get("protected")):
            flush()
            batches.append({**item, "label": str(item.get("heading") or "Protected reference section")})
            continue

        family_boundary = bool(current and current_family and item_family != current_family)
        would_overflow = bool(current and current_words + item_words > pack_limit)
        continuation = bool(item.get("continuation"))
        if current and (family_boundary or would_overflow or not pack_short or continuation):
            flush()
        current.append(item)
        current_words += item_words
        current_family = item_family
        if continuation:
            flush()
    flush()

    for index, batch in enumerate(batches, start=1):
        batch["index"] = index
        batch["word_count"] = _word_count(str(batch.get("text") or ""))
        batch["label"] = str(batch.get("label") or batch.get("heading") or "").strip() or f"Manuscript part {index}"
    return batches


def _source_records_for_section(label: str, source_records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    lower = str(label or "").lower()
    if any(token in lower for token in ["reference", "appendix", "declaration"]):
        return []
    if any(token in lower for token in ["method", "result", "finding", "analysis", "data"]):
        limit = _env_int("ARTICLEREADY_REVISION_METHOD_SOURCE_CONTEXT", 16, minimum=0, maximum=40)
    else:
        limit = _env_int("ARTICLEREADY_REVISION_SOURCE_CONTEXT", 36, minimum=0, maximum=60)
    return source_records[:limit]


def _deterministic_revision_plan(payload: dict[str, Any], batches: list[dict[str, Any]]) -> str:
    focus = _revision_focus(payload)
    headings = [str(item.get("label") or "") for item in batches]
    lines = [
        "# Revision Plan",
        "",
        "The revision will preserve confirmed evidence, numerical results, citations and the study design while strengthening the following areas:",
    ]
    lines.extend(f"- {item}." for item in focus)
    lines.extend(["", "## Manuscript sequence", ""])
    lines.extend(f"- {heading}" for heading in headings)
    lines.extend([
        "",
        "## Integrity constraints",
        "",
        "- Do not invent analyses, results, ethics approvals, citations or references.",
        "- Keep unperformed analyses as precise author actions.",
        "- Preserve the distinction between confirmed findings and recommendations for further work.",
    ])
    return "\n".join(lines).strip()


def _build_revision_plan(
    client: Any,
    *,
    payload: dict[str, Any],
    article_inputs: dict[str, Any],
    source_records: list[dict[str, Any]],
    batches: list[dict[str, Any]],
    provider_errors: list[Any],
) -> tuple[str, str]:
    fallback_plan = _deterministic_revision_plan(payload, batches)
    if str(os.getenv("ARTICLEREADY_REVISION_PLAN_ENABLED", "1") or "1").strip().lower() in {"0", "false", "no"}:
        return fallback_plan, "deterministic"

    plan_prompt = {
        "task": "Prepare a compact, section-specific revision plan for a journal article. Do not rewrite the manuscript in this pass.",
        "article_profile": {key: value for key, value in article_inputs.items() if key != "existing_article"},
        "revision_focus": _revision_focus(payload),
        "section_inventory": [
            {
                "section": item.get("label"),
                "word_count": item.get("word_count"),
                "continuation": bool(item.get("continuation")),
                "protected_reference_section": bool(item.get("protected")),
            }
            for item in batches
        ],
        "existing_article": _clip_text(article_inputs.get("existing_article", ""), _env_int("ARTICLEREADY_REVISION_PLAN_MAX_CHARS", 120000, minimum=20000, maximum=180000)),
        "source_record_titles": [
            {"key": item.get("key"), "title": item.get("title"), "year": item.get("year"), "relevance_tier": item.get("relevance_tier")}
            for item in source_records[:24]
        ],
        "rules": [
            "Identify the exact article-level problem, gap and contribution that need clarification.",
            "Identify method, analysis, validity, robustness, discussion and recommendation issues only where relevant to the supplied design.",
            "Separate corrections that can be made from the existing evidence from actions requiring author-supplied data or analysis.",
            "Do not invent evidence, citations, analyses, results or screening counts.",
            "Return a concise Markdown plan organised by manuscript section.",
        ],
    }
    try:
        plan, model_used, notes = _call_openai_response_with_fallback(
            client,
            primary_model=_analysis_model(),
            fallback_model=_revision_plan_model(),
            instructions=(
                "Act as a senior journal diagnostician. Identify the exact conceptual, methodological, analytical and publication-readiness corrections required, "
                "then produce a concise section-specific plan. Do not rewrite the manuscript in this pass."
            ),
            input_payload=json.dumps(plan_prompt, ensure_ascii=False, indent=2),
            max_output_tokens=_env_int("ARTICLEREADY_REVISION_PLAN_MAX_OUTPUT_TOKENS", 4500, minimum=1200, maximum=8000),
            reasoning_effort=os.getenv("OPENAI_ARTICLE_ANALYSIS_REASONING", "xhigh"),
            recovery_reasoning_effort=os.getenv("OPENAI_ARTICLE_ANALYSIS_RECOVERY_REASONING", "high"),
            fallback_reasoning_effort=os.getenv("OPENAI_ARTICLE_FAST_REASONING", "high"),
            request_timeout_seconds=min(_revision_timeout_seconds(), 420),
            include_configured_fallbacks=False,
            require_completed=True,
            purpose="revision_plan",
        )
        provider_errors.extend(notes)
        return _strip_code_fences(plan) or fallback_plan, model_used
    except Exception as exc:
        provider_errors.append(f"Revision plan model unavailable; deterministic plan used: {str(exc)[:180]}")
        return fallback_plan, "deterministic"


def _section_output_tokens(word_count: int) -> int:
    """Allow room for visible revision text and reasoning tokens."""
    hard_cap = _env_int(
        "ARTICLEREADY_REVISION_SECTION_MAX_OUTPUT_TOKENS",
        24000,
        minimum=8000,
        maximum=50000,
    )
    # max_output_tokens includes visible output and reasoning tokens. A small
    # multiple of the source word count is insufficient at high/xhigh effort.
    requested = int(max(350, word_count) * 3.6 + 6500)
    return max(8000, min(hard_cap, requested))


def _section_revision_is_usable(original: str, revised: str) -> tuple[bool, str]:
    original_clean = re.sub(r"\s+", " ", str(original or "")).strip()
    revised_clean = re.sub(r"\s+", " ", str(revised or "")).strip()
    if not revised_clean:
        return False, "empty section response"
    minimum_ratio = float(os.getenv("ARTICLEREADY_REVISION_MIN_SECTION_LENGTH_RATIO", "0.58") or 0.58)
    minimum_ratio = max(0.35, min(minimum_ratio, 0.90))
    if len(revised_clean) < max(1, int(len(original_clean) * minimum_ratio)):
        return False, "section response was substantially truncated"
    # A response ending in an unfinished Markdown marker or mid-sentence is a
    # strong sign that visible output was cut off even when the length ratio passes.
    tail = revised_clean[-180:]
    if tail.count("[") > tail.count("]") or tail.count("(") > tail.count(")"):
        return False, "section response ended with an unfinished structure"
    return True, ""


def _retry_sub_batches(batch: dict[str, Any], target_words: int) -> list[dict[str, Any]]:
    original = str(batch.get("text") or "").strip()
    sections = _split_revision_sections(original)
    pieces: list[dict[str, Any]] = []
    if len(sections) > 1:
        for section in sections:
            pieces.extend(_split_oversized_section(section, target_words))
    else:
        section = {
            "heading": str(batch.get("heading") or ""),
            "text": original,
            "protected": False,
        }
        pieces.extend(_split_oversized_section(section, target_words))

    # Ensure a hard split even for an unusual single-line section that resisted
    # heading and paragraph parsing.
    if len(pieces) <= 1 and _word_count(original) > target_words:
        words = original.split()
        pieces = []
        heading = str(batch.get("heading") or "")
        for index in range(0, len(words), target_words):
            chunk = " ".join(words[index:index + target_words]).strip()
            pieces.append({
                "heading": heading,
                "text": chunk,
                "protected": False,
                "continuation": index > 0,
            })

    results: list[dict[str, Any]] = []
    for index, piece in enumerate(pieces, start=1):
        text = str(piece.get("text") or "").strip()
        if not text:
            continue
        heading = str(piece.get("heading") or batch.get("heading") or "").strip()
        results.append({
            **piece,
            "label": heading or f"{batch.get('label') or 'Section'} retry part {index}",
            "word_count": _word_count(text),
        })
    return results


def _revise_section_batch(
    client: Any,
    *,
    payload: dict[str, Any],
    article_inputs: dict[str, Any],
    revision_plan: str,
    batch: dict[str, Any],
    batch_number: int,
    total_batches: int,
    source_records: list[dict[str, Any]],
    preceding_revised_context: str = "",
    following_original_context: str = "",
    retry_depth: int = 0,
    purpose_suffix: str = "",
) -> tuple[str, str, list[str]]:
    original = str(batch.get("text") or "").strip()
    label = str(batch.get("label") or f"Manuscript part {batch_number}")
    if bool(batch.get("protected")):
        return original, "preserved", []

    prompt = {
        "task": "Substantively revise one section of an existing journal article. Return only the revised section.",
        "article_profile": {key: value for key, value in article_inputs.items() if key != "existing_article"},
        "revision_focus": _revision_focus(payload),
        "revision_plan": revision_plan,
        "batch": {
            "number": batch_number,
            "total": total_batches,
            "section_label": label,
            "continuation_of_same_section": bool(batch.get("continuation")),
            "original_word_count": int(batch.get("word_count") or _word_count(original)),
            "retry_depth": retry_depth,
        },
        "scholarly_source_records": _source_records_for_section(label, source_records),
        "citation_density_requirements": _citation_density_requirements(payload),
        "human_like_writing_layer": _revision_human_writing_layer(payload),
        "author_voice_profile": article_inputs.get("author_voice_profile") or {},
        "continuity_context": {
            "preceding_revised_text_for_flow_only": _context_tail(preceding_revised_context),
            "following_original_text_for_flow_only": _context_head(following_original_context),
            "instruction": "Use these boundary excerpts only to maintain continuity. Do not repeat or rewrite them.",
        },
        "original_section": original,
        "rules": [
            "Revise the supplied section rather than replacing the study with a generic article.",
            "Preserve every confirmed fact, sample detail, date, coefficient, p-value, quotation, table value, citation and reference unless the supplied evidence establishes that it is wrong.",
            "Do not invent results, analyses, participant details, ethics approvals, permissions, citations, references, tables, figures or review counts.",
            "Where a necessary analysis or fact is missing, use one concise [Author action: ...] marker at the exact point of need.",
            "Strengthen conceptual logic, method fit, interpretation, contribution and publication focus only as relevant to this section.",
            "Use source records only where the title or abstract directly supports the claim. Never infer detailed findings from metadata alone.",
            "Preserve existing valid citations and place any verified added citation directly beside the claim it supports.",
            "Follow the global author_voice_profile so the manuscript reads as one scholarly article rather than separately generated batches.",
            "Use polished formal British English, varied sentence rhythm and natural transitions without artificial errors or commentary about humanisation.",
            "Do not force generic signposting, repeated paragraph summaries, balanced triples or a uniform paragraph template.",
            "Keep the section close to its original length unless clarity requires modest expansion or compression.",
            "Preserve the exact section heading. When continuation_of_same_section is true, do not repeat the heading.",
            "Finish the section completely. Do not stop mid-sentence, mid-table or inside an author-action marker.",
            "Return only the revised section in Markdown. Do not provide a report, preface, explanation or code fence.",
        ],
    }

    notes: list[str] = []
    failure: Exception | None = None
    revised = ""
    used_model = ""
    try:
        raw, used_model, provider_notes = _call_openai_response_with_fallback(
            client,
            primary_model=_revision_model() if retry_depth < 2 else _revision_recovery_model(),
            # GPT-5.5 writes the scholarly prose. Terra provides analytical
            # recovery, and Sol is reserved for exceptional smaller-batch escalation.
            fallback_model=(
                _revision_recovery_model()
                if retry_depth < 2
                else _revision_escalation_model()
            ),
            instructions=(
                "You are ArticleReady AI's senior scholarly editor. Revise only the supplied manuscript section in the author's established disciplinary voice. "
                "Preserve evidence and precision, avoid formulaic AI-style prose, never invent analysis or results, and return the complete section only."
            ),
            input_payload=json.dumps(prompt, ensure_ascii=False, indent=2),
            max_output_tokens=_section_output_tokens(int(batch.get("word_count") or _word_count(original))),
            reasoning_effort=_revision_reasoning() if retry_depth == 0 else _revision_recovery_reasoning(),
            recovery_reasoning_effort=_revision_recovery_reasoning(),
            fallback_reasoning_effort=os.getenv("OPENAI_ARTICLE_REVISION_RECOVERY_REASONING", "high") if retry_depth < 2 else os.getenv("OPENAI_ARTICLE_ESCALATION_REASONING", "high"),
            request_timeout_seconds=_revision_timeout_seconds(),
            include_configured_fallbacks=False,
            require_completed=True,
            purpose=f"revision_section_{batch_number}_of_{total_batches}{purpose_suffix}_depth_{retry_depth}",
        )
        notes.extend(provider_notes)
        revised = _strip_code_fences(raw)
        if "===REVISED_ARTICLE===" in revised:
            revised = revised.split("===REVISED_ARTICLE===", 1)[1].strip()
        usable, reason = _section_revision_is_usable(original, revised)
        if not usable:
            failure = RuntimeError(f"{label}: {reason}")
    except Exception as exc:
        failure = exc

    if failure is None:
        return revised, used_model, notes

    max_depth = _env_int("ARTICLEREADY_REVISION_RETRY_SPLIT_MAX_DEPTH", 3, minimum=1, maximum=5)
    min_retry_words = _env_int("ARTICLEREADY_REVISION_RETRY_MIN_WORDS", 280, minimum=120, maximum=800)
    original_words = _word_count(original)
    if retry_depth >= max_depth or original_words <= min_retry_words:
        raise RuntimeError(f"{label}: {str(failure)[:260]}") from failure

    target_words = max(min_retry_words, min(700, max(220, original_words // 2)))
    children = _retry_sub_batches(batch, target_words)
    if len(children) <= 1:
        raise RuntimeError(f"{label}: {str(failure)[:260]}") from failure

    LOGGER.warning(
        "Revision section split-retry label=%s original_words=%s child_count=%s retry_depth=%s error=%s",
        label,
        original_words,
        len(children),
        retry_depth + 1,
        str(failure)[:180],
    )
    notes.append(
        f"Section '{label[:120]}' was automatically divided into {len(children)} smaller recovery batches after an incomplete or truncated response."
    )
    revised_parts: list[str] = []
    models: list[str] = []
    child_preceding = preceding_revised_context
    for child_index, child in enumerate(children, start=1):
        next_child_original = str(children[child_index].get("text") or "") if child_index < len(children) else following_original_context
        child_text, child_model, child_notes = _revise_section_batch(
            client,
            payload=payload,
            article_inputs=article_inputs,
            revision_plan=revision_plan,
            batch=child,
            batch_number=batch_number,
            total_batches=total_batches,
            source_records=source_records,
            preceding_revised_context=child_preceding,
            following_original_context=next_child_original,
            retry_depth=retry_depth + 1,
            purpose_suffix=f"_sub_{child_index}_of_{len(children)}",
        )
        revised_parts.append(child_text.strip())
        child_preceding = child_text
        notes.extend(child_notes)
        if child_model and child_model not in models:
            models.append(child_model)

    combined = "\n\n".join(part for part in revised_parts if part).strip()
    usable, reason = _section_revision_is_usable(original, combined)
    if not usable:
        raise RuntimeError(f"{label}: split-retry assembly failed validation: {reason}")
    return combined, ", ".join(models) or _revision_model(), notes


def _fallback_completed_revision_report(
    *,
    revision_plan: str,
    batch_audit: list[dict[str, Any]],
    provider_errors: list[Any],
) -> str:
    changed = sum(1 for item in batch_audit if item.get("changed"))
    preserved = sum(1 for item in batch_audit if item.get("model") == "preserved")
    warnings = [str(item) for item in provider_errors[-6:] if item]
    warning_text = "\n".join(f"- {item[:220]}" for item in warnings) or "- No material provider warning was recorded."
    return f"""# Revision and Publishability Report

## Revision status

A substantive manuscript revision was completed across {len(batch_audit)} section batch(es). {changed} batch(es) changed materially, while {preserved} protected reference batch(es) were retained unchanged to avoid altering unverified bibliographic details. This report does not guarantee journal acceptance.

## Revision plan applied

{revision_plan}

## Additional analysis

Any analysis not supported by supplied output remains identified in the revised manuscript as an [Author action: ...] item. No suggested analysis is reported as completed.

## Provider and processing notes

{warning_text}

## Remaining author checks

- Verify every citation and reference against the final publisher record.
- Confirm all red author-action items before submission.
- Check the revised manuscript against the target journal's latest author instructions and word limit.
""".strip()


def _generate_revision_report(
    client: Any,
    *,
    payload: dict[str, Any],
    article_inputs: dict[str, Any],
    revision_plan: str,
    revised_article: str,
    batch_audit: list[dict[str, Any]],
    provider_errors: list[Any],
) -> tuple[str, str]:
    report_prompt = {
        "task": "Prepare the Revision and Publishability Report for a substantively revised journal article.",
        "article_profile": {key: value for key, value in article_inputs.items() if key != "existing_article"},
        "revision_plan": revision_plan,
        "batch_audit": batch_audit,
        "revised_article": _clip_text(revised_article, _env_int("ARTICLEREADY_REVISION_REPORT_ARTICLE_CHARS", 110000, minimum=20000, maximum=170000)),
        "requirements": [
            "Begin with an overall publication-readiness assessment without guaranteeing acceptance.",
            "List the most important revisions by section.",
            "Evaluate conceptualisation, theoretical positioning and contribution.",
            "Evaluate method fit, reporting completeness and whether the analysis supports the claims.",
            "Classify additional analyses as Essential, Strongly recommended or Optional. For each, state rationale, required data, suitable method, output to report and consequence of omission.",
            "Evaluate Results-Discussion alignment, implications, limitations and recommendations.",
            "Identify remaining author actions and reference-verification needs.",
            "Do not claim that an unperformed analysis was completed.",
            "Return the report only in Markdown without code fences.",
        ],
    }
    try:
        report, model_used, notes = _call_openai_response_with_fallback(
            client,
            primary_model=_analysis_model(),
            instructions="Write a precise, evidence-grounded journal revision report. Do not invent completed work.",
            input_payload=json.dumps(report_prompt, ensure_ascii=False, indent=2),
            max_output_tokens=_env_int("ARTICLEREADY_REVISION_REPORT_MAX_OUTPUT_TOKENS", 7000, minimum=2000, maximum=12000),
            reasoning_effort=os.getenv("OPENAI_ARTICLE_ANALYSIS_REASONING", os.getenv("OPENAI_ARTICLE_REVISION_REPORT_REASONING", "high")),
            recovery_reasoning_effort="high",
            fallback_reasoning_effort="high",
            request_timeout_seconds=_revision_timeout_seconds(),
            include_configured_fallbacks=False,
            require_completed=True,
            purpose="revision_report",
        )
        provider_errors.extend(notes)
        return _strip_code_fences(report), model_used
    except Exception as exc:
        provider_errors.append(f"Revision report model unavailable; truthful local report used: {str(exc)[:180]}")
        return _fallback_completed_revision_report(
            revision_plan=revision_plan,
            batch_audit=batch_audit,
            provider_errors=provider_errors,
        ), "deterministic"


def _completed_reviewer_matrix(review_comments: str) -> str:
    comments = [line.strip(" -\t") for line in str(review_comments or "").splitlines() if line.strip()]
    if not comments:
        return ""
    rows = ["| Reviewer comment | Revision made | Location | Remaining action |", "|---|---|---|---|"]
    for comment in comments[:40]:
        safe = comment.replace("|", "/")
        rows.append(f"| {safe} | Addressed where supported by the supplied manuscript and evidence. | Verify in the revised manuscript and report. | Confirm that the response fully satisfies the reviewer before submission. |")
    return "\n".join(rows)


def _generate_reviewer_matrix(
    client: Any,
    *,
    payload: dict[str, Any],
    revision_plan: str,
    revision_report: str,
    batch_audit: list[dict[str, Any]],
    provider_errors: list[Any],
) -> tuple[str, str]:
    comments = str(payload.get("review_comments") or "").strip()
    if not comments or not bool(payload.get("include_reviewer_response_matrix", True)):
        return "", "none"
    matrix_prompt = {
        "task": "Prepare a reviewer-response matrix grounded in the completed revision.",
        "reviewer_comments": comments,
        "revision_plan": revision_plan,
        "revision_report": revision_report,
        "section_audit": batch_audit,
        "rules": [
            "Create a Markdown table with columns Reviewer comment, Revision made, Location and Remaining action.",
            "Do not claim that missing evidence or unperformed analysis was supplied.",
            "Where the revision could not fully resolve a comment, state the exact remaining author action.",
            "Return the table only without code fences.",
        ],
    }
    try:
        matrix, model_used, notes = _call_openai_response_with_fallback(
            client,
            primary_model=_revision_plan_model(),
            instructions="Prepare a concise, accurate reviewer-response matrix from the completed revision evidence.",
            input_payload=json.dumps(matrix_prompt, ensure_ascii=False, indent=2),
            max_output_tokens=_env_int("ARTICLEREADY_REVIEWER_MATRIX_MAX_OUTPUT_TOKENS", 4500, minimum=1200, maximum=8000),
            reasoning_effort=os.getenv("OPENAI_ARTICLE_FAST_REASONING", "high"),
            recovery_reasoning_effort="high",
            fallback_reasoning_effort="high",
            request_timeout_seconds=min(_revision_timeout_seconds(), 420),
            include_configured_fallbacks=False,
            require_completed=True,
            purpose="reviewer_response_matrix",
        )
        provider_errors.extend(notes)
        return _strip_code_fences(matrix), model_used
    except Exception as exc:
        provider_errors.append(f"Reviewer-response model unavailable; verification matrix used: {str(exc)[:180]}")
        return _completed_reviewer_matrix(comments), "deterministic"

def revise_article(payload: dict[str, Any]) -> dict[str, Any]:
    article_text = _finalise_article_text(str(payload.get("article_text") or ""))
    if len(article_text) < 100:
        raise ValueError("Paste or upload the existing article before requesting revision.")

    payload = dict(payload)
    payload["article_text"] = article_text
    payload["academic_level"] = str(payload.get("academic_level") or "PhD")
    payload["source_thesis_title"] = ""
    payload["thesis_source_material"] = ""
    payload["objectives"] = str(payload.get("revision_goals") or "")
    payload["theory_or_framework"] = str(payload.get("contribution_claim") or "")
    payload["key_findings"] = str(payload.get("data_and_results") or "")
    payload["variables_constructs"] = str(payload.get("research_area") or "")

    sources, blocked, search_result = _search_sources(payload)
    source_limit = _env_int("ARTICLEREADY_REVISION_MAX_SOURCE_CONTEXT", 60, minimum=10, maximum=100)
    source_records = _source_context(sources)[:source_limit]
    provider_errors: list[Any] = list(search_result.get("provider_errors") or [])
    client = _safe_get_openai_client(_revision_timeout_seconds())
    model_used = "none"
    revision_models_used: list[str] = []
    batch_audit: list[dict[str, Any]] = []
    revision_plan = ""

    if not client or os.getenv("ARTICLEREADY_REVISION_USE_AI", "1").strip().lower() in {"0", "false", "no"}:
        provider_errors.append(
            "OpenAI revision client is unavailable. Confirm that OPENAI_API_KEY is configured and the openai package is installed."
        )
        if not _allow_non_substantive_fallback():
            raise _public_revision_failure(provider_errors)
        revised_article = article_text
        revision_report = _fallback_revision_report(payload, source_records, provider_errors)
        reviewer_matrix = _fallback_reviewer_matrix(str(payload.get("review_comments") or ""))
        mode = "metadata_fallback"
        batches: list[dict[str, Any]] = []
    else:
        article_inputs = {
            "article_title": str(payload.get("article_title") or "").strip(),
            "article_type": str(payload.get("article_type") or "Empirical research article").strip(),
            "research_area": str(payload.get("research_area") or "").strip(),
            "context": str(payload.get("context") or "").strip(),
            "target_journal": str(payload.get("target_journal") or "").strip(),
            "journal_scope": str(payload.get("journal_scope") or "").strip(),
            "author_guidelines": str(payload.get("author_guidelines") or "").strip(),
            "citation_style": str(payload.get("citation_style") or "APA 7th").strip(),
            "word_limit": str(payload.get("word_limit") or "").strip(),
            "methodology_declared_by_author": str(payload.get("methodology") or "").strip(),
            "confirmed_data_results_or_analysis": str(payload.get("data_and_results") or "").strip(),
            "current_contribution_claim": str(payload.get("contribution_claim") or "").strip(),
            "revision_level": str(payload.get("revision_level") or "Publication-readiness overhaul").strip(),
            "additional_revision_goals": str(payload.get("revision_goals") or "").strip(),
            "reviewer_comments": str(payload.get("review_comments") or "").strip(),
            "author_voice_profile": _build_author_voice_profile(article_text, payload),
            "existing_article": article_text,
        }
        candidate_batches = _split_revision_batches(article_text)
        batching_threshold = _env_int("ARTICLEREADY_REVISION_BATCH_THRESHOLD_WORDS", 4500, minimum=1800, maximum=12000)
        batching_applied = _word_count(article_text) >= batching_threshold
        batches = candidate_batches if batching_applied else [{
            "index": 1,
            "heading": "Complete manuscript",
            "label": "Complete manuscript",
            "text": article_text,
            "protected": False,
            "continuation": False,
            "word_count": _word_count(article_text),
        }]

        revision_plan, plan_model = _build_revision_plan(
            client,
            payload=payload,
            article_inputs=article_inputs,
            source_records=source_records,
            batches=batches,
            provider_errors=provider_errors,
        )
        if plan_model not in {"none", "deterministic"}:
            revision_models_used.append(plan_model)

        revised_parts: list[str] = []
        preceding_revised_context = ""
        try:
            for index, batch in enumerate(batches, start=1):
                following_original_context = str(batches[index].get("text") or "") if index < len(batches) else ""
                revised_part, used_model, notes = _revise_section_batch(
                    client,
                    payload=payload,
                    article_inputs=article_inputs,
                    revision_plan=revision_plan,
                    batch=batch,
                    batch_number=index,
                    total_batches=len(batches),
                    source_records=source_records,
                    preceding_revised_context=preceding_revised_context,
                    following_original_context=following_original_context,
                )
                provider_errors.extend(notes)
                if used_model not in {"none", "preserved", "deterministic"} and used_model not in revision_models_used:
                    revision_models_used.append(used_model)
                original_part = str(batch.get("text") or "").strip()
                revised_parts.append(revised_part.strip())
                preceding_revised_context = revised_part
                batch_audit.append({
                    "batch": index,
                    "section": batch.get("label"),
                    "continuation": bool(batch.get("continuation")),
                    "original_words": _word_count(original_part),
                    "revised_words": _word_count(revised_part),
                    "changed": re.sub(r"\s+", " ", original_part).strip() != re.sub(r"\s+", " ", revised_part).strip(),
                    "model": used_model,
                })
            revised_article = "\n\n".join(part for part in revised_parts if part).strip()
            if not _revision_changed(article_text, revised_article):
                raise RuntimeError("The completed section pipeline did not produce a substantive full-manuscript revision.")

            revision_report, report_model = _generate_revision_report(
                client,
                payload=payload,
                article_inputs=article_inputs,
                revision_plan=revision_plan,
                revised_article=revised_article,
                batch_audit=batch_audit,
                provider_errors=provider_errors,
            )
            if report_model not in {"none", "deterministic"} and report_model not in revision_models_used:
                revision_models_used.append(report_model)

            reviewer_matrix, matrix_model = _generate_reviewer_matrix(
                client,
                payload=payload,
                revision_plan=revision_plan,
                revision_report=revision_report,
                batch_audit=batch_audit,
                provider_errors=provider_errors,
            )
            if matrix_model not in {"none", "deterministic"} and matrix_model not in revision_models_used:
                revision_models_used.append(matrix_model)
            mode = "ai_revision"
            model_used = ", ".join(revision_models_used) or _revision_model()
        except Exception as exc:
            LOGGER.exception(
                "Article revision pipeline failed model=%s batches=%s article_chars=%s",
                _revision_model(), len(batches), len(article_text),
            )
            provider_errors.append(f"OpenAI article revision failed: {str(exc)[:220]}")
            if not _allow_non_substantive_fallback():
                raise _public_revision_failure(provider_errors) from exc
            revised_article = article_text
            revision_report = _fallback_revision_report(payload, source_records, provider_errors)
            reviewer_matrix = _fallback_reviewer_matrix(str(payload.get("review_comments") or ""))
            mode = "metadata_fallback_after_ai_error"
            batching_applied = False

    revised_article = _finalise_article_text(revised_article)
    humanizer_report: dict[str, Any] = {"mode": _humanizer_mode(payload), "applied": False}
    humanizer_models: list[str] = []
    if mode == "ai_revision":
        use_second_model_pass = str(
            os.getenv("ARTICLEREADY_REVISION_SECOND_HUMANIZER_MODEL_PASS", "1") or "1"
        ).strip().lower() in {"1", "true", "yes", "on"}
        if use_second_model_pass:
            revised_article, humanizer_report, humanizer_models = _humanize_article_with_model(
                client,
                revised_article,
                payload=payload,
                provider_errors=provider_errors,
            )
        else:
            revised_article, humanizer_report = humanize_scholarly_text(
                revised_article,
                mode=_humanizer_mode(payload),
            )
            humanizer_report["model_pass_applied"] = False
            humanizer_report["revision_section_prompts_applied_humanizer_rules"] = True
        revised_article = _finalise_article_text(revised_article)
    revision_report = _finalise_article_text(revision_report)
    reviewer_matrix = _finalise_article_text(reviewer_matrix) if reviewer_matrix else ""
    citation_density = _citation_density_report(revised_article)
    density_target = _citation_density_requirements(payload)["citation_occurrences_per_1000_words"]
    citation_density["minimum_target"] = int(density_target["minimum"])
    citation_density["preferred_target"] = int(density_target["target"])
    citation_density["meets_minimum"] = float(citation_density.get("citation_occurrences_per_1000_words") or 0) >= int(density_target["minimum"])

    return {
        "revised_article_text": revised_article,
        "revision_report": revision_report,
        "reviewer_response_matrix": reviewer_matrix,
        "revision_plan": revision_plan,
        "model_used": model_used if client else "none",
        "revision_models_used": revision_models_used,
        "reasoning_effort": _revision_reasoning() if client else "none",
        "mode": mode,
        "revision_batching_applied": bool(mode == "ai_revision" and batching_applied),
        "revision_batch_count": len(batch_audit),
        "revision_batch_audit": batch_audit,
        "source_records_used": source_records,
        "source_bank_count": len(source_records),
        "excluded_retracted_count": len(blocked),
        "excluded_retracted_titles": [str(item.get("title") or "Untitled") for item in blocked[:10]],
        "provider_errors": provider_errors,
        "revision_colour_note": "In the downloaded DOCX, wording added or changed by ArticleReady AI is shown in blue. Author-action items are shown in red. Exact unchanged wording remains black.",
        "human_like_writing_layer_applied": mode == "ai_revision" and _strong_humanisation_enabled(),
        "humanisation_strength": _humanisation_strength(),
        "humanizer_report": humanizer_report,
        "humanizer_models_used": humanizer_models,
        "citation_density_report": citation_density,
        "quality_filters": [
            "Long manuscripts are revised section by section to reduce truncation and timeout risk.",
            "GPT-5.5 performs the substantive scholarly rewriting, GPT-5.6 Terra provides analytical recovery and reporting, and GPT-5.6 Sol is reserved for exceptional escalation.",
            "GPT-5.6 Terra diagnoses the manuscript with xhigh reasoning before rewriting begins; GPT-5.6 Luna remains the lower-cost planning fallback and reviewer-response model.",
            "A global author-voice profile and section-boundary context keep batched revisions stylistically coherent.",
            "GPT-5.4 performs a selective preservation-gated final naturalness pass on sections that need it.",
            "The post-processing pass protects confirmed evidence, citations, tables, equations, placeholders and references.",
            "Confirmed results and numerical evidence are preserved unless the author supplies a correction.",
            "Suggested additional analyses are not presented as completed analyses.",
            "Scholarly records are subject to a relevance gate and retraction screening.",
            "The revision report does not guarantee journal acceptance.",
        ],
    }

def _plain_compare_text(text: str) -> str:
    text = re.sub(r"^#{1,6}\s+", "", str(text or "").strip())
    text = re.sub(r"^[-*•]\s+", "", text)
    text = re.sub(r"^\d+[.)]\s+", "", text)
    text = _plain_inline_text(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()


def _tokenise_for_diff(text: str) -> list[str]:
    return re.findall(r"\s+|[A-Za-z0-9]+(?:['’-][A-Za-z0-9]+)*|[^\w\s]", text, flags=re.UNICODE)


def _word_key(token: str) -> str:
    if token.isspace():
        return token
    return token.lower()


def _best_original_line(revised_line: str, original_lines: list[str], used: set[int]) -> tuple[str | None, float, int | None]:
    revised_norm = _plain_compare_text(revised_line)
    if not revised_norm:
        return None, 0.0, None
    best_line = None
    best_score = 0.0
    best_index = None
    for index, candidate in enumerate(original_lines):
        if index in used:
            continue
        candidate_norm = _plain_compare_text(candidate)
        if not candidate_norm:
            continue
        score = difflib.SequenceMatcher(None, candidate_norm, revised_norm, autojunk=False).ratio()
        if score > best_score:
            best_score = score
            best_line = candidate
            best_index = index
    return best_line, best_score, best_index


def _append_marked_run(
    paragraph,
    text: str,
    changed: bool,
    bold: bool = False,
    italic: bool = False,
    action_required: bool = False,
) -> None:
    if not text:
        return
    from docx.shared import RGBColor

    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if action_required:
        run.font.color.rgb = RGBColor(*_ACTION_RED)
    elif changed:
        run.font.color.rgb = RGBColor(*_REVISION_BLUE)


def _styled_diff_tokens(text: str) -> list[tuple[str, bool, bool, bool]]:
    tokens: list[tuple[str, bool, bool, bool]] = []
    for visible, bold, italic, action_required in _parse_inline_segments(text):
        for token in _tokenise_for_diff(visible):
            tokens.append((token, bold, italic, action_required))
    return tokens


def _add_diff_runs(paragraph, revised_text: str, original_text: str | None, changed_default: bool = True) -> None:
    """Render emphasis, keep revisions blue and override action-required text in red."""
    revised_tokens = _styled_diff_tokens(revised_text)
    if not original_text:
        for token, bold, italic, action_required in revised_tokens:
            _append_marked_run(
                paragraph, token, True if token else changed_default, bold, italic, action_required
            )
        return

    original_tokens = _styled_diff_tokens(original_text)
    matcher = difflib.SequenceMatcher(
        None,
        [_word_key(token[0]) for token in original_tokens],
        [_word_key(token[0]) for token in revised_tokens],
        autojunk=False,
    )
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        for offset, (token, bold, italic, action_required) in enumerate(revised_tokens[j1:j2]):
            changed = tag != "equal"
            if tag == "equal" and i1 + offset < i2:
                original_token = original_tokens[i1 + offset]
                # A formatting change is also a revision and should remain blue.
                if not token.isspace():
                    changed = (bold, italic) != (original_token[1], original_token[2])
            _append_marked_run(paragraph, token, changed, bold, italic, action_required)


def _add_black_inline_runs(paragraph, text: str, force_bold: bool = False) -> None:
    """Render Markdown emphasis in black, with action-required text in red."""
    for visible, bold, italic, action_required in _parse_inline_segments(text):
        _append_marked_run(
            paragraph,
            visible,
            False,
            bool(bold or force_bold),
            italic,
            action_required,
        )

def _revision_line_map(original_text: str, revised_lines: list[str]) -> dict[int, str | None]:
    original_lines = [line.rstrip() for line in _finalise_article_text(original_text).splitlines() if line.strip()]
    used: set[int] = set()
    mapping: dict[int, str | None] = {}
    exact_lookup: dict[str, list[int]] = {}
    for index, line in enumerate(original_lines):
        exact_lookup.setdefault(_plain_compare_text(line), []).append(index)

    for revised_index, line in enumerate(revised_lines):
        norm = _plain_compare_text(line)
        exact_candidates = exact_lookup.get(norm, [])
        exact_index = next((idx for idx in exact_candidates if idx not in used), None)
        if exact_index is not None:
            used.add(exact_index)
            mapping[revised_index] = original_lines[exact_index]
            continue
        candidate, score, candidate_index = _best_original_line(line, original_lines, used)
        if candidate is not None and score >= 0.36 and candidate_index is not None:
            used.add(candidate_index)
            mapping[revised_index] = candidate
        else:
            mapping[revised_index] = None
    return mapping


def _add_revision_table(doc, lines: list[str], original_text: str, colour_revisions: bool = True) -> None:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    original_norm = _plain_compare_text(original_text)
    for row_index, cells in enumerate(rows):
        row = table.add_row().cells
        for column in range(width):
            value = cells[column] if column < len(cells) else ""
            paragraph = row[column].paragraphs[0]
            paragraph.clear()
            if colour_revisions:
                visible_value = _plain_compare_text(value)
                changed = bool(visible_value and visible_value not in original_norm)
                for visible, bold, italic, action_required in _parse_inline_segments(value):
                    _append_marked_run(
                        paragraph,
                        visible,
                        changed,
                        bool(bold or row_index == 0),
                        italic,
                        action_required,
                    )
            else:
                _add_black_inline_runs(paragraph, value, force_bold=row_index == 0)

def _write_markdown_document(doc, markdown_text: str, original_text: str | None = None, colour_revisions: bool = False, heading_offset: int = 0) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    lines = _finalise_article_text(markdown_text).splitlines()
    line_map = _revision_line_map(original_text or "", [line for line in lines]) if colour_revisions else {}
    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False
    equation_buffer: list[str] = []
    in_equation = False

    def add_content(paragraph, content: str, index: int, original_override: str | None = None) -> None:
        original_line = original_override if original_override is not None else line_map.get(index)
        if colour_revisions:
            _add_diff_runs(paragraph, content, original_line)
        else:
            _add_black_inline_runs(paragraph, content)

    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                for code_line in code_buffer:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(code_line)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    if colour_revisions:
                        run.font.color.rgb = RGBColor(*_REVISION_BLUE)
                code_buffer = []
                in_code = False
            else:
                if table_buffer:
                    _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
                    table_buffer = []
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        if stripped == "$$":
            if in_equation:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(" ".join(equation_buffer))
                run.font.name = "Cambria Math"
                run.font.size = Pt(12)
                if colour_revisions and _plain_compare_text(run.text) not in _plain_compare_text(original_text or ""):
                    run.font.color.rgb = RGBColor(*_REVISION_BLUE)
                equation_buffer = []
                in_equation = False
            else:
                if table_buffer:
                    _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
                    table_buffer = []
                in_equation = True
            continue
        if in_equation:
            equation_buffer.append(stripped)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
            table_buffer = []

        if not stripped:
            continue
        if line.startswith("# "):
            level = min(3, 0 + heading_offset)
            paragraph = doc.add_heading("", level=level)
            if level == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_content(paragraph, line[2:].strip(), index)
        elif line.startswith("## "):
            paragraph = doc.add_heading("", level=min(3, 1 + heading_offset))
            add_content(paragraph, line[3:].strip(), index)
        elif line.startswith("### "):
            paragraph = doc.add_heading("", level=min(3, 2 + heading_offset))
            add_content(paragraph, line[4:].strip(), index)
        elif line.startswith("#### "):
            paragraph = doc.add_heading("", level=3)
            add_content(paragraph, line[5:].strip(), index)
        elif re.match(r"^[-*•]\s+", line):
            content = re.sub(r"^[-*•]\s+", "", line).strip()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_content(paragraph, content, index)
        elif re.match(r"^\d+[.)]\s+", line):
            content = re.sub(r"^\d+[.)]\s+", "", line).strip()
            paragraph = doc.add_paragraph(style="List Number")
            add_content(paragraph, content, index)
        else:
            paragraph = doc.add_paragraph()
            add_content(paragraph, line, index)

    if table_buffer:
        _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
    if code_buffer:
        for code_line in code_buffer:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code_line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            if colour_revisions:
                run.font.color.rgb = RGBColor(*_REVISION_BLUE)
    if equation_buffer:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(" ".join(equation_buffer))
        run.font.name = "Cambria Math"
        run.font.size = Pt(12)
        if colour_revisions:
            run.font.color.rgb = RGBColor(*_REVISION_BLUE)


def export_revised_article_docx(
    original_article_text: str,
    revised_article_text: str,
    title: str = "Revised Journal Article",
    revision_report: str = "",
    reviewer_response_matrix: str = "",
    include_revision_report: bool = True,
) -> tuple[io.BytesIO, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", (title or "revised_article")[:80]).strip("_") or "revised_article"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Revision display: wording added or changed by ArticleReady AI appears in blue. Author-action items appear in red. Exact unchanged wording remains black.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 98, 115)

    _write_markdown_document(
        doc,
        revised_article_text,
        original_text=original_article_text,
        colour_revisions=True,
    )

    if include_revision_report and (revision_report.strip() or reviewer_response_matrix.strip()):
        doc.add_page_break()
        heading = doc.add_heading("Revision and Publishability Report", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if revision_report.strip():
            clean_report = re.sub(r"^#\s+Revision and Publishability Report\s*", "", revision_report.strip(), count=1, flags=re.IGNORECASE)
            _write_markdown_document(doc, clean_report, colour_revisions=False, heading_offset=1)
        if reviewer_response_matrix.strip():
            doc.add_heading("Response to Reviewer Comments", level=1)
            _write_markdown_document(doc, reviewer_response_matrix, colour_revisions=False, heading_offset=1)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream, f"{safe_title}_polished_revision_blue.docx"
