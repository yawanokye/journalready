from app.article_ideas_service import generate_article_ideas
from app.article_service import draft_journal_article, export_article_docx


def test_article_ideas_fallback_is_article_focused(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    result = generate_article_ideas(
        {
            "research_area": "digital procurement and public expenditure",
            "source_mode": "Extract from a completed thesis or dissertation",
            "thesis_title": "Digitalisation, procurement governance and public expenditure",
            "thesis_material": "The study analysed public procurement systems and expenditure outcomes.",
            "context": "African public sectors",
            "article_type": "Empirical research article",
            "variables_or_themes": "e-procurement, procurement expenditure, governance",
            "max_ideas": 4,
            "include_source_search": False,
        }
    )
    assert len(result["ideas"]) == 4
    assert all(idea["objective"].startswith("To ") for idea in result["ideas"])
    assert all("scope_warning" in idea for idea in result["ideas"])


def test_article_draft_and_docx_fallback(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    result = draft_journal_article(
        {
            "article_title": "Digital procurement and expenditure intensity",
            "article_type": "Empirical research article",
            "academic_level": "Research Masters (e.g. MPhil)",
            "include_source_search": False,
            "thesis_source_material": "Study summary",
        }
    )
    assert "# Digital procurement and expenditure intensity" in result["article_text"]
    stream, filename = export_article_docx(result["article_text"], "Digital procurement")
    assert filename.endswith(".docx")
    assert len(stream.getvalue()) > 1000


def test_attached_source_bank_is_used_and_retracted_record_is_excluded(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    result = draft_journal_article(
        {
            "article_title": "Governance and digital procurement outcomes",
            "article_type": "Empirical research article",
            "include_source_search": False,
            "source_search_terms": "digital procurement governance",
            "source_bank": [
                {
                    "title": "Digital procurement and public-sector governance",
                    "authors": ["A. Researcher"],
                    "year": 2025,
                    "source": "Public Management Review",
                    "doi": "10.1000/example",
                    "abstract": "The study examines digital procurement and governance outcomes.",
                    "apa_hint": "Researcher, A. (2025). Digital procurement and public-sector governance.",
                },
                {
                    "title": "Retracted: Procurement systems and governance",
                    "authors": ["B. Author"],
                    "year": 2022,
                    "is_retracted": True,
                },
            ],
        }
    )
    assert result["attached_source_count"] == 1
    assert result["source_bank_count"] == 1
    assert result["excluded_retracted_count"] == 1
    assert result["source_records_used"][0]["title"] == "Digital procurement and public-sector governance"
    assert "## Source Use Audit" in result["article_text"]
    assert "Retracted: Procurement systems and governance" not in result["article_text"]


def test_article_ideas_list_secondary_data_sources(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    result = generate_article_ideas(
        {
            "research_area": "public procurement governance and expenditure",
            "source_mode": "Develop from an existing dataset",
            "article_type": "Empirical research article",
            "research_route": "Secondary data or existing dataset",
            "methodology": "cross-country panel data analysis",
            "variables_or_themes": "e-procurement, governance, public expenditure",
            "max_ideas": 3,
            "include_source_search": False,
            "include_research_resource_search": True,
        }
    )
    assert result["research_resources"]["data_sources"]
    assert any(idea["resource_guidance"]["possible_data_sources"] for idea in result["ideas"])


def test_independent_article_defaults_to_phd_and_stops_at_methods(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    result = draft_journal_article(
        {
            "article_title": "Financial literacy and retirement planning",
            "source_mode": "Develop as a new independent article",
            "draft_stage": "full_article",
            "academic_level": "Research Masters (e.g. MPhil)",
            "research_route": "Primary survey or questionnaire",
            "methodology": "cross-sectional questionnaire survey",
            "variables_constructs": "financial literacy, retirement planning",
            "include_source_search": False,
            "include_research_resource_search": True,
            "include_instrument_draft": True,
        }
    )
    assert result["draft_stage"] == "initial_to_methods"
    assert result["academic_level_used"] == "PhD"
    assert "## 4. Methods" in result["article_text"]
    assert "## 5. Results" not in result["article_text"]
    assert result["instrument_text"]
    assert result["research_resources"]["instrument_sources"]


def test_continuation_stage_requires_previous_or_results(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    try:
        draft_journal_article(
            {
                "article_title": "Article continuation",
                "source_mode": "Develop as a new independent article",
                "draft_stage": "continuation_after_results",
                "include_source_search": False,
            }
        )
    except ValueError as exc:
        assert "previous article sections" in str(exc)
    else:
        raise AssertionError("Expected continuation validation error")


def test_text_file_extraction():
    from app.file_extractor import extract_uploaded_text

    result = extract_uploaded_text("results.txt", b"Model 1 coefficient = 0.42\np-value = 0.01")
    assert result["filename"] == "results.txt"
    assert "0.42" in result["text"]


def test_article_ideas_use_deepseek_v4_pro_only(monkeypatch):
    import json
    from types import SimpleNamespace
    import app.article_ideas_service as ideas_service

    captured = {}
    payload_json = {
        "ideas": [
            {
                "title": "Digital procurement and expenditure outcomes",
                "article_type": "Empirical research article",
                "angle": "A focused empirical paper.",
                "gap": "Recent evidence remains limited.",
                "objective": "To examine digital procurement and expenditure outcomes.",
                "questions_or_hypotheses": ["How are the variables related?"],
                "contribution": "Provides focused evidence.",
                "method_and_data_route": "Use secondary panel data.",
                "journal_fit": "Assess against the selected journal.",
                "suggested_sections": ["Introduction", "Methods", "Results", "Discussion"],
                "keywords": ["digital procurement", "expenditure"],
                "evidence_needed": ["Verified data"],
                "scope_warning": "Keep one central contribution.",
                "readiness_score": 82,
                "research_route": "secondary_data",
            }
        ],
        "portfolio_note": "One focused paper.",
    }

    class FakeCompletions:
        def create(self, **kwargs):
            captured.update(kwargs)
            message = SimpleNamespace(content=json.dumps(payload_json))
            return SimpleNamespace(choices=[SimpleNamespace(message=message)])

    fake_client = SimpleNamespace(chat=SimpleNamespace(completions=FakeCompletions()))
    monkeypatch.setenv("DEEPSEEK_API_KEY", "test-deepseek-key")
    monkeypatch.setenv("OPENAI_API_KEY", "test-openai-key")
    monkeypatch.setattr(ideas_service, "_safe_get_deepseek_client", lambda: fake_client)

    result = ideas_service.generate_article_ideas(
        {
            "research_area": "digital procurement",
            "article_type": "Empirical research article",
            "max_ideas": 3,
            "include_source_search": False,
            "include_research_resource_search": False,
        }
    )

    assert captured["model"] == "deepseek-v4-pro"
    assert captured["extra_body"]["thinking"]["type"] == "enabled"
    assert captured["extra_body"]["reasoning_effort"] == "high"
    assert result["model_used"] == "deepseek-v4-pro"
    assert result["mode"] == "ai_generated"


def test_independent_term_structure_ideas_are_secondary_and_thesis_free(monkeypatch):
    import re

    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    result = generate_article_ideas(
        {
            "research_area": "term structure of interest rate in Ghana",
            "source_mode": "Develop as a new independent article",
            "discipline": "Finance",
            "context": "Ghana",
            "article_type": "Empirical research article",
            "research_route": "Auto",
            "max_ideas": 8,
            "include_source_search": False,
            "include_research_resource_search": True,
        }
    )
    rendered = " ".join(str(idea) for idea in result["ideas"])
    assert not re.search(r"\b(thesis|dissertation)\b", rendered, flags=re.IGNORECASE)
    assert all(idea["research_route"] == "secondary_data" for idea in result["ideas"])
    assert result["research_resources"]["data_sources"]
    assert not result["research_resources"]["instrument_sources"]
    assert result["ideas"][0]["readiness_score"] < 70


def test_topic_source_filter_rejects_weak_keyword_and_country_matches():
    from app.article_ideas_service import _filter_topic_sources

    payload = {
        "research_area": "term structure of interest rate in Ghana",
        "context": "Ghana",
    }
    sources = [
        {"title": "Testing the expectations hypothesis of the term structure of interest rate: the case of Ghana"},
        {"title": "The Term Structure of Interest Rates"},
        {"title": "Effects of interest rate on bank dividends in Ghana"},
        {"title": "Psychometric properties of an African career interest inventory in Ghana"},
    ]
    retained, excluded = _filter_topic_sources(sources, payload, 10)
    titles = [item["title"] for item in retained]
    assert excluded == 2
    assert "The Term Structure of Interest Rates" in titles
    assert all("career interest" not in title.lower() for title in titles)
    assert all("bank dividends" not in title.lower() for title in titles)


def test_source_query_deduplicates_repeated_topic_and_context():
    from app.source_finder import build_source_query

    query = build_source_query(
        {
            "title": "term structure of interest rate in Ghana",
            "research_area": "term structure of interest rate in Ghana",
            "study_context": "Ghana",
            "objectives": [],
        },
        "term structure of interest rate in Ghana",
    )
    assert query.lower().count("term structure of interest rate in ghana") == 1


def test_article_revision_fallback_preserves_article_and_reports_actions(monkeypatch):
    from app.article_revision_service import revise_article

    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    original = (
        "# Digital procurement and expenditure\n\n"
        "## Introduction\n\nDigital procurement may affect public expenditure outcomes. "
        "The study examines this relationship using cross-country data.\n\n"
        "## Methods\n\nThe study uses panel regression and governance controls.\n\n"
        "## Results\n\nThe confirmed coefficient is -0.12 with p = 0.04."
    )
    result = revise_article(
        {
            "article_title": "Digital procurement and expenditure",
            "article_text": original,
            "review_comments": "Clarify the theoretical contribution.\nAdd robustness analysis.",
            "include_source_search": False,
        }
    )
    assert result["revised_article_text"] == original
    assert "Revision and Publishability Report" in result["revision_report"]
    assert "Clarify the theoretical contribution" in result["reviewer_response_matrix"]
    assert result["mode"] == "metadata_fallback"


def test_revision_docx_colours_changed_text_blue():
    from docx import Document
    from app.article_revision_service import export_revised_article_docx

    original = "# Article Title\n\nThe model explains procurement outcomes."
    revised = "# Article Title\n\nThe revised model more clearly explains public procurement outcomes."
    stream, filename = export_revised_article_docx(
        original_article_text=original,
        revised_article_text=revised,
        title="Article Title",
        revision_report="# Report\n\nAdditional robustness analysis is recommended.",
    )
    assert filename.endswith("_polished_revision_blue.docx")
    document = Document(stream)
    colours = []
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.color.rgb is not None:
                colours.append(str(run.font.color.rgb))
    assert "0070C0" in colours


def test_revision_package_parser_handles_report_and_matrix():
    from app.article_revision_service import _split_revision_package

    revised, report, matrix = _split_revision_package(
        "===REVISED_ARTICLE===\n# Revised\nText\n"
        "===REVISION_REPORT===\n# Report\nDetails\n"
        "===REVIEWER_RESPONSE_MATRIX===\n| Comment | Action |\n|---|---|\n| A | B |"
    )
    assert revised.startswith("# Revised")
    assert report.startswith("# Report")
    assert "| A | B |" in matrix


def test_revision_docx_renders_bold_italic_and_red_actions():
    from docx import Document
    from app.article_revision_service import export_revised_article_docx

    original = "# Article Title\n\nThe model explains procurement outcomes."
    revised = (
        "# Article Title\n\n"
        "The **revised model** offers *careful interpretation* of procurement outcomes. "
        "[confirm robustness test before submission]"
    )
    stream, _ = export_revised_article_docx(
        original_article_text=original,
        revised_article_text=revised,
        title="Article Title",
    )
    document = Document(stream)
    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "**" not in all_text
    assert "*careful interpretation*" not in all_text

    runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
    assert any(run.bold and run.text == "revised" for run in runs)
    assert any(run.bold and run.text == "model" for run in runs)
    assert any(run.italic and run.text == "careful" for run in runs)
    assert any(run.italic and run.text == "interpretation" for run in runs)
    red_action_text = "".join(
        run.text for run in runs
        if run.font.color.rgb is not None and str(run.font.color.rgb) == "C00000"
    )
    assert "confirm robustness test" in red_action_text
    assert any(
        "revised" in run.text
        and run.font.color.rgb is not None
        and str(run.font.color.rgb) == "0070C0"
        for run in runs
    )


def test_article_docx_renders_markdown_and_red_actions():
    from docx import Document
    from app.article_service import export_article_docx

    stream, _ = export_article_docx(
        "# **Article Title**\n\nThis is *important*. [insert verified source]",
        title="Article Title",
    )
    document = Document(stream)
    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "**" not in all_text
    assert "*important*" not in all_text
    runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
    assert any(run.bold and "Article Title" in run.text for run in runs)
    assert any(run.italic and "important" in run.text for run in runs)
    assert any(
        "insert verified source" in run.text
        and run.font.color.rgb is not None
        and str(run.font.color.rgb) == "C00000"
        for run in runs
    )


def test_long_article_length_plan_and_token_estimate(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    result = draft_journal_article(
        {
            "article_title": "Graduate employability and talent pipelines",
            "article_type": "Empirical research article",
            "include_source_search": False,
            "word_limit": "7000-9000",
            "target_word_count": 8000,
            "long_write_mode": "batch",
            "article_structure": "Introduction 1200\nLiterature review and hypotheses 1800\nMethods 1400\nResults 1400\nDiscussion 1600\nConclusion 600",
            "thesis_source_material": "Completed study summary with results.",
        }
    )
    assert result["article_length_plan"]["target_words"] == 8000
    assert result["article_length_plan"]["structure_source"] == "user_supplied"
    assert result["token_budget_estimate"]["target_words"] == 8000
    assert result["token_budget_estimate"]["drafting_passes"] >= 2
    assert result["batch_drafting_applied"] is False  # no API key, so fallback text is used


def test_gpt56_terra_sol_model_routing(monkeypatch):
    from app.article_service import _select_article_model
    monkeypatch.setenv('OPENAI_ARTICLE_TERRA_MODEL', 'gpt-5.6-terra')
    monkeypatch.setenv('OPENAI_ARTICLE_SOL_MODEL', 'gpt-5.6-sol')
    assert _select_article_model('Bachelors', 'Empirical research article', {'target_word_count': 7000}) == 'gpt-5.6-terra'
    assert _select_article_model('Research Masters (e.g. MPhil)', 'Empirical research article', {'target_word_count': 7000}) == 'gpt-5.6-sol'
    assert _select_article_model('Bachelors', 'Systematic review article', {'target_word_count': 7000}) == 'gpt-5.6-sol'
    assert _select_article_model('Bachelors', 'Empirical research article', {'target_word_count': 11000}) == 'gpt-5.6-sol'
    monkeypatch.delenv('OPENAI_ARTICLE_TERRA_MODEL', raising=False)
    monkeypatch.setenv('OPENAI_ARTICLE_BACHELOR_MODEL', 'gpt-5.4')
    assert _select_article_model('Bachelors', 'Empirical research article', {'target_word_count': 7000}) == 'gpt-5.6-terra'


def test_thesisready_humanizer_preserves_scholarly_evidence():
    from app.scholarly_humanizer import humanize_scholarly_text, validate_humanizer_preservation
    original = (
        '# Discussion\n\n'
        'The evidence indicates that procurement digitisation is associated with transparency (Adam, 2024). '
        'The coefficient is 0.42 and the p-value is 0.01. [Author action: Confirm the robustness specification.]\n\n'
        '## References\n\nAdam, A. (2024). Digital procurement evidence. https://doi.org/10.1000/example'
    )
    candidate, report = humanize_scholarly_text(original, mode='balanced')
    valid, issues = validate_humanizer_preservation(original, candidate, max_word_change_ratio=0.06)
    assert valid, issues
    assert '(Adam, 2024)' in candidate
    assert '0.42' in candidate and '0.01' in candidate
    assert '[Author action: Confirm the robustness specification.]' in candidate
    assert report['mode'] == 'balanced'
