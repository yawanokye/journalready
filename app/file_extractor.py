from __future__ import annotations

import csv
import io
import os
import re
import zipfile
from typing import Any

MAX_UPLOAD_BYTES = int(os.getenv("ARTICLEREADY_MAX_UPLOAD_BYTES", str(15 * 1024 * 1024)))
MAX_EXTRACTED_CHARS = int(os.getenv("ARTICLEREADY_MAX_EXTRACTED_CHARS", "190000"))


MAX_ARCHIVE_FILES = int(os.getenv("ARTICLEREADY_MAX_ARCHIVE_FILES", "5000"))
MAX_ARCHIVE_UNCOMPRESSED_BYTES = int(os.getenv("ARTICLEREADY_MAX_ARCHIVE_UNCOMPRESSED_BYTES", str(120 * 1024 * 1024)))
MAX_ARCHIVE_COMPRESSION_RATIO = float(os.getenv("ARTICLEREADY_MAX_ARCHIVE_COMPRESSION_RATIO", "200"))
MAX_PDF_PAGES = int(os.getenv("ARTICLEREADY_MAX_PDF_PAGES", "500"))


async def read_upload_limited(file: Any, max_bytes: int) -> bytes:
    """Read an UploadFile incrementally and reject oversized chunked uploads."""
    chunks: list[bytes] = []
    total = 0
    while True:
        chunk = await file.read(min(1024 * 1024, max_bytes + 1 - total))
        if not chunk:
            break
        total += len(chunk)
        if total > max_bytes:
            raise ValueError(f"The uploaded file exceeds the {max_bytes // (1024 * 1024)} MB limit.")
        chunks.append(chunk)
    return b"".join(chunks)


def _validate_zip_container(content: bytes, label: str) -> None:
    """Reject path traversal and decompression-bomb style Office archives."""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            members = archive.infolist()
            if len(members) > MAX_ARCHIVE_FILES:
                raise ValueError(f"The {label} contains too many internal files.")
            total_uncompressed = 0
            total_compressed = 0
            for member in members:
                name = str(member.filename or "").replace("\\", "/")
                if name.startswith("/") or "../" in f"/{name}":
                    raise ValueError(f"The {label} contains an unsafe internal path.")
                total_uncompressed += max(0, int(member.file_size or 0))
                total_compressed += max(0, int(member.compress_size or 0))
                if total_uncompressed > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                    raise ValueError(f"The {label} expands beyond the safe extraction limit.")
            ratio = total_uncompressed / max(1, total_compressed)
            if ratio > MAX_ARCHIVE_COMPRESSION_RATIO and total_uncompressed > 10 * 1024 * 1024:
                raise ValueError(f"The {label} has an unsafe compression ratio.")
    except zipfile.BadZipFile as exc:
        raise ValueError(f"The uploaded {label} is not a valid Office file.") from exc


def _normalise(text: str) -> str:
    text = str(text or "").replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_EXTRACTED_CHARS]


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_docx(content: bytes) -> str:
    from docx import Document

    _validate_zip_container(content, "DOCX file")
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError(f"The PDF exceeds the {MAX_PDF_PAGES}-page safety limit.")
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            pages.append(f"[Page {index}]\n{page_text.strip()}")
    if not pages:
        raise ValueError("No selectable text was found in the PDF. Upload a text-based PDF or paste the content manually.")
    return "\n\n".join(pages)


def _extract_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook

    _validate_zip_container(content, "XLSX file")
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    output: list[str] = []
    for sheet in workbook.worksheets:
        output.append(f"[Worksheet: {sheet.title}]")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if not any(value.strip() for value in values):
                continue
            output.append(" | ".join(values))
            row_count += 1
            if row_count >= 2000:
                output.append("[Worksheet truncated after 2,000 non-empty rows]")
                break
    return "\n".join(output)


def _extract_csv(content: bytes) -> str:
    text = _decode_text(content)
    rows: list[str] = []
    reader = csv.reader(io.StringIO(text))
    for index, row in enumerate(reader):
        rows.append(" | ".join(str(cell) for cell in row))
        if index >= 2000:
            rows.append("[CSV truncated after 2,001 rows]")
            break
    return "\n".join(rows)


def extract_uploaded_text(filename: str, content: bytes) -> dict[str, Any]:
    if not content:
        raise ValueError("The uploaded file is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError(f"The uploaded file exceeds the {MAX_UPLOAD_BYTES // (1024 * 1024)} MB limit.")

    safe_name = os.path.basename(str(filename or "upload"))
    extension = os.path.splitext(safe_name.lower())[1]
    if extension == ".docx":
        text = _extract_docx(content)
    elif extension == ".pdf":
        text = _extract_pdf(content)
    elif extension in {".xlsx", ".xlsm"}:
        text = _extract_xlsx(content)
    elif extension == ".csv":
        text = _extract_csv(content)
    elif extension in {".txt", ".md", ".rtf", ".log", ".json"}:
        text = _decode_text(content)
    else:
        raise ValueError("Unsupported file type. Upload DOCX, PDF, XLSX, CSV, TXT, MD, RTF, LOG or JSON files.")

    cleaned = _normalise(text)
    if not cleaned:
        raise ValueError("No usable text could be extracted from the uploaded file.")
    return {
        "filename": safe_name,
        "extension": extension,
        "character_count": len(cleaned),
        "text": cleaned,
        "truncated": len(str(text or "")) > len(cleaned),
    }
