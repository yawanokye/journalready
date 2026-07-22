from __future__ import annotations

import hashlib
import io
import json
import os
import random
import re
from datetime import datetime
from typing import Any

try:
    from app.source_finder import search_literature_sources
except Exception:  # pragma: no cover
    search_literature_sources = None

from app.research_resources import discover_research_resources, infer_research_route
from app.review_protocol import build_review_protocol_documentation
from app.scholarly_humanizer import (
    analyse_scholarly_style,
    build_humanizer_batches,
    humanize_scholarly_text,
    humanizer_variation_profile,
    scholarly_humanizer_prompt_rules,
    validate_humanizer_preservation,
    variation_targets_met,
)

# Dynamic source context defaults are controlled through helpers below.  Keep this
# reasonably high because journal articles need deeper citation coverage than a
# thesis chapter subsection.
MAX_SOURCE_CONTEXT = int(os.getenv("ARTICLEREADY_ARTICLE_MAX_SOURCE_CONTEXT", "100"))

_RETRACTION_TERMS = re.compile(
    r"\b(retracted|retraction\s+notice|withdrawn|removed\s+article|expression\s+of\s+concern|erratum\s+to\s+retracted)\b",
    flags=re.IGNORECASE,
)

_ATTENTION_RE = re.compile(
    r"\[(?:insert|verify|confirm|provide|supply|complete|replace|check|add|update|obtain|state|specify|include|revise|review|conduct|perform|run|collect|clarify|report|resolve|address|identify|upload|attach|calculate|test|assess|determine|seek|action|required|author\s+action)\b[^\]]*\]",
    flags=re.IGNORECASE,
)

_ACTION_LABEL_RE = re.compile(
    r"\b(?:action required|author action|required action|remaining action|attention required|user action)\s*:\s*",
    flags=re.IGNORECASE,
)

_ACTION_RED = (192, 0, 0)

_ADAM_2020_REFERENCE = (
    "Adam, A. M. (2020). Sample size determination in survey research. "
    "Journal of Scientific Research and Reports, 26(5), 90-97. "
    "https://journaljsrr.com/index.php/JSRR/article/view/1154"
)


def _safe_get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        return None
    try:
        from openai import OpenAI
        return OpenAI(api_key=api_key)
    except Exception:
        return None


def _normalise_gpt56_model(value: str, fallback: str) -> str:
    """Restrict ArticleReady OpenAI routing to GPT-5.6 Terra or Sol."""
    model = str(value or "").strip().lower()
    allowed = {"gpt-5.6-terra", "gpt-5.6-sol"}
    return model if model in allowed else fallback


def _select_article_model(level: str, article_type: str = "", payload: dict[str, Any] | None = None) -> str:
    """Route all ArticleReady OpenAI work through GPT-5.6 Terra or Sol.

    Terra handles cost-balanced drafting. Sol handles doctoral, review, long,
    synthesis-heavy and other high-complexity work. Environment overrides are
    retained, but their defaults are restricted to the GPT-5.6 family.
    """
    payload = payload or {}
    level_l = (level or "").strip().lower()
    type_l = (article_type or "").strip().lower()
    stage = str(payload.get("draft_stage") or "").strip().lower()
    target_words = int(payload.get("target_word_count") or 0)
    long_mode = str(payload.get("long_write_mode") or "auto").strip().lower()

    is_doctoral = any(token in level_l for token in ["phd", "doctor", "dba", "ded", "professional doctorate"])
    is_research_masters = "research masters" in level_l or "mphil" in level_l
    is_review_article = any(token in type_l for token in [
        "systematic", "scoping", "meta-analysis", "meta analysis",
        "review article", "literature review", "conceptual", "theory", "bibliometric", "scientometric",
    ])
    is_long = target_words > 9000 or long_mode == "batch"
    is_completion = stage == "continuation_after_results"

    terra = _normalise_gpt56_model(
        os.getenv("OPENAI_ARTICLE_TERRA_MODEL")
        or os.getenv("OPENAI_ARTICLE_MASTERS_MODEL")
        or os.getenv("OPENAI_ARTICLE_BACHELOR_MODEL")
        or "",
        "gpt-5.6-terra",
    )
    sol = _normalise_gpt56_model(
        os.getenv("OPENAI_ARTICLE_SOL_MODEL")
        or os.getenv("OPENAI_ARTICLE_DOCTORAL_MODEL")
        or os.getenv("OPENAI_ARTICLE_RESEARCH_MODEL")
        or "",
        "gpt-5.6-sol",
    )

    if is_doctoral or is_research_masters or is_review_article or is_long or is_completion:
        return sol
    return terra


def _openai_model_candidates(primary_model: str, *, fallback_model: str = "") -> list[str]:
    """Return an ordered GPT-5.6-only model chain with duplicates removed."""
    terra = _normalise_gpt56_model(os.getenv("OPENAI_ARTICLE_TERRA_MODEL", ""), "gpt-5.6-terra")
    sol = _normalise_gpt56_model(os.getenv("OPENAI_ARTICLE_SOL_MODEL", ""), "gpt-5.6-sol")
    configured_fallback = fallback_model or os.getenv("OPENAI_ARTICLE_FALLBACK_MODEL") or ""
    candidates: list[str] = []
    for raw_model, default in [
        (primary_model, terra),
        (configured_fallback, terra),
        (terra, "gpt-5.6-terra"),
        (sol, "gpt-5.6-sol"),
    ]:
        model = _normalise_gpt56_model(str(raw_model or ""), default)
        if model not in candidates:
            candidates.append(model)
    return candidates


def _call_openai_response_with_fallback(
    client: Any,
    *,
    primary_model: str,
    instructions: str,
    input_payload: Any,
    max_output_tokens: int | None = None,
    fallback_model: str = "",
) -> tuple[str, str, list[str]]:
    """Call the Responses API with a Terra/Sol fallback chain.

    Returns response text, the model that succeeded, and non-fatal attempt notes.
    """
    errors: list[str] = []
    last_exc: Exception | None = None
    for model in _openai_model_candidates(primary_model, fallback_model=fallback_model):
        kwargs: dict[str, Any] = {
            "model": model,
            "instructions": instructions,
            "input": input_payload,
        }
        if max_output_tokens:
            kwargs["max_output_tokens"] = int(max_output_tokens)
        try:
            response = client.responses.create(**kwargs)
            text = _extract_text(response)
            if text:
                return text, model, errors
            errors.append(f"{model} returned no usable text.")
        except Exception as exc:  # pragma: no cover - provider behaviour varies
            last_exc = exc
            errors.append(f"{model}: {str(exc)[:180]}")
    if last_exc:
        raise RuntimeError("; ".join(errors)) from last_exc
    raise RuntimeError("No configured GPT-5.6 model returned usable text.")


def _article_kind(article_type: str) -> str:
    t = (article_type or "").lower()
    if any(x in t for x in ["bibliometric", "scientometric", "science mapping", "co-citation", "co citation"]):
        return "bibliometric_or_scientometric_article"
    if any(x in t for x in ["systematic", "meta-analysis", "meta analysis"]):
        return "systematic_review_or_meta_analysis"
    if any(x in t for x in ["conceptual", "theory article", "theoretical article"]):
        return "conceptual_or_theory_article"
    if any(x in t for x in ["review", "literature review", "scoping"]):
        return "review_article_or_literature_review"
    if any(x in t for x in ["short", "brief", "communication"]):
        return "short_communication_or_brief_report"
    if "conference" in t:
        return "conference_paper"
    return "standard_research_article"


def _is_full_synthesis_article(article_type: str) -> bool:
    """Return True for article types that can be developed as full evidence-synthesis papers.

    These designs do not require primary data collection. Systematic, scoping and
    bibliometric articles still require transparent corpus/search evidence, so any
    missing counts or software-derived results remain explicit author actions.
    """
    return _article_kind(article_type) in {
        "systematic_review_or_meta_analysis",
        "review_article_or_literature_review",
        "conceptual_or_theory_article",
        "bibliometric_or_scientometric_article",
    }


def _article_reference_expectations(article_type: str) -> dict[str, Any]:
    """Return article-type citation-depth guidance without encouraging reference padding."""
    kind = _article_kind(article_type)
    matrix: dict[str, dict[str, Any]] = {
        "standard_research_article": {
            "target_range": "40-60 references",
            "default_search_limit": 60,
            "guidance": "Original research should support the background, theory, methods, results interpretation and discussion with substantial but selective coverage.",
        },
        "review_article_or_literature_review": {
            "target_range": "60-150+ references",
            "default_search_limit": 140,
            "guidance": "Review articles require broad coverage of the relevant literature, including competing streams, methodological differences and unresolved debates.",
        },
        "systematic_review_or_meta_analysis": {
            "target_range": "80-300+ references",
            "default_search_limit": 180,
            "guidance": "Systematic reviews and meta-analyses require exhaustive, transparent coverage consistent with the review protocol and eligibility criteria.",
        },
        "conceptual_or_theory_article": {
            "target_range": "60-140+ references",
            "default_search_limit": 120,
            "guidance": "Conceptual articles require broad, critical and theory-led coverage that supports construct clarification, integration, propositions and a defensible research agenda.",
        },
        "bibliometric_or_scientometric_article": {
            "target_range": "80-250+ references",
            "default_search_limit": 180,
            "guidance": "Bibliometric and scientometric articles require a transparent corpus, reproducible search and cleaning rules, performance analysis, science mapping and careful interpretation of software-derived outputs.",
        },
        "short_communication_or_brief_report": {
            "target_range": "10-25 references",
            "default_search_limit": 30,
            "guidance": "Short communications should use a narrow, high-value reference base because of strict word limits.",
        },
        "conference_paper": {
            "target_range": "15-40 references",
            "default_search_limit": 45,
            "guidance": "Conference papers need enough current literature to justify novelty without overcrowding a short manuscript.",
        },
    }
    out = dict(matrix[kind])
    out["article_kind"] = kind
    out["anti_padding_rule"] = (
        "Do not pad the reference list. Cite only sources that directly support a claim, method, theory, measure, comparison or interpretation. "
        "If available verified sources are fewer than the expected range, add an attention placeholder asking for additional verified literature rather than inventing sources."
    )
    return out


def _article_source_limit(payload: dict[str, Any]) -> int:
    env_limit = os.getenv("ARTICLEREADY_ARTICLE_SOURCE_LIMIT", "").strip()
    if env_limit.isdigit():
        return max(1, int(env_limit))
    return int(_article_reference_expectations(str(payload.get("article_type") or "")).get("default_search_limit", 60))


def _looks_retracted(src: dict[str, Any]) -> bool:
    fields = [
        src.get("title"), src.get("type"), src.get("subtype"), src.get("status"), src.get("publication_status"),
        src.get("update_type"), src.get("relation_type"), src.get("abstract"), src.get("note"), src.get("warning"),
    ]
    combined = " ".join(str(x or "") for x in fields)
    if _RETRACTION_TERMS.search(combined):
        return True
    flags = ["is_retracted", "retracted", "has_retraction", "is_withdrawn", "withdrawn", "removed", "expression_of_concern"]
    return any(bool(src.get(flag)) for flag in flags)


def _build_search_profile(payload: dict[str, Any]) -> dict[str, Any]:
    objectives = []
    for raw in str(payload.get("objectives") or "").split("\n"):
        item = raw.strip(" -;,")
        if item:
            objectives.append(item)
    return {
        "title": str(payload.get("article_title") or payload.get("source_thesis_title") or "").strip(),
        "research_area": str(payload.get("research_area") or payload.get("article_title") or payload.get("source_thesis_title") or "").strip(),
        "study_context": str(payload.get("context") or "").strip(),
        "level": str(payload.get("academic_level") or "Research Masters (e.g. MPhil)").strip(),
        "research_approach": str(payload.get("methodology") or "").strip(),
        "data_type": str(payload.get("article_type") or "").strip(),
        "objectives": objectives[:8],
        "notes": " ".join([
            str(payload.get("target_journal") or ""),
            str(payload.get("variables_constructs") or ""),
            str(payload.get("key_findings") or ""),
            str(payload.get("theory_or_framework") or ""),
            str(payload.get("source_thesis_title") or ""),
            str(payload.get("extraction_focus") or ""),
        ]).strip(),
    }


def _source_key(src: dict[str, Any]) -> str:
    doi = str(src.get("doi") or "").strip().lower()
    if doi:
        doi = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", doi)
        return "doi:" + doi
    title = re.sub(r"[^a-z0-9]+", "", str(src.get("title") or "").lower())[:120]
    return "title:" + title


def _merge_source_banks(*banks: list[dict[str, Any]], limit: int = 120) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    seen: set[str] = set()
    for bank in banks:
        for src in bank or []:
            if not isinstance(src, dict):
                continue
            key = _source_key(src)
            if not key or key == "title:" or key in seen:
                continue
            seen.add(key)
            merged.append(dict(src))
            if len(merged) >= limit:
                return merged
    return merged


def _payload_attached_sources(payload: dict[str, Any]) -> list[dict[str, Any]]:
    direct = payload.get("source_bank") or []
    retrieved = payload.get("retrieved_sources") or {}
    retrieved_items = retrieved.get("sources") or [] if isinstance(retrieved, dict) else []
    attached = _merge_source_banks(direct, retrieved_items, limit=120)
    output: list[dict[str, Any]] = []
    for src in attached:
        record = dict(src)
        record.setdefault("attachment_origin", "attached_before_drafting")
        output.append(record)
    return output


def find_article_sources(payload: dict[str, Any]) -> dict[str, Any]:
    """Search scholarly metadata for the article and return records that can be attached before drafting."""
    if search_literature_sources is None:
        raise RuntimeError("The scholarly source search service is unavailable.")
    profile = _build_search_profile(payload)
    query = str(payload.get("query") or payload.get("source_search_terms") or "").strip()
    result = search_literature_sources(
        profile=profile,
        query=query,
        max_results=int(payload.get("max_results") or 12),
        include_older_foundational=bool(payload.get("include_older_foundational", True)),
    )
    raw_sources = [s for s in (result.get("sources") or []) if isinstance(s, dict)]
    blocked = [s for s in raw_sources if _looks_retracted(s)]
    usable: list[dict[str, Any]] = []
    for src in raw_sources:
        if _looks_retracted(src):
            continue
        record = dict(src)
        record.setdefault("attachment_origin", "manual_source_search")
        usable.append(record)
    return {
        **result,
        "sources": usable,
        "source_bank": usable,
        "source_bank_count": len(usable),
        "excluded_retracted_count": int(result.get("excluded_retracted_count") or 0) + len(blocked),
        "excluded_retracted_titles": list(dict.fromkeys([
            *(result.get("excluded_retracted_titles") or []),
            *(str(src.get("title") or "Untitled") for src in blocked),
        ]))[:20],
        "attachment_note": "These records are attached to the article evidence bank in the browser and will be sent with the drafting request.",
    }


def _search_sources(payload: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    attached_raw = _payload_attached_sources(payload)
    attached_blocked = [src for src in attached_raw if _looks_retracted(src)]
    attached_usable = [src for src in attached_raw if not _looks_retracted(src)]

    search_result: dict[str, Any] = {
        "provider_errors": [],
        "query": str(payload.get("source_search_terms") or ""),
        "sources": [],
        "quality_filters": [],
    }
    searched_usable: list[dict[str, Any]] = []
    searched_blocked: list[dict[str, Any]] = []

    if payload.get("include_source_search", True) and search_literature_sources is not None:
        search_payload = dict(payload)
        search_payload["query"] = str(payload.get("source_search_terms") or "").strip()
        search_payload["max_results"] = _article_source_limit(payload)
        try:
            search_result = find_article_sources(search_payload)
            for src in search_result.get("sources") or []:
                record = dict(src)
                record["attachment_origin"] = "automatic_draft_search"
                searched_usable.append(record)
            searched_blocked = [
                {"title": title, "attachment_origin": "automatic_draft_search"}
                for title in search_result.get("excluded_retracted_titles") or []
            ]
        except Exception as exc:
            search_result = {
                "provider_errors": [{"provider": "source_search", "error": str(exc)[:220]}],
                "query": str(payload.get("source_search_terms") or ""),
                "sources": [],
                "quality_filters": [],
            }

    merged = _merge_source_banks(attached_usable, searched_usable, limit=120)
    blocked = _merge_source_banks(attached_blocked, searched_blocked, limit=40)
    search_result["attached_source_count"] = len(attached_usable)
    search_result["automatic_source_count"] = len(searched_usable)
    search_result["source_bank_count"] = len(merged)
    search_result["sources"] = merged
    search_result["attachment_priority"] = "Sources explicitly attached before drafting are ordered before sources found automatically during drafting."
    return merged, blocked, search_result


def _source_context(sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    records = []
    max_context = max(1, int(os.getenv("ARTICLEREADY_ARTICLE_MAX_SOURCE_CONTEXT", str(MAX_SOURCE_CONTEXT))))
    for idx, src in enumerate(sources[:max_context], start=1):
        authors = src.get("authors") or []
        if isinstance(authors, str):
            authors = [authors]
        records.append({
            "key": f"S{idx}",
            "title": src.get("title", ""),
            "authors": authors,
            "year": src.get("year", ""),
            "source": src.get("source", ""),
            "doi": src.get("doi", ""),
            "url": src.get("url", ""),
            "abstract": str(src.get("abstract") or "")[: int(os.getenv("ARTICLEREADY_ARTICLE_ABSTRACT_CHARS", "700"))],
            "database": src.get("database", ""),
            "relevance_tier": src.get("relevance_tier", ""),
            "citation_count": src.get("citation_count", ""),
            "reference_entry_hint": src.get("apa_hint") or src.get("reference_entry_hint") or "",
            "attachment_origin": src.get("attachment_origin", ""),
        })
    return records


def _extract_text(response: Any) -> str:
    return str(getattr(response, "output_text", "") or "").strip()


def _strip_code_fences(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"^```(?:markdown|md)?\s*|\s*```$", "", text.strip(), flags=re.IGNORECASE | re.MULTILINE).strip()


def _is_survey_research(payload: dict[str, Any]) -> bool:
    haystack = " ".join(
        str(payload.get(k) or "")
        for k in ["article_type", "methodology", "context", "data", "data_results", "data_and_results", "variables_constructs", "objectives", "research_area", "thesis_source_material"]
    ).lower()
    return any(token in haystack for token in ["survey", "questionnaire", "likert", "respondent", "pls-sem", "pls sem", "sem", "cross-sectional"])


def _has_sample_size_determination(payload: dict[str, Any]) -> bool:
    haystack = " ".join(
        str(payload.get(k) or "")
        for k in ["methodology", "data", "data_results", "data_and_results", "context", "objectives", "key_findings", "thesis_source_material"]
    ).lower()
    return any(token in haystack for token in ["sample size", "yamane", "krejcie", "morgan", "cochran", "adam (2020)", "power analysis", "g*power", "gpower"])


def _survey_method_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    survey = _is_survey_research(payload)
    if not survey:
        return {"is_survey_research": False, "rules": []}
    rules = [
        "For survey research, include how common method variance/common method bias was addressed in the Methods section, using prose rather than a bare checklist.",
        "Discuss procedural remedies where appropriate: anonymity, voluntary participation, careful item wording, separation of predictor and outcome sections, varied scale anchors where defensible, reduced evaluation apprehension, and clear instructions.",
        "Discuss statistical remedies where appropriate: Harman's single-factor test, marker-variable approach, unmeasured latent method factor, full collinearity VIF, or another method justified by the selected analysis technique.",
        "If PLS-SEM or SEM is used, distinguish measurement-model quality from common-method-bias assessment.",
        "Do not report CMV/CMB statistics unless the user supplied them. Insert [insert CMV/CMB test result] where needed.",
    ]
    if not _has_sample_size_determination(payload):
        rules.append(
            "If the user did not state a sample-size determination method, justify the required sample size using Adam (2020) for survey research and include the Adam (2020) reference. Do not fabricate the final sample size; use [confirm population size and required sample size] where the population size or confidence level is missing."
        )
    return {
        "is_survey_research": True,
        "sample_size_source_if_missing": _ADAM_2020_REFERENCE if not _has_sample_size_determination(payload) else "User supplied another sample-size determination approach.",
        "rules": rules,
    }


def _human_article_style_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "style_goal": "publishable, strongly human-supervised academic prose",
        "rules": [
            "Write in formal British English with natural academic rhythm, not template-like AI prose.",
            "Use controlled high burstiness and strong lexical variation: combine concise emphasis with developed analytical sentences, without creating fragments or artificial mistakes.",
            "Vary sentence and paragraph length according to argument needs, but keep grammar, citations, headings, tables, equations and references clean.",
            "Build paragraphs around claim, evidence, interpretation, qualification and contribution. Avoid paragraphs that merely list authors.",
            "Vary paragraph openings and transitions so consecutive paragraphs do not begin with the same formula.",
            "Use precise verbs such as indicates, qualifies, complicates, supports, constrains, extends and contradicts.",
            "Avoid filler phrases such as 'in today's world', 'it is important to note', 'delve into', 'plays a crucial role', and repeated 'moreover' or 'furthermore'.",
            "Use cautious scholarly judgement rather than promotional language. Do not overclaim practical or theoretical contributions.",
            "Maintain the author's voice through reasoned transitions, context-specific qualifiers and close links between the study objectives, methods, results and contribution.",
            "Preserve all confirmed facts, statistics, citations, technical terms, placeholders and section logic during stylistic revision.",
            "Do not introduce typos, random grammar errors, lower-case sentence starts, broken citations, invented evidence or artificial mistakes.",
        ],
    }


def _prose_objective_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "objective_style": "prose-first",
        "rules": [
            "Write the article objective, purpose and contribution in prose. Avoid a bare bullet list of objectives unless the target journal explicitly requires it.",
            "In the Introduction, write objectives as a paragraph such as: 'Accordingly, this article examines..., assesses..., and explains...' rather than listing Objective 1, Objective 2 and Objective 3.",
            "Write all major article sections in prose first. Tables may support the argument, but they must not replace the narrative.",
            "If research questions or hypotheses are needed, introduce them through a short prose bridge before listing or tabulating them.",
        ],
    }


def _expert_professor_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    research_area = str(payload.get("research_area") or payload.get("article_title") or "the article's discipline").strip()
    return {
        "research_area": research_area,
        "voice_standard": "senior professor and experienced journal editor in the exact field",
        "rules": [
            f"Write with the conceptual command, disciplinary judgement and methodological precision expected of a leading professor and journal editor in {research_area}.",
            "Do not announce or claim a professorial identity in the manuscript. Demonstrate expertise through accurate conceptual distinctions, method fit, evidence use and restrained scholarly judgement.",
            "Interrogate construct definitions, causal assumptions, identification, measurement validity, analytical alternatives and boundary conditions rather than accepting weak formulations at face value.",
            "State the article's contribution as a precise advance over the cited literature, not as a broad claim of novelty or location alone.",
            "Use field-appropriate terminology consistently and explain technical choices at the depth expected in a strong peer-reviewed journal.",
        ],
    }


def _citation_density_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    article_type = str(payload.get("article_type") or "").lower()
    synthesis = any(token in article_type for token in [
        "systematic", "scoping", "review", "meta-analysis", "meta analysis",
        "conceptual", "theory", "bibliometric", "scientometric",
    ])
    if synthesis:
        overall = {"minimum": 16, "target": 22}
    elif any(token in article_type for token in ["short", "brief", "communication"]):
        overall = {"minimum": 7, "target": 10}
    else:
        overall = {"minimum": 10, "target": 14}
    return {
        "citation_occurrences_per_1000_words": overall,
        "section_guidance": {
            "Introduction": "10-14 citation occurrences per 1,000 words, concentrated on factual context, problem evidence, gap and contribution claims.",
            "Literature Review or Theoretical Background": "16-22 citation occurrences per 1,000 words, organised around claims, tensions and gaps rather than author-by-author summaries.",
            "Review, Conceptual or Bibliometric Synthesis": "18-26 citation occurrences per 1,000 words where the argument, evidence map or conceptual integration depends directly on prior scholarship.",
            "Methods": "6-10 citation occurrences per 1,000 words for design, measurement, sampling, analytical and reporting choices that require authority.",
            "Discussion": "10-15 citation occurrences per 1,000 words, positioned directly beside comparisons, mechanisms, contradictions and boundary-condition claims.",
        },
        "rules": [
            "Place a verified citation in the same sentence as the claim it supports, or immediately after the supported clause. Do not leave several factual or theoretical claims under one distant citation.",
            "Cite every substantive factual, theoretical, methodological and empirical claim that is not common knowledge or supplied as the study's own confirmed result.",
            "Use multiple directly relevant sources when a claim represents a contested debate, broad evidence base or methodological standard.",
            "Distribute citations across the paragraph. Avoid placing one large citation cluster only at the end after several independently supportable claims.",
            "Before returning the manuscript, audit citation coverage section by section and strengthen under-cited claims using only verified records in the supplied source bank.",
            "Do not cite a source merely because a keyword matches. The title, abstract, method, context or reported finding must directly support the claim.",
            "Do not fabricate references or pad the manuscript to hit a numerical target. When the verified evidence bank is insufficient, insert [Author action: Add a verified source that directly supports this claim.]",
            "Keep citation density lower in pure results reporting, but cite method authorities in Methods and theory or prior evidence in Discussion.",
        ],
    }


def _tense_and_action_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "tense_rule": "No future tense anywhere in the article or instrument package.",
        "rules": [
            "Do not use 'will', 'shall', 'is going to', 'are going to', 'will be' or equivalent future constructions in any section, table, note, checklist, declaration or instrument guidance.",
            "Describe an independent Stage 1 study in present tense, for example: 'The study adopts...', 'Data collection uses...', 'The analysis applies...', and 'The questionnaire contains...'.",
            "Describe completed work in past tense and confirmed findings in past or present tense as appropriate.",
            "Do not disguise future tense as proposal language. Rewrite 'will be measured' as 'is measured', 'will be collected' as 'is collected', and 'will assess' as 'assesses'.",
            "Every item that requires author attention, a decision, permission, verification, missing evidence, additional analysis or a later action must appear in one square-bracketed instruction beginning exactly with '[Author action:'.",
            "Do not place advice, recommendations to the author, next-stage instructions, licensing checks, ethics actions, data-access requirements or unresolved methodological decisions in ordinary black prose.",
            "Do not nest square brackets. Consolidate the full action into one bracketed instruction, for example [Author action: Confirm the accessible population, confidence level, margin of error and final sample size.].",
        ],
    }


def _equation_and_framework_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "equation_rules": [
            "When equations or statistical models are required, place each equation in a display equation block using $$ equation $$. Use clean Word-friendly mathematical notation with Unicode Greek letters and subscripts where possible.",
            "Define every symbol directly below the equation in prose. Do not leave equations unexplained.",
            "Use article-appropriate model notation, for example Yᵢ = β₀ + β₁X₁ᵢ + β₂X₂ᵢ + εᵢ, but adapt variable names to the user's actual study.",
            "Do not invent coefficients, p-values or model results. Use [insert model estimates] where results are missing.",
        ],
        "conceptual_framework_rules": [
            "If a conceptual framework is required, structure it as: framework narrative, variable architecture table, hypothesised path table, figure caption, and optional Mermaid or Graphviz code block.",
            "Classify variables as IV, DV, MED, MOD and CV where applicable. Do not force mediation or moderation if not supplied or theoretically justified.",
            "For moderation, show the moderator pointing to the relationship being moderated, not as a vague ordinary predictor of the dependent variable.",
            "Avoid messy ASCII diagrams. Use a concise Mermaid flowchart or Graphviz DOT/Python Graphviz code block where a diagram is useful.",
            "Do not hard-code sample variables. Use only variables, constructs and relationships supplied by the user or clearly inferable from the current article input.",
        ],
        "graphviz_guide": {
            "node_types": ["IV", "DV", "MED", "MOD", "CV"],
            "preferred_orientation": "LR for simple causal frameworks, TB for layered or process frameworks",
            "style_note": "Use publication-style boxes, clear arrows and short labels. If the diagram cannot be rendered in the app, output a Graphviz-compatible code block for the user to render externally.",
        },
    }


def _strong_humanisation_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "mode": "strong_human_supervised",
        "enabled_by_default": True,
        "variation_profile": humanizer_variation_profile(),
        "rules": [
            *scholarly_humanizer_prompt_rules(),
            "Apply controlled high burstiness and high lexical variation while preserving clarity, evidence and disciplinary precision.",
            "Allow sentence length to vary naturally. Split overloaded sentences at defensible clause boundaries, but do not create fragments.",
            "Vary paragraph shape and openings according to the argument rather than through random transition words.",
            "Retain the user's confirmed facts, citations, statistics, quotations, placeholders, equations, tables and reference details exactly unless a substantive revision is justified by supplied evidence.",
            "Remove generic filler and repetitive academic templates. Replace them with context-specific reasoning and precise verbs.",
            "Do not randomise paragraph order, inject unrelated tangents, fabricate citations, or add deliberate grammatical errors.",
            "Do not mention humanisation, AI detection or detector scores in the manuscript.",
        ],
    }


def _article_prompt_quality_pack(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        "expert_professor_requirements": _expert_professor_requirements(payload),
        "human_article_style_requirements": _human_article_style_requirements(payload),
        "strong_humanisation_requirements": _strong_humanisation_requirements(payload),
        "prose_objective_requirements": _prose_objective_requirements(payload),
        "survey_method_requirements": _survey_method_requirements(payload),
        "equation_and_framework_requirements": _equation_and_framework_requirements(payload),
        "citation_density_requirements": _citation_density_requirements(payload),
        "tense_and_action_requirements": _tense_and_action_requirements(payload),
        "reference_depth_requirements": _article_reference_expectations(str(payload.get("article_type") or "")),
    }


def _light_human_article_polish(text: str) -> str:
    """Safe text polish adapted from ai_service.py, without artificial errors."""
    if not text:
        return text
    replacements = {
        r"\bin today's world\b": "in the present context",
        r"\bit is important to note that\b": "",
        r"\bdelve into\b": "examine",
        r"\bplays a crucial role\b": "is important",
        r"\bvarious factors\b": "specific factors",
        r"\bsignificant impact\b": "meaningful effect",
        r"\bthis highlights the importance of\b": "this indicates why",
        r"\bmoreover,\s+moreover\b": "moreover",
        r"\bfurthermore,\s+furthermore\b": "furthermore",
    }
    out = text
    for pattern, repl in replacements.items():
        def case_replacement(match: re.Match[str], replacement: str = repl) -> str:
            if not replacement:
                return replacement
            return replacement[:1].upper() + replacement[1:] if match.group(0)[:1].isupper() else replacement
        out = re.sub(pattern, case_replacement, out, flags=re.IGNORECASE)
    out = re.sub(r"\n{3,}", "\n\n", out)
    out = re.sub(r"[ \t]{2,}", " ", out)
    return out.strip()


_HUMANISATION_OFF_VALUES = {"0", "false", "no", "off"}
_SENTENCE_ABBREVIATIONS = (
    "e.g.", "i.e.", "et al.", "Dr.", "Prof.", "Mr.", "Mrs.", "Ms.",
    "Fig.", "Eq.", "No.", "Vol.", "pp.", "p.", "vs.", "etc.",
)


def _strong_humanisation_enabled() -> bool:
    return os.getenv("ARTICLEREADY_STRONG_HUMANISATION", "1").strip().lower() not in _HUMANISATION_OFF_VALUES


def _humanisation_strength() -> str:
    value = os.getenv("ARTICLEREADY_HUMANISATION_STRENGTH", "strong").strip().lower()
    return value if value in {"light", "standard", "strong"} else "strong"


def _humanisation_rng(text: str, seed_text: str = "") -> random.Random:
    salt = os.getenv("ARTICLEREADY_HUMANISATION_SEED_SALT", "articleready-v1")
    digest = hashlib.sha256(f"{salt}|{seed_text}|{text[:4000]}".encode("utf-8", errors="ignore")).hexdigest()
    return random.Random(int(digest[:16], 16))


def _split_sentences_safe(text: str) -> list[str]:
    if not text.strip():
        return []
    protected = text
    token = "<AR_DOT>"
    for abbr in _SENTENCE_ABBREVIATIONS:
        protected = re.sub(
            re.escape(abbr),
            lambda match: match.group(0).replace(".", token),
            protected,
            flags=re.IGNORECASE,
        )
    protected = re.sub(r"(?<=\d)\.(?=\d)", token, protected)
    parts = re.split(r'(?<=[.!?])\s+(?=[A-Z0-9\[\(“"])', protected)
    return [part.replace(token, ".").strip() for part in parts if part.strip()]


def _sentence_word_count(sentence: str) -> int:
    return len(re.findall(r"\b[\w’'-]+\b", sentence))


def _capitalise_sentence_start(text: str) -> str:
    match = re.search(r"[A-Za-z]", text)
    if not match:
        return text
    idx = match.start()
    return text[:idx] + text[idx].upper() + text[idx + 1:]


def _best_sentence_split(sentence: str) -> tuple[str, str] | None:
    if _sentence_word_count(sentence) < 26:
        return None
    candidates: list[tuple[int, int]] = []
    for match in re.finditer(r";\s+|:\s+|,\s+(?=(?:but|yet|however)\b)", sentence, flags=re.IGNORECASE):
        left = sentence[:match.start()].strip()
        right = sentence[match.end():].strip()
        if _sentence_word_count(left) >= 9 and _sentence_word_count(right) >= 8:
            candidates.append((match.start(), match.end()))
    if not candidates:
        return None
    midpoint = len(sentence) / 2
    start, end = min(candidates, key=lambda pair: abs(pair[0] - midpoint))
    left = sentence[:start].rstrip(" ,;:") + "."
    right = sentence[end:].strip()
    separator = sentence[start:end].strip().lower()
    if right.lower().startswith("but "):
        right = "Even so, " + right[4:]
    elif right.lower().startswith("yet "):
        right = "Even so, " + right[4:]
    elif separator.startswith(","):
        conjunction = re.sub(r"^[,;:]\s*", "", separator).strip()
        if conjunction and not right.lower().startswith(conjunction + " "):
            right = conjunction.capitalize() + " " + right
    right = _capitalise_sentence_start(right)
    return left, right


def _increase_natural_variation(text: str, rng: random.Random | None = None) -> str:
    """Improve rhythm and remove repetitive stock phrases without adding new claims."""
    if not text:
        return text
    rng = rng or _humanisation_rng(text)
    replacements = {
        r"\bin order to\b": "to",
        r"\bdue to the fact that\b": "because",
        r"\bthe fact that\b": "that",
        r"\bas a result of\b": "because of",
        r"\bit is important to note that\b": "",
        r"\bit should be noted that\b": "",
        r"\bthe present study seeks to\b": "this article",
    }
    out = text
    for pattern, replacement in replacements.items():
        def case_replacement(match: re.Match[str], value: str = replacement) -> str:
            if not value:
                return value
            return value[:1].upper() + value[1:] if match.group(0)[:1].isupper() else value
        out = re.sub(pattern, case_replacement, out, flags=re.IGNORECASE)
    sentences = _split_sentences_safe(out)
    if len(sentences) >= 2:
        lengths = [_sentence_word_count(sentence) for sentence in sentences]
        for index in sorted(range(len(sentences)), key=lambda i: lengths[i], reverse=True):
            if lengths[index] < 34:
                break
            split = _best_sentence_split(sentences[index])
            if split:
                sentences[index:index + 1] = list(split)
                break
        repeated_openers = {
            "this study": ["The present article", "The analysis", "This investigation"],
            "the findings": ["These results", "The evidence", "The reported findings"],
            "the results": ["These results", "The evidence", "The estimates"],
        }
        previous = ""
        counts: dict[str, int] = {}
        for index, sentence in enumerate(sentences):
            match = re.match(r"([A-Za-z]+(?:\s+[A-Za-z]+)?)", sentence)
            opener = match.group(1).lower() if match else ""
            if opener == previous and opener in repeated_openers:
                count = counts.get(opener, 0)
                replacement = repeated_openers[opener][count % len(repeated_openers[opener])]
                counts[opener] = count + 1
                sentences[index] = re.sub(rf"^{re.escape(match.group(1))}", replacement, sentence, count=1, flags=re.IGNORECASE)
                opener = replacement.lower()
            previous = opener
        out = " ".join(sentences)
    return re.sub(r"[ \t]{2,}", " ", out).strip()


def _enforce_burstiness(
    text: str,
    target_std_dev: float = 10.0,
    max_uniform: int = 3,
    rng: random.Random | None = None,
) -> str:
    """Increase sentence-length variation by splitting overloaded sentences at safe boundaries."""
    if not text or len(text) < 180:
        return text
    rng = rng or _humanisation_rng(text)
    sentences = _split_sentences_safe(text)
    if len(sentences) < 2:
        return text

    def spread(items: list[str]) -> float:
        lengths = [_sentence_word_count(item) for item in items]
        mean = sum(lengths) / max(1, len(lengths))
        return (sum((value - mean) ** 2 for value in lengths) / max(1, len(lengths))) ** 0.5

    attempts = 0
    while spread(sentences) < target_std_dev and attempts < 3:
        attempts += 1
        lengths = [_sentence_word_count(sentence) for sentence in sentences]
        candidates = [i for i, length in enumerate(lengths) if length >= 28]
        if not candidates:
            break
        index = max(candidates, key=lambda i: lengths[i])
        split = _best_sentence_split(sentences[index])
        if not split:
            break
        sentences[index:index + 1] = list(split)

    # Break long runs of similarly sized sentences only when a safe split is available.
    run = 1
    lengths = [_sentence_word_count(sentence) for sentence in sentences]
    for index in range(1, len(sentences)):
        if abs(lengths[index] - lengths[index - 1]) <= 3:
            run += 1
        else:
            run = 1
        if run > max_uniform and lengths[index] >= 28:
            split = _best_sentence_split(sentences[index])
            if split:
                sentences[index:index + 1] = list(split)
                break
    return " ".join(sentences).strip()


def _add_drafting_artefacts(
    text: str,
    probability_per_500_words: float = 0.35,
    rng: random.Random | None = None,
) -> str:
    """Add restrained scholarly texture by removing formulaic transitions, without adding errors or claims."""
    if not text or len(text.split()) < 120:
        return text
    rng = rng or _humanisation_rng(text)
    if rng.random() > probability_per_500_words:
        return text
    substitutions = [
        (r"^Furthermore,\s+", ""),
        (r"^Moreover,\s+", ""),
        (r"^It is also important to note that\s+", ""),
        (r"^In addition,\s+", "More specifically, "),
    ]
    out = text
    for pattern, replacement in substitutions:
        if re.search(pattern, out, flags=re.IGNORECASE):
            out = re.sub(pattern, replacement, out, count=1, flags=re.IGNORECASE)
            return _capitalise_sentence_start(out)
    return out


def _boost_lexical_richness(
    text: str,
    replacement_probability: float = 0.35,
    rng: random.Random | None = None,
) -> str:
    """Reduce repetitive academic phrasing using meaning-preserving substitutions."""
    if not text or len(text.split()) < 40:
        return text
    rng = rng or _humanisation_rng(text)
    replacements = [
        (r"\bshows that\b", "indicates that"),
        (r"\bsuggests that\b", "points to"),
        (r"\bdemonstrates that\b", "provides evidence that"),
        (r"\bimportant role\b", "substantive role"),
        (r"\bmany studies\b", "a substantial body of research"),
        (r"\bfor example\b", "for instance"),
        (r"\bin contrast\b", "by contrast"),
        (r"\bhas been shown to\b", "has been found to"),
    ]
    out = text
    applied = False
    def replace_once(source: str, pattern: str, replacement: str) -> str:
        def case_replacement(match: re.Match[str]) -> str:
            return replacement[:1].upper() + replacement[1:] if match.group(0)[:1].isupper() else replacement
        return re.sub(pattern, case_replacement, source, count=1, flags=re.IGNORECASE)

    for pattern, replacement in replacements:
        if rng.random() <= replacement_probability and re.search(pattern, out, flags=re.IGNORECASE):
            out = replace_once(out, pattern, replacement)
            applied = True
    if not applied and replacement_probability >= 0.30:
        for pattern, replacement in replacements:
            if re.search(pattern, out, flags=re.IGNORECASE):
                out = replace_once(out, pattern, replacement)
                break
    return out


def _cluster_citations(text: str) -> str:
    """Compatibility hook: citations are preserved and never fabricated or moved mechanically."""
    return text or ""


def _vary_paragraph_openings(text: str) -> str:
    """Reduce repeated paragraph openings with conservative, article-appropriate alternatives."""
    paragraphs = re.split(r"(\n\s*\n)", text or "")
    previous_opener = ""
    alternatives = {
        "this study": ["The present article", "The analysis", "This investigation"],
        "the findings": ["These results", "The evidence", "The reported findings"],
        "the results": ["These results", "The evidence", "The estimates"],
    }
    counters: dict[str, int] = {}
    for index in range(0, len(paragraphs), 2):
        paragraph = paragraphs[index]
        stripped = paragraph.lstrip()
        if not stripped or stripped.startswith(("#", "|", "- ", "* ")):
            continue
        match = re.match(r"([A-Za-z]+(?:\s+[A-Za-z]+)?)", stripped)
        opener = match.group(1).lower() if match else ""
        if opener == previous_opener and opener in alternatives:
            count = counters.get(opener, 0)
            replacement = alternatives[opener][count % len(alternatives[opener])]
            counters[opener] = count + 1
            leading = paragraph[: len(paragraph) - len(stripped)]
            stripped = re.sub(rf"^{re.escape(match.group(1))}", replacement, stripped, count=1, flags=re.IGNORECASE)
            paragraphs[index] = leading + stripped
            opener = replacement.lower()
        previous_opener = opener
    return "".join(paragraphs)


def _force_short_sentences(text: str, target_every_n_words: int = 220) -> str:
    """Create occasional concise sentences by safely splitting existing long sentences."""
    words = re.findall(r"\b[\w’'-]+\b", text or "")
    if len(words) < target_every_n_words:
        return text
    sentences = _split_sentences_safe(text)
    if any(_sentence_word_count(sentence) <= 9 for sentence in sentences):
        return text
    candidates = sorted(
        ((index, _sentence_word_count(sentence)) for index, sentence in enumerate(sentences)),
        key=lambda item: item[1],
        reverse=True,
    )
    for index, length in candidates:
        if length < 30:
            break
        split = _best_sentence_split(sentences[index])
        if split:
            sentences[index:index + 1] = list(split)
            return " ".join(sentences)
    return text


def _add_human_noise(text: str, error_probability: float = 0.0) -> str:
    """Compatibility hook. ArticleReady never introduces deliberate mistakes or citation damage."""
    return text or ""


def _protect_humanisation_regions(text: str) -> tuple[str, dict[str, str]]:
    protected: dict[str, str] = {}

    def replace(match: re.Match[str]) -> str:
        key = f"<AR_PROTECTED_{len(protected)}>"
        protected[key] = match.group(0)
        return key

    out = re.sub(r"```.*?```", replace, text, flags=re.DOTALL)
    out = re.sub(r"(?ms)^\$\$\s*$.*?^\$\$\s*$", replace, out)
    return out, protected


def _restore_humanisation_regions(text: str, protected: dict[str, str]) -> str:
    out = text
    for key, value in protected.items():
        out = out.replace(key, value)
    return out


def _is_humanisable_block(block: str) -> bool:
    stripped = block.strip()
    if not stripped or stripped.startswith("<AR_PROTECTED_"):
        return False
    if stripped.startswith(("#", "|", "- ", "* ")):
        return False
    if re.match(r"^\d+[.)]\s+", stripped):
        return False
    if any(line.lstrip().startswith("|") for line in stripped.splitlines()):
        return False
    if re.match(r"^(Funding|Conflict of interest|Ethics approval|Data availability):", stripped, flags=re.IGNORECASE):
        return False
    return len(re.findall(r"\b[\w’'-]+\b", stripped)) >= 18


def _apply_strong_article_humanisation(
    text: str,
    *,
    seed_text: str = "",
    preserve_reference_section: bool = True,
) -> str:
    """Apply the shared ThesisReady scholarly humaniser deterministically.

    The shared layer removes formulaic phrasing and improves controlled lexical
    and rhythmic variation while preserving headings, evidence, citations,
    numbers, tables, equations, URLs and action placeholders.
    """
    if not text or not _strong_humanisation_enabled():
        return text
    strength = _humanisation_strength()
    mode = {"light": "light", "standard": "balanced", "strong": "deep"}.get(strength, "deep")
    candidate, _report = humanize_scholarly_text(text, mode=mode)
    valid, _issues = validate_humanizer_preservation(
        text,
        candidate,
        max_word_change_ratio=float(humanizer_variation_profile()["model_word_change_limit"]),
    )
    return candidate if valid else text


def _humanizer_model() -> str:
    return _normalise_gpt56_model(
        os.getenv("OPENAI_ARTICLE_HUMANIZER_MODEL") or os.getenv("OPENAI_ARTICLE_TERRA_MODEL") or "",
        "gpt-5.6-terra",
    )


def _humanizer_mode(payload: dict[str, Any] | None = None) -> str:
    requested = str((payload or {}).get("humanizer_mode") or "").strip().lower()
    configured = requested or str(os.getenv("ARTICLEREADY_HUMANIZER_MODE", "balanced") or "balanced").strip().lower()
    return configured if configured in {"off", "light", "balanced", "deep"} else "balanced"


def _humanizer_batch_output_tokens(word_count: int) -> int:
    words = max(250, int(word_count or 0))
    return max(1800, min(9000, int(words * 2.1)))


def _humanize_article_with_model(
    client: Any,
    text: str,
    *,
    payload: dict[str, Any],
    provider_errors: list[Any],
) -> tuple[str, dict[str, Any], list[str]]:
    """Run the same preservation-gated section-batched humaniser used in ThesisReady.

    The deterministic pass always runs when enabled. The optional Terra pass
    touches only weak sections in balanced mode and all eligible sections in
    deep mode. Failure never invalidates the completed article.
    """
    mode = _humanizer_mode(payload)
    local_text, local_report = humanize_scholarly_text(text, mode=mode)
    models_used: list[str] = []
    if mode in {"off", "light"} or not client or not local_text.strip():
        return local_text, local_report, models_used
    if os.getenv("ARTICLEREADY_HUMANIZER_MODEL_PASS", "1").strip().lower() in {"0", "false", "no", "off"}:
        return local_text, local_report, models_used

    variation_profile = humanizer_variation_profile()
    threshold = int(os.getenv("ARTICLEREADY_HUMANIZER_MODEL_THRESHOLD", "97") or 97)
    style_context = any(str(payload.get(key) or "").strip() for key in (
        "author_guidelines", "target_journal", "article_structure", "references_notes", "revision_goals"
    ))
    if (
        mode == "balanced"
        and not style_context
        and int(local_report.get("score") or 100) >= threshold
        and variation_targets_met(local_report, variation_profile)
    ):
        return local_text, local_report, models_used

    batch_words = int(os.getenv("ARTICLEREADY_HUMANIZER_BATCH_WORDS", "1800") or 1800)
    batches = build_humanizer_batches(local_text, max_words=batch_words)
    eligible = [
        index for index, batch in enumerate(batches)
        if not batch.get("protected")
        and int((batch.get("diagnostic") or {}).get("word_count") or 0) >= 120
        and (
            mode == "deep"
            or style_context
            or int((batch.get("diagnostic") or {}).get("score") or 100) < threshold
            or not variation_targets_met(batch.get("diagnostic") or {}, variation_profile)
        )
    ]
    if mode == "balanced":
        eligible.sort(key=lambda index: int((batches[index].get("diagnostic") or {}).get("score") or 100))
        eligible = eligible[:max(1, int(os.getenv("ARTICLEREADY_HUMANIZER_MAX_BATCHES_BALANCED", "6") or 6))]
    else:
        eligible = eligible[:max(1, int(os.getenv("ARTICLEREADY_HUMANIZER_MAX_BATCHES_DEEP", "16") or 16))]
    if not eligible:
        return local_text, local_report, models_used

    chosen = set(eligible)
    output: list[str] = []
    for index, batch in enumerate(batches):
        original = str(batch.get("text") or "")
        if index not in chosen:
            output.append(original)
            continue
        prompt = {
            "task": "Refine this journal article section for natural scholarly flow without changing its substance.",
            "article_type": str(payload.get("article_type") or ""),
            "research_area": str(payload.get("research_area") or payload.get("article_title") or ""),
            "style_diagnostic": batch.get("diagnostic") or {},
            "variation_profile": variation_profile,
            "rules": [
                "Revise rather than restart.",
                *scholarly_humanizer_prompt_rules(),
                "Preserve every heading, citation, reference, date, statistic, objective, question, hypothesis, table, equation, URL and bracketed author-action item exactly.",
                "Do not add evidence, citations, findings, examples, interpretations, recommendations or new sections.",
                "Preserve the order of ideas and the strength of claims.",
                "Keep the word count within six percent of the supplied section.",
                "Return only the revised section with its headings and no report.",
            ],
            "section": original,
        }
        try:
            candidate, used_model, attempt_notes = _call_openai_response_with_fallback(
                client,
                primary_model=_humanizer_model(),
                fallback_model=os.getenv("OPENAI_ARTICLE_SOL_MODEL", "gpt-5.6-sol"),
                instructions="Perform an evidence-preserving, high-variation scholarly naturalness edit. Return only the revised section.",
                input_payload=json.dumps(prompt, ensure_ascii=False, indent=2),
                max_output_tokens=_humanizer_batch_output_tokens(int(batch.get("word_count") or 0)),
            )
            provider_errors.extend(attempt_notes)
            candidate, _ = humanize_scholarly_text(candidate, mode="balanced") if candidate else (original, {})
            valid, _issues = validate_humanizer_preservation(
                original,
                candidate,
                max_word_change_ratio=float(variation_profile["model_word_change_limit"]),
            )
            output.append(candidate if candidate and valid else original)
            if used_model and used_model not in models_used:
                models_used.append(used_model)
        except Exception as exc:  # optional stage
            provider_errors.append(f"Humaniser pass skipped for one section: {str(exc)[:180]}")
            output.append(original)

    candidate = "\n\n".join(part.strip() for part in output if part.strip()).strip()
    valid, issues = validate_humanizer_preservation(
        local_text,
        candidate,
        max_word_change_ratio=float(variation_profile["model_word_change_limit"]),
    )
    final_text = candidate if valid else local_text
    final_text, final_report = humanize_scholarly_text(final_text, mode=mode)
    final_report["model_pass_applied"] = bool(models_used)
    final_report["model_pass_models"] = models_used
    if issues and not valid:
        final_report["preservation_issues"] = issues
    return final_text, final_report, models_used


def _is_independent_article(payload: dict[str, Any]) -> bool:
    return "new independent article" in str(payload.get("source_mode") or "").strip().lower()



# ----------------------------------------------------------------------
# LONG ARTICLE / BATCH WRITING CONTROLS
# ----------------------------------------------------------------------


def _normalise_int(text: Any) -> int | None:
    value = str(text or "").replace(",", "").strip()
    if not value:
        return None
    match = re.search(r"\d+", value)
    return int(match.group(0)) if match else None


def _article_word_range(payload: dict[str, Any]) -> tuple[int, int, int]:
    """Return minimum, maximum and target words requested by the user."""
    explicit = payload.get("target_word_count")
    if explicit:
        words = max(1200, min(int(explicit), 30000))
        span = max(250, int(words * 0.08))
        return max(800, words - span), words + span, words

    raw = str(payload.get("word_limit") or "").replace(",", "")
    numbers = [int(x) for x in re.findall(r"\d{3,5}", raw)]
    if len(numbers) >= 2:
        low, high = min(numbers[0], numbers[1]), max(numbers[0], numbers[1])
        return max(800, low), max(low, high), int(round((low + high) / 2))
    if len(numbers) == 1:
        words = max(1200, min(numbers[0], 30000))
        span = max(250, int(words * 0.10))
        return max(800, words - span), words + span, words

    stage = str(payload.get("draft_stage") or "full_article")
    if stage == "initial_to_methods":
        return 4500, 6000, 5200
    if _article_kind(str(payload.get("article_type") or "")) == "short_communication_or_brief_report":
        return 2500, 4000, 3200
    return 6000, 8000, 7000


def _default_article_sections(payload: dict[str, Any], target_words: int) -> list[dict[str, Any]]:
    stage = str(payload.get("draft_stage") or "full_article")
    article_type = str(payload.get("article_type") or "Empirical research article")
    article_kind = _article_kind(article_type)
    is_review = article_kind in {"review_article_or_literature_review", "systematic_review_or_meta_analysis"}
    is_conceptual = article_kind == "conceptual_or_theory_article"
    is_bibliometric = article_kind == "bibliometric_or_scientometric_article"
    if stage == "initial_to_methods":
        sections = [
            ("Title, protocol-style abstract and keywords", 0.08),
            ("Introduction", 0.20),
            ("Literature review and theoretical positioning", 0.27),
            ("Conceptual or analytical framework and hypotheses/propositions", 0.15),
            ("Methods", 0.25),
            ("Methods readiness checklist, next-stage note and references", 0.05),
        ]
    elif stage == "continuation_after_results":
        sections = [
            ("Updated abstract and article alignment", 0.08),
            ("Earlier sections revised for result alignment", 0.14),
            ("Results or findings", 0.22),
            ("Discussion", 0.26),
            ("Conclusion, contribution, implications and limitations", 0.18),
            ("Declarations, references and source use audit", 0.12),
        ]
    elif is_bibliometric:
        sections = [
            ("Title, abstract and keywords", 0.05),
            ("Introduction and review questions", 0.12),
            ("Bibliometric corpus, search strategy and analytical methods", 0.15),
            ("Descriptive performance analysis", 0.16),
            ("Science mapping and intellectual structure", 0.22),
            ("Thematic evolution, emerging fronts and geographic or collaboration patterns", 0.14),
            ("Discussion, contribution and research agenda", 0.10),
            ("Conclusion, limitations, declarations, references and source audit", 0.06),
        ]
    elif is_conceptual:
        sections = [
            ("Title, abstract and keywords", 0.06),
            ("Introduction and conceptual problem", 0.15),
            ("Conceptual foundations and construct clarification", 0.20),
            ("Critical synthesis and theoretical tensions", 0.22),
            ("Integrative framework and propositions", 0.18),
            ("Theoretical contribution, boundary conditions and research agenda", 0.12),
            ("Conclusion, limitations, declarations and references", 0.07),
        ]
    elif is_review:
        sections = [
            ("Title, abstract and keywords", 0.06),
            ("Introduction", 0.14),
            ("Review protocol and methods", 0.15),
            ("Descriptive profile of the evidence base", 0.10),
            ("Thematic or evidence synthesis", 0.29),
            ("Discussion, conceptual contribution and research agenda", 0.17),
            ("Conclusion, limitations, declarations and references", 0.09),
        ]
    else:
        sections = [
            ("Title, abstract and keywords", 0.06),
            ("Introduction", 0.16),
            ("Literature review and theoretical background", 0.20),
            ("Conceptual or analytical framework and hypotheses", 0.08),
            ("Methods", 0.14),
            ("Results or findings", 0.14),
            ("Discussion", 0.16),
            ("Conclusion, contributions, implications, limitations, declarations and references", 0.06),
        ]
    return [
        {"heading": heading, "target_words": max(250, int(round((target_words * weight) / 50) * 50))}
        for heading, weight in sections
    ]


def _parse_user_article_structure(payload: dict[str, Any], target_words: int) -> list[dict[str, Any]]:
    raw = str(payload.get("article_structure") or "").strip()
    if not raw:
        return _default_article_sections(payload, target_words)
    rows: list[dict[str, Any]] = []
    for line in raw.splitlines():
        clean = line.strip(" -\t")
        if not clean:
            continue
        words = None
        word_match = re.search(r"(?:\(|\b)(\d{3,5})\s*(?:words?|wds?)?\)?\s*$", clean, flags=re.IGNORECASE)
        if word_match:
            words = int(word_match.group(1))
            clean = clean[: word_match.start()].strip(" -:;,()")
        clean = re.sub(r"^\d+(?:\.\d+)*[.)]?\s*", "", clean).strip()
        if clean:
            rows.append({"heading": clean, "target_words": max(150, min(words or 0, 8000)) if words else 0})
    if not rows:
        return _default_article_sections(payload, target_words)
    specified = sum(item["target_words"] for item in rows if item["target_words"])
    missing = [item for item in rows if not item["target_words"]]
    remainder = max(0, target_words - specified)
    if missing:
        share = max(200, int(round((remainder / len(missing)) / 50) * 50))
        for item in missing:
            item["target_words"] = share
    return rows[:18]


def _article_length_structure_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    min_words, max_words, target_words = _article_word_range(payload)
    sections = _parse_user_article_structure(payload, target_words)
    default_threshold = int(os.getenv("ARTICLEREADY_BATCH_DRAFT_WORD_THRESHOLD", "6500") or 6500)
    synthesis_threshold = int(os.getenv("ARTICLEREADY_SYNTHESIS_BATCH_DRAFT_WORD_THRESHOLD", "9500") or 9500)
    batch_threshold = max(default_threshold, synthesis_threshold) if _is_full_synthesis_article(str(payload.get("article_type") or "")) else default_threshold
    return {
        "requested_word_limit": str(payload.get("word_limit") or ""),
        "minimum_words": min_words,
        "target_words": target_words,
        "maximum_words": max_words,
        "structure_source": "user_supplied" if str(payload.get("article_structure") or "").strip() else "article_type_default",
        "long_write_mode": str(payload.get("long_write_mode") or "auto"),
        "batch_threshold_words": batch_threshold,
        "sections": sections,
        "length_rules": [
            "Treat the word limit as a depth and structure requirement, not permission to add filler.",
            "Meet the length through sharper conceptualisation, stronger literature synthesis, method justification, transparent analysis reporting and rigorous discussion.",
            "Follow the user's article structure when supplied. Do not merge or omit requested sections unless the selected stage forbids them.",
            "For long articles, develop each section to its allocated target and avoid ending with a short outline when a full manuscript is requested.",
            "In Auto mode, synthesis articles up to the synthesis batch threshold use one complete drafting pass to reduce web-request timeout risk. Explicit Batch mode still drafts section by section.",
            "If evidence is missing for a requested section, keep the section but insert a red [Author action: ...] item rather than inventing facts, results or citations.",
        ],
    }


def _should_batch_draft(payload: dict[str, Any], length_plan: dict[str, Any]) -> bool:
    mode = str(payload.get("long_write_mode") or "auto").strip().lower()
    if mode == "single_pass":
        return False
    if mode == "batch":
        return True
    target_words = int(length_plan.get("target_words") or 0)
    threshold = int(length_plan.get("batch_threshold_words") or 6500)
    return target_words >= threshold


def _max_output_tokens_for_article(target_words: int, *, section_batch: bool = False) -> int:
    """Set a practical output ceiling from the target prose length."""
    target_words = max(500, int(target_words or 0))
    multiplier = 2.4 if section_batch else 1.9
    base = 1400 if section_batch else 3000
    estimated = int(target_words * multiplier + base)
    default_cap = int(os.getenv("ARTICLEREADY_ARTICLE_MAX_OUTPUT_TOKENS", "24000") or 24000)
    hard_cap = int(os.getenv("ARTICLEREADY_ARTICLE_HARD_OUTPUT_CAP", "60000") or 60000)
    return max(2500, min(estimated, default_cap, hard_cap))


def _article_token_estimate(payload: dict[str, Any], source_records: list[dict[str, Any]], length_plan: dict[str, Any]) -> dict[str, Any]:
    """Estimate token demand for pricing and status reporting."""
    target_words = int(length_plan.get("target_words") or 7000)
    output_tokens = int(round(target_words * 1.35))
    input_chars = 0
    for key in [
        "thesis_source_material", "previous_sections", "continuation_material", "author_guidelines", "data_and_results",
        "research_problem", "objectives", "theory_or_framework", "variables_constructs", "key_findings", "references_notes",
    ]:
        input_chars += len(str(payload.get(key) or ""))
    input_chars += sum(len(str(record.get("title") or "")) + len(str(record.get("abstract") or "")) + 250 for record in source_records)
    base_prompt_tokens = 4500
    input_tokens_single = int(round(input_chars / 4)) + base_prompt_tokens
    batches = len(length_plan.get("sections") or []) if _should_batch_draft(payload, length_plan) else 1
    if batches > 1:
        # Batch mode repeats the core context for each section, but section targets keep output controlled.
        input_tokens = int(round(input_tokens_single * min(batches, 8) * 0.55))
    else:
        input_tokens = input_tokens_single
    return {
        "target_words": target_words,
        "estimated_output_tokens": output_tokens,
        "estimated_input_tokens": max(1000, input_tokens),
        "estimated_total_tokens": max(1000, input_tokens) + output_tokens,
        "drafting_passes": batches,
        "pricing_note": "Estimate only. Actual usage depends on source context, uploaded material, tables, equations, references and model response length.",
    }


def _call_responses_api(
    client: Any,
    *,
    model: str,
    instructions: str,
    prompt: dict[str, Any],
    max_output_tokens: int,
    fallback_model: str = "",
) -> tuple[str, str, list[str]]:
    return _call_openai_response_with_fallback(
        client,
        primary_model=model,
        fallback_model=fallback_model,
        instructions=instructions,
        input_payload=json.dumps(prompt, ensure_ascii=False, indent=2),
        max_output_tokens=max_output_tokens,
    )


def _draft_article_in_batches(
    client: Any,
    *,
    model: str,
    base_prompt: dict[str, Any],
    instructions: str,
    length_plan: dict[str, Any],
    payload: dict[str, Any],
) -> tuple[str, str, list[str], list[str]]:
    """Draft a long article section by section and return article, instrument, warnings and models used."""
    sections = list(length_plan.get("sections") or []) or _default_article_sections(payload, int(length_plan.get("target_words") or 7000))
    generated_sections: list[str] = []
    warnings: list[str] = []
    models_used: list[str] = []
    prior_outline = ""
    for index, section in enumerate(sections, start=1):
        section_prompt = dict(base_prompt)
        section_prompt["task"] = "Draft one section of a longer publishable journal article."
        section_prompt["batch_drafting"] = {
            "enabled": True,
            "current_section_number": index,
            "total_sections": len(sections),
            "current_section": section,
            "all_sections": sections,
            "previous_section_headings_already_drafted": prior_outline,
            "rules": [
                "Write only the current section. Do not repeat earlier sections.",
                "Use the exact current section heading unless journal guidance requires a minor wording adjustment.",
                "Keep citations, tables, equations and author-action placeholders valid within this section.",
                "Do not add final References unless this is the last planned section or the current section explicitly asks for references.",
                "If this is the last planned section, include References and Source Use Audit where needed.",
                "Respect the selected writing stage. Stage 1 stops at Methods and readiness guidance.",
            ],
        }
        section_words = int(section.get("target_words") or max(500, int(length_plan.get("target_words") or 7000) / len(sections)))
        section_prompt["article_length_and_structure"] = {**length_plan, "current_section_target_words": section_words}
        try:
            raw, used_model, attempt_notes = _call_responses_api(
                client,
                model=model,
                fallback_model=os.getenv("OPENAI_ARTICLE_TERRA_MODEL", "gpt-5.6-terra"),
                instructions=instructions,
                prompt=section_prompt,
                max_output_tokens=_max_output_tokens_for_article(section_words, section_batch=True),
            )
            warnings.extend(attempt_notes)
            if used_model and used_model not in models_used:
                models_used.append(used_model)
            section_text, _instrument = _split_draft_package(raw)
            section_text = _strip_code_fences(section_text or raw)
            if section_text:
                generated_sections.append(section_text.strip())
                prior_outline += f"\n- {section.get('heading', 'Section')}"
            else:
                warnings.append(f"Batch section {index} returned no text.")
        except Exception as exc:
            warnings.append(f"Batch section {index} failed: {str(exc)[:180]}")
    article_text = "\n\n".join(part for part in generated_sections if part.strip())
    return article_text, "", warnings, models_used

def _prepare_workflow_payload(payload: dict[str, Any]) -> dict[str, Any]:
    prepared = dict(payload)
    independent = _is_independent_article(prepared)
    stage = str(prepared.get("draft_stage") or "full_article").strip()
    if stage not in {"full_article", "initial_to_methods", "continuation_after_results"}:
        stage = "full_article"
    synthesis_full_article = _is_full_synthesis_article(str(prepared.get("article_type") or ""))
    if independent and stage == "full_article" and not synthesis_full_article:
        stage = "initial_to_methods"
    prepared["draft_stage"] = stage
    prepared["full_synthesis_article"] = synthesis_full_article
    if independent:
        # Independent articles should not depend on thesis/dissertation/project material.
        prepared["source_thesis_title"] = ""
        prepared["thesis_source_material"] = ""
        current_level = str(prepared.get("academic_level") or "").strip()
        if not current_level or current_level == "Research Masters (e.g. MPhil)":
            prepared["academic_level"] = "PhD"
    return prepared


def _resource_markdown(resources: dict[str, Any]) -> str:
    if not resources:
        return ""
    lines = ["## Research Resource Guidance", "", resources.get("search_note") or "Verify every proposed research resource before use."]
    data_sources = resources.get("data_sources") or []
    instruments = resources.get("instrument_sources") or []
    if data_sources:
        lines.extend(["", "### Possible Secondary Data Sources", ""])
        for item in data_sources[:8]:
            line = f"- **{item.get('name', 'Unnamed source')}**, {item.get('provider', '')}. {item.get('coverage', '')} {item.get('suitability', '')}"
            if item.get("url"):
                line += f" Access: {item['url']}"
            if item.get("access_note"):
                line += f" Check: {item['access_note']}"
            lines.append(line.strip())
    if instruments:
        lines.extend(["", "### Possible Questionnaire or Instrument Sources", ""])
        for item in instruments[:8]:
            line = f"- **{item.get('name', 'Unnamed instrument source')}**, {item.get('provider', '')}. {item.get('purpose', '')} {item.get('suitability', '')}"
            if item.get("url"):
                line += f" Access: {item['url']}"
            if item.get("permission_note"):
                line += f" Use condition: {item['permission_note']}"
            lines.append(line.strip())
    return "\n".join(lines).strip()


def _fallback_instrument(payload: dict[str, Any], resources: dict[str, Any]) -> str:
    if not payload.get("include_instrument_draft"):
        return ""
    route = infer_research_route(payload)
    requirements = str(payload.get("instrument_requirements") or "").strip()
    if route == "qualitative_instrument":
        return f"""# Provisional Interview Guide

## Use and Adaptation Note

This guide is provisional and must be aligned with the article objective, ethics approval and study context. {requirements or '[specify the participant group, interview setting and expected duration]'}

## Opening Script

[insert consent, confidentiality, voluntary participation and recording statement]

## Section A: Participant Context

1. Please describe your role or experience in relation to the study topic.
2. What aspects of the context are most important for understanding this issue?

## Section B: Core Article Constructs or Themes

1. How do you understand or experience [insert focal construct or phenomenon]?
2. What factors shape this experience or outcome?
3. Can you describe a specific example?
4. Under what conditions does the pattern become stronger, weaker or different?

## Section C: Implications

1. What changes would improve the situation?
2. What constraints could prevent those changes?

## Probes

Use neutral probes such as “Could you explain further?”, “What happened next?”, “Why was that important?” and “Were there exceptions?”

## Validation Actions

[conduct expert review, cognitive testing or pilot interviews, revise ambiguous questions, and document the final guide]
""".strip()
    return f"""# Provisional Questionnaire and Measurement Plan

## Use and Adaptation Note

This is an original provisional instrument framework, not a reproduction of a proprietary scale. Confirm whether a validated instrument listed in the research-resource guidance is suitable before creating new items. {requirements or '[specify the target population, constructs and preferred response scale]'}

## Section A: Screening and Consent

1. [insert eligibility question]
2. [insert informed-consent statement]

## Section B: Respondent Profile

Include only demographic or organisational variables needed for the analysis and justified by the article.

## Section C: Focal Construct Measures

For each construct, create a clearly labelled item block with a consistent response scale. Draft at least three content-valid items per reflective construct unless the original validated scale specifies otherwise.

| Construct | Proposed item focus | Response format | Original or adapted source | Action required |
|---|---|---|---|---|
| [Construct 1] | [define the content domain] | [e.g. 1 strongly disagree to 5 strongly agree] | [insert verified source] | [expert review and pilot test] |
| [Construct 2] | [define the content domain] | [insert scale] | [insert verified source] | [translation/adaptation check] |
| [Outcome] | [define the outcome domain] | [insert scale] | [insert verified source] | [validity and reliability assessment] |

## Section D: Open Comment

Please provide any additional information that would help explain your responses.

## Instrument Development and Validation Plan

1. Map every item to an article objective, hypothesis or construct.
2. Verify permission and licensing for adopted scales.
3. Use expert review for content validity and cognitive interviewing for comprehension.
4. Pilot the instrument with respondents similar to the target population.
5. Assess reliability and the measurement model using methods appropriate to the study design.
6. Document all adaptations, translations, removed items and scoring changes.
""".strip()


def _fallback_article(payload: dict[str, Any], sources: list[dict[str, Any]], resources: dict[str, Any] | None = None) -> str:
    title = str(payload.get("article_title") or "Article Draft").strip()
    article_type = str(payload.get("article_type") or "Empirical research article").strip()
    citation_style = str(payload.get("citation_style") or "APA 7th").strip()
    stage = str(payload.get("draft_stage") or "full_article")
    ref_expectation = _article_reference_expectations(article_type)
    survey_rules = _survey_method_requirements(payload)
    source_note = ""
    source_audit = ""
    if sources:
        source_note = "\n\nAn attached evidence bank is available for this article, including: " + "; ".join(
            f"S{i+1}: {s.get('title', 'Untitled')} ({s.get('year', 'n.d.')})" for i, s in enumerate(sources[:8])
        ) + ". Each record must pass a relevance check before citation."
        source_audit = "\n\n## Source Use Audit\n\n" + "\n".join(
            f"- S{i+1}: {s.get('title', 'Untitled')} ({s.get('year', 'n.d.')}), attached but not yet integrated because the fallback draft contains placeholders. Verify relevance and cite it only where it directly supports a claim."
            for i, s in enumerate(sources[:12])
        )
    survey_note = ""
    if survey_rules.get("is_survey_research"):
        survey_note = (
            "\n\nFor survey research, this Methods section should explain sample-size determination, questionnaire development, "
            "validity, reliability and common method variance/common method bias remedies. If no sample-size method has been supplied, use Adam (2020) as the sample-size determination source and verify the final population size and required sample size."
        )
    resource_text = _resource_markdown(resources or {})

    if stage == "initial_to_methods":
        return f"""# {title}

## Protocol Status

This is a PhD-depth independent article development draft. The manuscript body intentionally stops at the Methods section. Results, Discussion and Conclusion must be completed only after the analysis is uploaded.

## Abstract

[insert a protocol-style abstract covering the problem, objective, proposed method and intended contribution. Do not report results that do not yet exist]

## Keywords

[insert 4-6 keywords]

## 1. Introduction

This section should establish the focused research problem, article-level gap, context and intended contribution. Write the objective in prose. {source_note}

## 2. Literature Review and Theoretical Positioning

[insert a critical synthesis using verified current and foundational literature]

## 3. Conceptual or Analytical Framework

[insert framework narrative, variable architecture, hypotheses or propositions, and a clean framework figure where appropriate]

## 4. Methods

[insert the proposed design, setting, population or units of analysis, data source or instrument, sample strategy, measures, analysis plan, validity or trustworthiness procedures, ethics, data management and reproducibility steps]{survey_note}

## Methods Readiness Checklist

- [confirm data access or participant access]
- [confirm variable availability or instrument permission]
- [confirm sampling frame and sample-size approach]
- [confirm analysis software and model specification]
- [confirm ethics approval or exemption before data collection]

{resource_text}

## Next Stage

When analysis is ready, select **Stage 2: Complete article after results**, upload this draft together with the results or analysis output, and generate the Results, Discussion, Conclusion, updated Abstract, Declarations and final References.

## References

[insert {citation_style} references cited in Sections 1-4 only]{source_audit}
""".strip()

    if stage == "continuation_after_results":
        return f"""# {title}

## Continuation Draft Status

The supplied previous sections and analysis should be integrated into one coherent article. Preserve accurate material from the earlier draft and revise it only where the new results require alignment.

## Abstract

[update the abstract with the confirmed sample, method, key results and contribution]

## 1. Introduction to Methods

[insert or preserve the supplied earlier article sections]

## 5. Results

[write the results from the uploaded analysis only. Insert tables and figures where supported. Do not invent statistics or themes]

## 6. Discussion

[interpret each confirmed finding against theory, prior studies and the study context]

## 7. Conclusion

[write the conclusion, contribution, implications, limitations and future research]

## Declarations

Funding: [confirm funding statement]

Conflict of interest: [confirm conflict-of-interest statement]

Ethics approval: [confirm ethics approval or exemption]

Data availability: [confirm data availability statement]

## References

[merge and verify all {citation_style} references cited in the completed article]{source_audit}
""".strip()

    article_kind = _article_kind(article_type)
    if article_kind == "conceptual_or_theory_article":
        return f"""# {title}

## Abstract

[Author action: Confirm the final abstract after the conceptual argument, propositions and contribution have been reviewed.]

## Keywords

[Author action: Confirm four to six indexing terms that represent the article's central concepts.]

## 1. Introduction

This full conceptual article establishes the conceptual problem, identifies theoretical tensions and develops a focused contribution from the verified evidence base. {source_note}

## 2. Conceptual Foundations and Construct Clarification

[Author action: Verify the definitions, theoretical lineage and boundaries of each focal construct against the original sources.]

## 3. Critical Synthesis and Theoretical Tensions

[Author action: Complete the critical synthesis using the attached evidence bank and remove any claim that lacks direct source support.]

## 4. Integrative Framework and Propositions

[Author action: Confirm the logic, boundary conditions and wording of each proposition before submission.]

## 5. Discussion and Contribution

[Author action: Confirm the precise theoretical contribution, alternative explanations and conditions under which the framework may not hold.]

## 6. Research Agenda and Practical Implications

[Author action: Retain only implications that follow directly from the conceptual synthesis.]

## 7. Conclusion and Limitations

[Author action: Confirm the scope limitations and the claims that remain provisional.]

## Declarations

Funding: [Author action: Confirm the funding statement.]

Conflict of interest: [Author action: Confirm the conflict-of-interest statement.]

## References

[Author action: Insert and verify all {citation_style} references cited in the article.]{source_audit}
""".strip()

    if article_kind in {"systematic_review_or_meta_analysis", "review_article_or_literature_review"}:
        return f"""# {title}

## Abstract

[Author action: Confirm the review objective, databases, search date, eligibility criteria, evidence-base size, principal synthesis and contribution.]

## Keywords

[Author action: Confirm four to six review and subject indexing terms.]

## 1. Introduction

This full review article defines the review problem, explains why synthesis is needed and states the review questions. {source_note}

## 2. Review Methods

[Author action: Confirm the databases, complete search strings, search dates, inclusion and exclusion criteria, screening process, quality appraisal and synthesis method.]

## 3. Profile of the Evidence Base

[Author action: Insert the verified search, deduplication, screening and inclusion counts. Do not infer PRISMA counts from metadata search results.]

## 4. Evidence Synthesis

[Author action: Complete the thematic, narrative, integrative or quantitative synthesis from the verified included-study corpus.]

## 5. Discussion

[Author action: Confirm how the synthesis resolves, qualifies or exposes tensions in the literature and where evidence remains weak.]

## 6. Contribution and Research Agenda

[Author action: Confirm the conceptual, empirical, methodological and policy contribution and prioritise the research agenda.]

## 7. Conclusion and Limitations

[Author action: Confirm limitations arising from database coverage, language, period, screening and appraisal decisions.]

## Declarations

Funding: [Author action: Confirm the funding statement.]

Conflict of interest: [Author action: Confirm the conflict-of-interest statement.]

Data availability: [Author action: Confirm where the search records, screening file, extraction sheet and review protocol are available.]

## References

[Author action: Insert and verify all {citation_style} references cited in the article.]{source_audit}
""".strip()

    if article_kind == "bibliometric_or_scientometric_article":
        return f"""# {title}

## Abstract

[Author action: Confirm the database, search period, final corpus size, analytical software, principal mapping results and contribution.]

## Keywords

[Author action: Confirm four to six bibliometric and subject indexing terms.]

## 1. Introduction

This full bibliometric article establishes the knowledge-domain problem, review questions and value of performance analysis and science mapping. {source_note}

## 2. Data and Bibliometric Methods

[Author action: Confirm the database, complete search string, search date, document types, language restrictions, deduplication rules, final corpus and software settings.]

## 3. Descriptive Performance Analysis

[Author action: Insert verified annual production, citation, source, author, institution and country statistics from the bibliometric output.]

## 4. Science Mapping and Intellectual Structure

[Author action: Insert verified co-citation, bibliographic coupling, co-authorship and co-word results, including thresholds, normalisation and cluster labels.]

## 5. Thematic Evolution and Emerging Research Fronts

[Author action: Insert verified thematic-map, trend-topic or temporal-network results and explain how the field has changed.]

## 6. Discussion and Contribution

[Author action: Confirm the interpretation of the mapped structures against the underlying publications. Network proximity alone does not establish conceptual agreement or causality.]

## 7. Research Agenda, Limitations and Conclusion

[Author action: Confirm the research agenda and limitations arising from database coverage, citation lag, search design, disambiguation and software choices.]

## Declarations

Funding: [Author action: Confirm the funding statement.]

Conflict of interest: [Author action: Confirm the conflict-of-interest statement.]

Data availability: [Author action: Confirm where the exported corpus, cleaning log, code and network files are available.]

## References

[Author action: Insert and verify all {citation_style} references cited in the article.]{source_audit}
""".strip()

    return f"""# {title}

## Article Type and Target

This manuscript is being prepared as a {article_type}. The target journal is {payload.get('target_journal') or '[insert target journal]'}. The expected reference depth for this article type is {ref_expectation['target_range']}, but the final reference list must include only sources cited in the manuscript.

## Abstract

[insert 180-250 word structured or unstructured abstract after confirming the final results, sample, method and contribution]

## Keywords

[insert 4-6 keywords]

## 1. Introduction

This section should establish the research problem, the current scholarly conversation, the study context and the article's contribution. The article objective should be written in prose rather than as a bare list. {source_note}

## 2. Literature Review and Theoretical Positioning

[insert focused literature synthesis using verified current sources and foundational theory where appropriate]

## 3. Conceptual or Analytical Framework

[insert framework narrative, variable architecture table, hypothesised path table, and Mermaid or Graphviz code block where applicable]

## 4. Methods

[insert article-ready methods section, including design, setting, population/sample, data source, measures, analysis technique, validity/reliability or trustworthiness, and ethics]{survey_note}

## 5. Results

[insert analysed results, tables and figures. Do not invent statistics, coefficients, themes or p-values]

## 6. Discussion

[insert interpretation of findings against theory, prior studies and the study context]

## 7. Conclusion

[insert concise conclusion, contribution, limitations and future research]

## Declarations

Funding: [confirm funding statement]

Conflict of interest: [confirm conflict-of-interest statement]

Ethics approval: [confirm ethics approval or exemption]

Data availability: [confirm data availability statement]

## References

[insert {citation_style} references for sources cited in the article only]{source_audit}
""".strip()

_ACTION_SECTION_RE = re.compile(
    r"\b(?:methods? readiness checklist|author actions?|required actions?|remaining actions?|next stage(?: note)?|research resource guidance|validation actions?|instrument development and validation plan|data requirements?|analysis requirements?)\b",
    flags=re.IGNORECASE,
)

_ACTION_VERBS = (
    "confirm|verify|obtain|calculate|select|decide|provide|upload|insert|revise|report|check|seek|conduct|run|collect|complete|clarify|specify|include|add|replace|adapt|test|assess|determine|use|measure|retain|justify|administer|pilot|translate|document|classify|code|compare|estimate|disclose|cite|remove|address|resolve|avoid|ensure|comply|consider|preserve|explain|state|define|establish|describe|develop|align|integrate|interpret|merge|write|generate"
)

_ACTION_SENTENCE_RE = re.compile(
    rf"(?:\b(?:the authors?|the research team|the study team|the user)\s+(?:should|must|needs?\s+to|is\s+required\s+to)\b|"
    rf"\b(?:is|are)\s+(?:recommended|required|needed)\b|"
    rf"\b(?:should|must)\s+be\s+(?:obtained|confirmed|verified|calculated|selected|provided|uploaded|inserted|revised|reported|checked|conducted|collected|completed|clarified|specified|included|added|replaced|adapted|tested|assessed|determined|used|measured|retained|justified|administered|piloted|translated|documented|classified|coded|compared|estimated|disclosed|cited|removed|addressed)\b|"
    rf"\bsubject\s+to\s+(?:permission|approval|licensing|validation|verification)\b|"
    rf"^(?:do\s+not\s+|(?:{_ACTION_VERBS})\b)|"
    rf"^if\b[^.!?]{{0,180}}?,\s*(?:{_ACTION_VERBS})\b|"
    rf"\[(?:author\s+action|insert|verify|confirm|provide|supply|complete|replace|check|add|update|obtain|state|specify|include|revise|review|conduct|perform|run|collect|clarify|report|resolve|address|identify|upload|attach|calculate|test|assess|determine|seek)\b)",
    flags=re.IGNORECASE,
)

_FUTURE_BASE_VERBS = {
    "adopt": "adopts", "analyse": "analyses", "analyze": "analyzes", "apply": "applies",
    "assess": "assesses", "calculate": "calculates", "capture": "captures", "classify": "classifies",
    "collect": "collects", "compare": "compares", "comply": "complies", "conduct": "conducts",
    "consider": "considers", "contain": "contains", "control": "controls", "define": "defines",
    "describe": "describes", "distinguish": "distinguishes", "document": "documents", "estimate": "estimates",
    "evaluate": "evaluates", "examine": "examines", "exclude": "excludes", "focus": "focuses",
    "handle": "handles", "include": "includes", "interpret": "interprets", "measure": "measures",
    "model": "models", "obtain": "obtains", "present": "presents", "proceed": "proceeds",
    "provide": "provides", "receive": "receives", "recruit": "recruits", "remove": "removes",
    "report": "reports", "retain": "retains", "select": "selects", "specify": "specifies",
    "test": "tests", "use": "uses", "verify": "verifies", "write": "writes",
}

_PLURAL_SUBJECT_WORDS = {
    "we", "they", "you", "data", "results", "participants", "respondents", "graduates", "authors",
    "researchers", "measures", "items", "indicators", "variables", "constructs", "responses", "records",
    "sources", "controls", "coefficients", "effects", "analyses", "findings", "hypotheses", "models",
    "messages", "criteria", "procedures", "dimensions", "relationships", "statistics", "estimates",
}

_SINGULAR_SUBJECT_WORDS = {
    "study", "article", "analysis", "model", "questionnaire", "survey", "sample", "research", "framework",
    "construct", "measure", "procedure", "design", "team", "approval", "consent", "psychological capital",
    "career adaptability", "gig readiness", "effectiveness", "software", "instrument", "population",
    "bias", "status", "process", "assessment", "screening", "bootstrapping", "analysis",
}


def _subject_is_plural(prefix: str) -> bool:
    raw = re.sub(r"(?<=\d),(?=\d)", "", prefix.lower())
    clean = re.sub(r"[^A-Za-z,;: ]+", " ", raw)
    clean = re.sub(r"\s+", " ", clean).strip()
    if not clean:
        return False

    # A coordinated noun list is plural, even when the final item is singular.
    if clean.count(",") >= 1 and " and " in f" {clean} ":
        return True

    # Remove introductory adjuncts and work with the local grammatical subject.
    clause = re.split(r"[;:]", clean)[-1].strip()
    if "," in clause:
        clause = clause.rsplit(",", 1)[-1].strip()
    clause = re.sub(r"^(?:however|therefore|consequently|statistically|empirically|conceptually|methodologically)\s+", "", clause)
    words = [word for word in clause.split() if word]
    if not words:
        return False

    determiner_plural = words[0] in {"these", "those", "several", "multiple", "many", "both", "various"}
    if determiner_plural:
        return True
    while words and words[0] in {"the", "a", "an", "this", "that", "each", "every"}:
        words.pop(0)
    if not words:
        return False

    # Keep only the noun head before a trailing prepositional or participial phrase.
    stop_words = {"with", "of", "in", "for", "among", "across", "using", "based", "through", "within", "from"}
    head: list[str] = []
    for word in words:
        if word in stop_words and head:
            break
        head.append(word)
    head_text = " ".join(head[-5:])
    if " and " in f" {head_text} ":
        return True

    last = head[-1] if head else words[-1]
    # The grammatical head takes precedence over modifiers such as "data" in "data screening".
    if last.endswith("ing") or last in _SINGULAR_SUBJECT_WORDS:
        return False
    if last in _PLURAL_SUBJECT_WORDS:
        return True
    if any(re.search(rf"\b{re.escape(item)}\b", head_text) for item in _SINGULAR_SUBJECT_WORDS):
        return False
    if any(re.search(rf"\b{re.escape(item)}\b", head_text) for item in _PLURAL_SUBJECT_WORDS):
        return True
    if last.endswith("s") and last not in {"analysis", "hypothesis", "effectiveness", "progress", "bias", "status", "process"}:
        return True
    return False


def _third_person_present(verb: str) -> str:
    lower = verb.lower()
    if lower in _FUTURE_BASE_VERBS:
        converted = _FUTURE_BASE_VERBS[lower]
    elif lower == "be":
        converted = "is"
    elif lower == "have":
        converted = "has"
    elif re.search(r"[^aeiou]y$", lower):
        converted = lower[:-1] + "ies"
    elif re.search(r"(?:s|x|z|ch|sh|o)$", lower):
        converted = lower + "es"
    else:
        converted = lower + "s"
    return converted[:1].upper() + converted[1:] if verb[:1].isupper() else converted


def _convert_future_tense_sentence(sentence: str) -> str:
    """Convert common academic future constructions to present tense without changing evidence."""
    value = sentence

    first_future = re.search(r"\b(?:will|shall)\b", value, flags=re.IGNORECASE)
    coordinated_plural = _subject_is_plural(value[: first_future.start()]) if first_future else False

    def coordinated(match: re.Match[str]) -> str:
        conjunction = match.group("conjunction")
        negation = bool(match.group("not"))
        verb = match.group("verb")
        if verb.lower() == "be":
            present = "are" if coordinated_plural else "is"
        elif verb.lower() == "have":
            present = "have" if coordinated_plural else "has"
        else:
            present = verb.lower() if coordinated_plural else _third_person_present(verb)
        if negation:
            auxiliary = "do not" if coordinated_plural else "does not"
            present = verb.lower()
            return f"{conjunction} {auxiliary} {present}"
        return f"{conjunction} {present}"

    value = re.sub(
        r"\b(?P<conjunction>and|but)\s+(?:will|shall)\s+(?P<not>not\s+)?(?P<verb>[A-Za-z]+)",
        coordinated,
        value,
        flags=re.IGNORECASE,
    )

    # Passive and copular constructions.
    def passive(match: re.Match[str]) -> str:
        prefix = match.group("prefix")
        auxiliary = "are" if _subject_is_plural(prefix) else "is"
        return f"{prefix}{auxiliary} "

    value = re.sub(
        r"(?P<prefix>(?:^|(?<=[.!?]\s))[^.!?]{0,140}?)\b(?:will|shall)\s+be\s+",
        passive,
        value,
        flags=re.IGNORECASE,
    )

    # Active constructions.
    def active(match: re.Match[str]) -> str:
        prefix = match.group("prefix")
        negation = bool(match.group("not"))
        verb = match.group("verb")
        plural = _subject_is_plural(prefix)
        if verb.lower() == "have":
            present = "have" if plural else "has"
        elif verb.lower() == "be":
            present = "are" if plural else "is"
        else:
            present = verb.lower() if plural else _third_person_present(verb)
        if negation:
            auxiliary = "do not" if plural else "does not"
            present = verb.lower()
            return f"{prefix}{auxiliary} {present}"
        return f"{prefix}{present}"

    value = re.sub(
        r"(?P<prefix>(?:^|(?<=[.!?]\s))[^.!?]{0,140}?)\b(?:will|shall)\s+(?P<not>not\s+)?(?P<verb>[A-Za-z]+)",
        active,
        value,
        flags=re.IGNORECASE,
    )

    # Going-to future.
    def going_to(match: re.Match[str]) -> str:
        prefix = match.group("prefix")
        verb = match.group("verb")
        plural = _subject_is_plural(prefix)
        present = verb.lower() if plural else _third_person_present(verb)
        return f"{prefix}{present}"

    value = re.sub(
        r"(?P<prefix>(?:^|(?<=[.!?]\s))[^.!?]{0,140}?)\b(?:is|are)\s+going\s+to\s+(?P<verb>[A-Za-z]+)",
        going_to,
        value,
        flags=re.IGNORECASE,
    )
    return value


def _normalise_action_text(text: str) -> str:
    value = str(text or "").strip()
    # Flatten existing or nested action brackets into one clean instruction.
    value = re.sub(r"\[\s*(?:Author action\s*:\s*)?", "", value, flags=re.IGNORECASE)
    value = value.replace("]", " ")
    value = re.sub(r"\s+", " ", value).strip()
    if re.match(r"^insert\s+", value, flags=re.IGNORECASE):
        value = re.sub(r"^insert\s+", "Insert ", value, count=1, flags=re.IGNORECASE)
    else:
        value = re.sub(r"\binsert\s+", "the confirmed ", value, flags=re.IGNORECASE)
    value = _convert_future_tense_sentence(value)
    value = re.sub(r"^(?:Action required|Required action|Attention required|User action)\s*:\s*", "", value, flags=re.IGNORECASE)

    negative_actor = re.match(
        r"^(?:The authors?|The research team|The study team|The user|The study|This study)\s+(?:should|must)\s+not\s+(?P<rest>.+)$",
        value,
        flags=re.IGNORECASE,
    )
    if negative_actor:
        value = "Do not " + negative_actor.group("rest")
    else:
        value = re.sub(
            r"^(?:The authors?|The research team|The study team|The user|The study|This study)\s+(?:should|must|needs?\s+to|is\s+required\s+to)\s+",
            "",
            value,
            flags=re.IGNORECASE,
        )
    value = value[:1].upper() + value[1:] if value else value
    value = value.rstrip()
    if value and value[-1] not in ".!?":
        value += "."
    return f"[Author action: {value}]"


def _sentence_requires_action(sentence: str) -> bool:
    value = sentence.strip()
    if not value or value.startswith("[") and value.lower().startswith("[author action:"):
        return False
    if re.match(r"^(?:H\d+[a-z]?|RQ\d+|P\d+)\s*:", value, flags=re.IGNORECASE):
        return False
    return bool(_ACTION_SENTENCE_RE.search(value))


def _process_action_paragraph(text: str, force_action: bool = False, method_context: bool = False) -> str:
    value = _convert_future_tense_sentence(text)
    if force_action:
        return _normalise_action_text(value)

    protected_actions: dict[str, str] = {}

    def protect(match: re.Match[str]) -> str:
        token = match.group(0)
        if not _ATTENTION_RE.fullmatch(token) and not token.lower().startswith("[author action:"):
            return token
        key = f"<AR_ACTION_{len(protected_actions)}>"
        protected_actions[key] = token
        return key

    value = re.sub(r"\[[^\]\n]+\]", protect, value)
    sentences = _split_sentences_safe(value)
    if not sentences:
        result = value
    else:
        processed: list[str] = []
        for sentence in sentences:
            keys_in_sentence = [key for key in protected_actions if key in sentence]
            restored_sentence = sentence
            for key in keys_in_sentence:
                restored_sentence = restored_sentence.replace(key, protected_actions[key])
            method_advice = bool(
                method_context
                and re.search(
                    r"\b(?:should|must|needs?\s+to|need\s+to|is\s+required|are\s+required|is\s+recommended|are\s+recommended|subject\s+to\s+(?:permission|approval|licensing|validation|verification)|may\s+be\s+considered|could\s+be\s+considered)\b",
                    restored_sentence,
                    flags=re.IGNORECASE,
                )
            )
            if keys_in_sentence or _sentence_requires_action(restored_sentence) or method_advice:
                processed.append(_normalise_action_text(restored_sentence))
            else:
                processed.append(restored_sentence)
        result = " ".join(processed)
    return result


def _process_action_table_line(
    line: str,
    force_action: bool,
    action_columns: set[int] | None = None,
    method_context: bool = False,
) -> str:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    if not cells or all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
        return line
    action_columns = set(action_columns or set())
    for index, cell in enumerate(cells):
        if not cell:
            continue
        if force_action and index > 0:
            cells[index] = _normalise_action_text(cell)
        elif index in action_columns:
            cells[index] = _normalise_action_text(cell)
        else:
            cells[index] = _process_action_paragraph(cell, method_context=method_context)
    return "| " + " | ".join(cells) + " |"


def _enforce_article_writer_output_rules(text: str) -> str:
    """Enforce present/past tense and red-ready bracketed author actions after generation."""
    if not text:
        return text
    lines = text.splitlines()
    output: list[str] = []
    in_code = False
    in_equation = False
    in_references = False
    action_section = False
    method_context = False
    in_table = False
    table_action_columns: set[int] = set()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            output.append(line)
            continue
        if stripped == "$$":
            in_equation = not in_equation
            output.append(line)
            continue
        if in_code or in_equation:
            output.append(line)
            continue
        if re.match(r"^#{1,6}\s+", stripped):
            heading = re.sub(r"^#{1,6}\s+", "", stripped)
            in_references = bool(re.match(r"^(?:references|source use audit)\b", heading, flags=re.IGNORECASE))
            action_section = bool(_ACTION_SECTION_RE.search(heading)) and not in_references
            method_context = bool(
                re.search(
                    r"\b(?:methods?|methodology|research design|study context|population|sampling|eligibility|measures?|measurement|questionnaire|instrument|pre-testing|data collection|common method|response bias|data preparation|analytical strategy|analysis strategy|moderation testing|robustness checks?|subgroup analysis|ethics)\b",
                    heading,
                    flags=re.IGNORECASE,
                )
            ) and not in_references
            in_table = False
            table_action_columns = set()
            output.append(line)
            continue
        if in_references:
            line = re.sub(
                r"\[[^\]\n]+\]",
                lambda m: _normalise_action_text(m.group(0)) if (_ATTENTION_RE.fullmatch(m.group(0)) or m.group(0).lower().startswith("[author action:")) else m.group(0),
                line,
            )
            output.append(line)
            continue
        if not stripped:
            in_table = False
            table_action_columns = set()
            output.append(line)
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            is_separator = bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)
            if not in_table and not is_separator:
                in_table = True
                table_action_columns = {
                    index
                    for index, cell in enumerate(cells)
                    if re.search(r"\b(?:action|required|attention|guidance|adaptation|permission|check|readiness|author input|information needed|decision\s+(?:needed|required)|(?:needed|required)\s+decision)\b", cell, flags=re.IGNORECASE)
                }
                output.append(line)
            elif is_separator:
                output.append(line)
            else:
                output.append(
                    _process_action_table_line(
                        line,
                        action_section,
                        action_columns=table_action_columns,
                        method_context=method_context,
                    )
                )
            continue
        in_table = False
        table_action_columns = set()
        list_match = re.match(r"^(?P<prefix>\s*(?:[-*•]|\d+[.)])\s+)(?P<body>.+)$", line)
        if list_match:
            body = _process_action_paragraph(
                list_match.group("body"),
                force_action=action_section,
                method_context=method_context,
            )
            output.append(list_match.group("prefix") + body)
            continue
        output.append(_process_action_paragraph(line, force_action=action_section, method_context=method_context))

    result = "\n".join(output)
    reference_match = re.search(r"(?im)^#{1,6}\s+(references|source use audit)\b", result)
    body = result[:reference_match.start()] if reference_match else result
    tail = result[reference_match.start():] if reference_match else ""
    body = re.sub(r"\b(?:will|shall)\s+be\b", "is", body, flags=re.IGNORECASE)
    body = re.sub(r"\b(?:will|shall)\s+have\b", "has", body, flags=re.IGNORECASE)
    body = re.sub(r"\b(?:will|shall)\s+(?P<verb>[A-Za-z]+)", lambda m: m.group("verb"), body, flags=re.IGNORECASE)
    body = re.sub(r"[ \t]{2,}", " ", body)
    body = re.sub(r"\n +", "\n", body)
    return (body.rstrip() + ("\n\n" + tail.lstrip() if tail else "")).strip()


def _citation_density_report(text: str) -> dict[str, Any]:
    body = text or ""
    reference_match = re.search(r"(?im)^#{1,6}\s+(references|source use audit)\b", body)
    if reference_match:
        body = body[:reference_match.start()]
    words = len(re.findall(r"\b[\w’'-]+\b", re.sub(r"```.*?```", " ", body, flags=re.DOTALL)))
    parenthetical = re.findall(r"\([^)]*\b(?:19|20)\d{2}[a-z]?\b[^)]*\)", body)
    narrative = re.findall(r"\b[A-Z][A-Za-z'’-]+(?:\s+(?:and|&|et\s+al\.)\s+[A-Z][A-Za-z'’-]+)?\s+\((?:19|20)\d{2}[a-z]?\)", body)
    count = len(parenthetical) + len(narrative)
    density = round((count * 1000 / words), 2) if words else 0.0
    return {"word_count": words, "citation_occurrences": count, "citation_occurrences_per_1000_words": density}


def _finalise_article_text(text: str) -> str:
    text = _strip_code_fences(text or "")
    text = re.sub(r"<span\s+[^>]*>(.*?)</span>", r"\1", text, flags=re.I | re.S)
    text = text.replace("—", ", ").replace(" – ", ", ").replace("–", "-").replace("‑", "-")
    text = _light_human_article_polish(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _split_draft_package(text: str) -> tuple[str, str]:
    raw = str(text or "").strip()
    article_marker = "===ARTICLE_DRAFT==="
    instrument_marker = "===INSTRUMENT_DRAFT==="
    if article_marker in raw:
        raw = raw.split(article_marker, 1)[1]
    if instrument_marker in raw:
        article, instrument = raw.split(instrument_marker, 1)
        return article.strip(), instrument.strip()
    return raw, ""


def _enforce_initial_scope(text: str) -> str:
    # Stop an accidental continuation if the model writes later empirical sections.
    pattern = re.compile(r"(?im)^#{1,3}\s*(?:\d+\.?\s*)?(results|findings|discussion|conclusion|implications)\b")
    match = pattern.search(text or "")
    if not match:
        return text
    kept = (text or "")[:match.start()].rstrip()
    return kept + "\n\n## Next Stage\n\nResults, Discussion and Conclusion were intentionally withheld. Upload the completed analysis with the previous sections and use Stage 2 to finish the article."


def draft_journal_article(payload: dict[str, Any]) -> dict[str, Any]:
    payload = _prepare_workflow_payload(dict(payload))
    material_limit = int(os.getenv("ARTICLEREADY_ARTICLE_MATERIAL_CHARS", "120000") or 120000)
    continuation_limit = int(os.getenv("ARTICLEREADY_ARTICLE_CONTINUATION_CHARS", "140000") or 140000)
    guideline_limit = int(os.getenv("ARTICLEREADY_AUTHOR_GUIDELINE_CHARS", "30000") or 30000)
    data_limit = int(os.getenv("ARTICLEREADY_ARTICLE_DATA_CHARS", "120000") or 120000)
    payload["thesis_source_material"] = str(payload.get("thesis_source_material") or "")[:material_limit]
    payload["previous_sections"] = str(payload.get("previous_sections") or "")[:continuation_limit]
    payload["continuation_material"] = str(payload.get("continuation_material") or "")[:continuation_limit]
    payload["author_guidelines"] = str(payload.get("author_guidelines") or "")[:guideline_limit]
    payload["data_and_results"] = str(payload.get("data_and_results") or "")[:data_limit]
    if not str(payload.get("article_title") or "").strip():
        raise ValueError("Article title or working topic is required.")
    if payload["draft_stage"] == "continuation_after_results" and not (
        str(payload.get("previous_sections") or "").strip() or str(payload.get("continuation_material") or "").strip() or str(payload.get("data_and_results") or "").strip()
    ):
        raise ValueError("Upload or paste the previous article sections and the completed results or analysis before using Stage 2.")

    sources, blocked, search_result = _search_sources(payload)
    source_records = _source_context(sources)
    review_protocol_text, review_protocol_audit = build_review_protocol_documentation(
        payload, search_result, source_records
    )
    payload["review_protocol_documentation"] = review_protocol_text
    length_plan = _article_length_structure_requirements(payload)
    token_estimate = _article_token_estimate(payload, source_records, length_plan)
    resources = payload.get("research_resources") or {}
    if payload.get("include_research_resource_search", True) and not (resources.get("data_sources") or resources.get("instrument_sources")):
        resources = discover_research_resources(
            payload,
            extra_text=" ".join([str(payload.get("research_problem") or ""), str(payload.get("objectives") or "")]),
            max_results=8,
            include_live_search=bool(payload.get("include_source_search", True)),
        )
    payload["research_route"] = infer_research_route(payload)
    model = _select_article_model(str(payload.get("academic_level") or ""), str(payload.get("article_type") or ""), payload)
    client = _safe_get_openai_client()
    provider_errors = list(search_result.get("provider_errors") or []) if isinstance(search_result, dict) else []
    provider_errors.extend(resources.get("provider_errors") or [])
    instrument_text = ""
    model_used = "none"

    if not client or os.getenv("ARTICLEREADY_ARTICLE_USE_AI", "1").strip().lower() in {"0", "false", "no"}:
        article_text = _fallback_article(payload, sources, resources)
        instrument_text = _fallback_instrument(payload, resources)
        mode = "metadata_fallback"
    else:
        current_year = datetime.now().year
        quality_pack = _article_prompt_quality_pack(payload)
        quality_pack["article_length_and_structure_requirements"] = length_plan
        article_inputs = {
            key: value
            for key, value in payload.items()
            if key not in {"source_bank", "retrieved_sources", "research_resources", "review_protocol_documentation"}
        }
        article_inputs["attached_source_count"] = int(search_result.get("attached_source_count") or 0)
        article_inputs["automatic_source_count"] = int(search_result.get("automatic_source_count") or 0)
        stage = payload["draft_stage"]
        stage_rules = {
            "initial_to_methods": [
                "This is Stage 1 for a new independent article. Draft the manuscript body only from the Title through the Methods section.",
                "Use PhD-level depth by default. Develop a strong article-level gap, theory or framework, questions or hypotheses, and a defensible proposed method.",
                "Do not write Results, Findings, Discussion, Conclusion, completed declarations or result-based claims.",
                "The abstract must be protocol-style and must not imply that data have been analysed.",
                "After Methods, add a Methods Readiness Checklist and a short Next Stage note. References and Source Use Audit may follow.",
            ],
            "continuation_after_results": [
                "This is Stage 2. Use previous_sections as the earlier manuscript and continuation_material/data_and_results as the completed analysis evidence.",
                "Integrate the earlier sections and new results into one coherent full article. Preserve accurate prior prose but revise objectives, methods and framing when the analysis requires alignment.",
                "Write Results, Discussion, Conclusion, updated Abstract, implications, limitations, declarations and the final References from supplied evidence only.",
                "Do not invent missing coefficients, p-values, sample details, themes, quotations, tables or figures.",
            ],
            "full_article": (
                [
                    "Draft a complete evidence-synthesis article. Formal primary data collection is not required for this article type.",
                    "Use the attached source bank, verified references, supplied corpus information and any review or bibliometric output as the evidential base.",
                    "For a conceptual article, develop the full conceptual argument, construct clarification, integrative framework, propositions, boundary conditions, contribution and research agenda without inventing empirical findings.",
                    "For a systematic or scoping review, write all full-article sections but report search counts, screening decisions, quality appraisal and included-study results only when supplied. Missing review outputs remain [Author action: ...] items.",
                    "For a bibliometric article, write all full-article sections but report corpus statistics, citation indicators, networks, clusters and thematic evolution only when supplied by verified bibliometric output. Never invent bibliometric results.",
                ]
                if _is_full_synthesis_article(str(payload.get("article_type") or ""))
                else [
                    "Draft the full article because the user supplied a completed thesis, project, dataset analysis or study evidence.",
                ]
            ),
        }[stage]
        prompt = {
            "task": "Draft the requested stage of a publishable journal article and, where requested, a separate provisional instrument package.",
            "draft_stage": stage,
            "article_inputs": article_inputs,
            "current_year": current_year,
            "source_records": source_records,
            "research_resource_guidance": resources,
            "review_protocol_documentation": review_protocol_text,
            "review_protocol_audit": review_protocol_audit,
            "quality_pack": quality_pack,
            "article_length_and_structure": length_plan,
            "token_budget_estimate": token_estimate,
            "stage_rules": stage_rules,
            "strict_rules": [
                "Use supplied target-journal guidance and article_length_and_structure as structural rules. If absent, use an article structure appropriate to the article type and current stage.",
                "Respect the target word range and section allocation. A 7,000-9,000 word article must be developed as a full manuscript, not as a short protocol outline.",
                "When user-supplied article_structure is provided, follow it closely and preserve all requested sections unless the selected stage forbids them.",
                "Treat an independent article as a new study, not as a disguised thesis extraction. Thesis, dissertation and project fields are intentionally blank in independent mode.",
                "Systematic, scoping, conceptual, theory-led and bibliometric articles may be drafted as full independent articles because their evidence base is literature or publication metadata rather than new primary data. Do not force these article types to stop at Methods.",
                "Use review_protocol_documentation as the authoritative completion aid for review-method details. Integrate confirmed details into Methods, Conceptual Approach or Data and Bibliometric Methods rather than repeating one broad author-action paragraph.",
                "Keep ArticleReady metadata discovery separate from formal database searching. Never report its source-bank count as records screened, included studies or the final corpus unless the user supplied and verified those figures independently.",
                "For a conceptual article, preserve integrative theory-building positioning unless the user explicitly selects and documents a genuine systematic or scoping protocol.",
                "When protocol details remain missing, use focused [Author action: ...] items for the specific missing field. Never invent databases, search strings, dates, reviewer numbers, screening decisions, quality scores, PRISMA counts, software settings or final corpus size.",
                "Do not guarantee publication and do not fabricate evidence, results, citations, permissions, ethics approvals, data access or declarations.",
                "Use bracketed attention placeholders for missing details.",
                "Apply a relevance gate to all attached scholarly sources and research resources.",
                "Candidate data sources and instruments are possibilities only. Explain variable coverage, population fit, period, access, ethics, licensing and validation checks before recommending adoption.",
                "Do not reproduce proprietary questionnaire items. When a scale may be copyrighted, identify the source and state that permission or licensing must be checked.",
                "If include_instrument_draft is true, draft a separate original provisional questionnaire, interview guide or measurement plan aligned with the objectives. Do not present it as validated until it has been tested.",
                "Write with the conceptual authority, disciplinary command and methodological judgement expected of a leading professor and experienced journal editor in the stated research area, without announcing that role in the manuscript.",
                "Never use future tense anywhere in the article or instrument package. Use present tense for a proposed Stage 1 design and past tense for completed work.",
                "Every author decision, missing detail, permission check, ethics requirement, additional analysis, next-stage instruction or unresolved issue must appear as one square-bracketed instruction beginning '[Author action:'. No advice may remain in ordinary prose.",
                "Cite substantive claims densely and locally. Place verified citations in the same sentence as the factual, theoretical, methodological or empirical claim they support, while avoiding citation padding and irrelevant sources.",
                "Use the citation-density targets in the quality pack as a section-level minimum. Before returning the manuscript, audit each substantive section and strengthen under-cited claims using only verified source records.",
                "Write in polished formal British English, minimise long dashes, use prose-led objectives and maintain a focused article contribution.",
                "Apply the strong_humanisation_requirements in the quality pack: use controlled high burstiness, varied paragraph shape, precise lexical variation and natural transitions while preserving all evidence, citations, tables, equations and placeholders.",
                "Do not randomise paragraph order, inject tangents, introduce deliberate mistakes, or mention AI detection or humanisation in the article.",
                "Use only confirmed results in Stage 2 or full-article mode. Metadata abstracts do not justify detailed claims about a paper's findings.",
                "References must contain only cited and verified sources. Include a Source Use Audit when source records are supplied.",
            ],
            "output_format": [
                "Return plain Markdown separated by exact markers.",
                "Start with ===ARTICLE_DRAFT=== and then the article draft.",
                "If include_instrument_draft is true, add ===INSTRUMENT_DRAFT=== followed by the separate instrument package. Otherwise omit the instrument marker.",
                "Do not wrap the response in code fences.",
            ],
        }
        instructions = (
            "You are ArticleReady AI's senior disciplinary professor and journal editor for the user's exact research field. Respect the selected stage. "
            "Write with expert conceptual judgement, rigorous method fit and publication-level analytical precision without claiming a professorial identity in the manuscript. "
            "Use no future tense. Stage 1 methods use present tense, and completed work uses past tense. "
            "Place every unresolved author action, permission check, missing evidence item, additional-analysis need or next-stage instruction inside one [Author action: ...] bracket. "
            "Attach verified citations closely to substantive claims and meet the section-level citation-density minimums in the supplied quality pack without padding the text with weak or irrelevant sources. "
            "Apply the supplied strong human-supervised academic writing layer: use natural sentence-length variation, varied paragraph openings, "
            "precise disciplinary language and evidence-led reasoning while preserving citations, technical terms, tables, equations and placeholders. "
            "Do not add deliberate errors, unrelated tangents, or commentary about AI detection or humanisation. "
            "For a new independent empirical study, stop the article body at Methods and provide data-source or instrument guidance without inventing access or validated items. "
            "For a new independent systematic, scoping, conceptual or bibliometric article, draft the full article while keeping unsupplied screening, corpus and software outputs as [Author action: ...] items. "
            "Use the supplied review protocol documentation to complete verified method details, but keep ArticleReady metadata discovery distinct from a formal systematic search and never invent record-flow counts. "
            "For conceptual articles, retain integrative theory-building positioning unless a genuine systematic or scoping protocol is explicitly supplied. "
            "For Stage 2, use the uploaded previous sections and results to complete the manuscript."
        )
        try:
            if _should_batch_draft(payload, length_plan):
                article_text, instrument_text, batch_warnings, batch_models = _draft_article_in_batches(
                    client,
                    model=model,
                    base_prompt=prompt,
                    instructions=instructions,
                    length_plan=length_plan,
                    payload=payload,
                )
                provider_errors.extend(batch_warnings)
                model_used = ", ".join(batch_models) if batch_models else model
                if payload.get("include_instrument_draft") and not instrument_text:
                    instrument_text = _fallback_instrument(payload, resources)
                mode = "ai_batch_draft"
            else:
                raw_text, model_used, attempt_notes = _call_responses_api(
                    client,
                    model=model,
                    fallback_model=os.getenv("OPENAI_ARTICLE_FALLBACK_MODEL", ""),
                    instructions=instructions,
                    prompt=prompt,
                    max_output_tokens=_max_output_tokens_for_article(int(length_plan.get("target_words") or 7000)),
                )
                provider_errors.extend(attempt_notes)
                article_text, instrument_text = _split_draft_package(raw_text)
                if payload.get("include_instrument_draft") and not instrument_text:
                    instrument_text = _fallback_instrument(payload, resources)
                mode = "ai_draft"
            if not article_text:
                article_text = _fallback_article(payload, sources, resources)
        except Exception as exc:
            provider_errors.append(f"OpenAI article drafting failed: {str(exc)[:180]}")
            article_text = _fallback_article(payload, sources, resources)
            instrument_text = _fallback_instrument(payload, resources)
            mode = "metadata_fallback_after_ai_error"

    if review_protocol_text and mode.startswith("metadata_fallback"):
        article_text = f"{article_text.rstrip()}\n\n---\n\n{review_protocol_text}"

    article_text = _finalise_article_text(article_text)
    instrument_text = _finalise_article_text(instrument_text) if instrument_text else ""
    if payload["draft_stage"] == "initial_to_methods":
        article_text = _enforce_initial_scope(article_text)

    humanizer_report: dict[str, Any] = {"mode": _humanizer_mode(payload), "applied": False}
    humanizer_models: list[str] = []
    if mode in {"ai_draft", "ai_batch_draft"}:
        article_text, humanizer_report, humanizer_models = _humanize_article_with_model(
            client,
            article_text,
            payload=payload,
            provider_errors=provider_errors,
        )
        article_text = _finalise_article_text(article_text)

    article_text = _enforce_article_writer_output_rules(article_text)
    instrument_text = _enforce_article_writer_output_rules(instrument_text) if instrument_text else ""
    citation_density = _citation_density_report(article_text)
    density_target = _citation_density_requirements(payload)["citation_occurrences_per_1000_words"]
    citation_density["minimum_target"] = int(density_target["minimum"])
    citation_density["preferred_target"] = int(density_target["target"])
    citation_density["meets_minimum"] = float(citation_density.get("citation_occurrences_per_1000_words") or 0) >= int(density_target["minimum"])

    return {
        "article_text": article_text,
        "instrument_text": instrument_text,
        "draft_stage": payload["draft_stage"],
        "academic_level_used": payload.get("academic_level") or "PhD",
        "research_route": payload.get("research_route") or "undetermined",
        "research_resources": resources,
        "review_protocol_text": review_protocol_text,
        "review_protocol_audit": review_protocol_audit,
        "model_used": model_used if client else "none",
        "mode": mode,
        "source_records_used": source_records,
        "attached_source_count": int(search_result.get("attached_source_count") or 0),
        "automatic_source_count": int(search_result.get("automatic_source_count") or 0),
        "source_bank_count": len(source_records),
        "source_search_terms": str(search_result.get("query") or payload.get("source_search_terms") or ""),
        "excluded_retracted_count": len(blocked),
        "excluded_retracted_titles": [str(s.get("title") or "Untitled") for s in blocked[:10]],
        "provider_errors": provider_errors,
        "reference_depth_guidance": _article_reference_expectations(str(payload.get("article_type") or "")),
        "citation_density_report": citation_density,
        "article_length_plan": length_plan,
        "token_budget_estimate": token_estimate,
        "batch_drafting_applied": mode == "ai_batch_draft",
        "drafting_passes": token_estimate.get("drafting_passes", 1),
        "expert_professor_standard_applied": True,
        "future_tense_guard_applied": True,
        "author_action_bracketing_applied": True,
        "strong_humanisation_applied": mode in {"ai_draft", "ai_batch_draft"} and _strong_humanisation_enabled(),
        "humanisation_strength": _humanisation_strength(),
        "humanizer_report": humanizer_report,
        "humanizer_models_used": humanizer_models,
        "quality_filters": [
            "The strong human-supervised writing layer varies sentence rhythm, paragraph shape and wording without changing confirmed evidence or article structure.",
            "Long-article mode allows user-controlled word targets and section structures, with batch drafting used automatically for long manuscripts unless single-pass mode is selected.",
            "Independent-article mode disables thesis, dissertation and project source fields.",
            "Stage 1 stops the article body at Methods. Stage 2 requires previous sections and completed results or analysis.",
            "Candidate secondary datasets and instruments must be checked for fit, access, permission and validity.",
            "Retracted, withdrawn, removed and expression-of-concern records are excluded where detectable.",
            "Attached scholarly records are filtered through a relevance gate and cannot replace the user's study evidence.",
            "The article is written with the conceptual and methodological judgement expected of a senior professor and experienced journal editor in the field.",
            "Future-tense constructions are converted to present or past tense before the article is returned.",
            "Author advice, missing information and required actions are consolidated into [Author action: ...] brackets for red DOCX formatting.",
            "Verified citations are placed close to the substantive claims they support, subject to source relevance and integrity checks.",
            "Review-protocol and evidence-base documentation is generated separately for synthesis articles using confirmed inputs only.",
            "ArticleReady metadata discovery remains distinct from formal database searching, screening and final-corpus counts.",
            "Missing article details are rendered as bracketed attention placeholders.",
        ],
    }

def _normalise_inline_markdown(text: str) -> str:
    """Normalise common Markdown emphasis typos before DOCX rendering."""
    value = str(text or "")
    # Common model-output typo: one extra closing marker, for example **text***.
    value = re.sub(r"(?<!\*)\*\*(?!\*)([^*\n]+)\*{3}(?=$|[\s.,;:!?])", r"**\1**", value)
    value = re.sub(r"(?<!\*)\*(?!\*)([^*\n]+)\*{2}(?=$|[\s.,;:!?])", r"*\1*", value)
    value = re.sub(r"(?<!_)__(?!_)([^_\n]+)_{3}(?=$|[\s.,;:!?])", r"__\1__", value)
    value = re.sub(r"(?<!_)_(?!_)([^_\n]+)_{2}(?=$|[\s.,;:!?])", r"_\1_", value)
    return value


def _strip_stray_inline_markers(text: str) -> str:
    """Remove unmatched emphasis markers at word boundaries without touching multiplication signs."""
    value = str(text or "")
    value = re.sub(r"(?<!\w)\*{1,3}(?=\w)", "", value)
    value = re.sub(r"(?<=\w)\*{1,3}(?=$|[\s.,;:!?])", "", value)
    value = re.sub(r"(?<!\w)_{1,3}(?=\w)", "", value)
    value = re.sub(r"(?<=\w)_{1,3}(?=$|[\s.,;:!?])", "", value)
    return value


def _split_action_segments(text: str, bold: bool = False, italic: bool = False) -> list[tuple[str, bool, bool, bool]]:
    """Split visible text into normal and author-action segments."""
    value = str(text or "")
    if not value:
        return []
    segments: list[tuple[str, bool, bool, bool]] = []
    position = 0
    bracket_re = re.compile(r"\[[^\]\n]+\]")
    for match in bracket_re.finditer(value):
        if match.start() > position:
            prefix = value[position:match.start()]
            label = _ACTION_LABEL_RE.search(prefix)
            if label:
                if label.start() > 0:
                    segments.append((prefix[:label.start()], bold, italic, False))
                segments.append((prefix[label.start():], bold, italic, True))
            else:
                segments.append((prefix, bold, italic, False))
        token = match.group(0)
        segments.append((token, bold, italic, bool(_ATTENTION_RE.fullmatch(token))))
        position = match.end()
    if position < len(value):
        tail = value[position:]
        label = _ACTION_LABEL_RE.search(tail)
        if label:
            if label.start() > 0:
                segments.append((tail[:label.start()], bold, italic, False))
            segments.append((tail[label.start():], bold, italic, True))
        else:
            segments.append((tail, bold, italic, False))
    if not segments:
        label = _ACTION_LABEL_RE.search(value)
        if label:
            if label.start() > 0:
                segments.append((value[:label.start()], bold, italic, False))
            segments.append((value[label.start():], bold, italic, True))
        else:
            segments.append((value, bold, italic, False))
    return [(part, b, i, action) for part, b, i, action in segments if part]


def _parse_inline_segments(text: str) -> list[tuple[str, bool, bool, bool]]:
    """Parse Markdown emphasis and action placeholders into DOCX-ready segments.

    The returned tuples are: visible text, bold, italic and action-required.
    """
    value = _normalise_inline_markdown(text)
    style_re = re.compile(
        r"(?P<bolditalic>\*\*\*(?P<bolditalic_text>.+?)\*\*\*)"
        r"|(?P<bold>\*\*(?P<bold_text>.+?)\*\*)"
        r"|(?P<italic>(?<!\*)\*(?!\*)(?P<italic_text>.+?)(?<!\*)\*(?!\*))"
        r"|(?P<bolditalic_u>___(?P<bolditalic_u_text>.+?)___)"
        r"|(?P<bold_u>__(?P<bold_u_text>.+?)__)"
        r"|(?P<italic_u>(?<!\w)_(?!_)(?P<italic_u_text>.+?)(?<!_)_(?!\w))",
        flags=re.DOTALL,
    )
    segments: list[tuple[str, bool, bool, bool]] = []
    position = 0
    for match in style_re.finditer(value):
        if match.start() > position:
            plain = _strip_stray_inline_markers(value[position:match.start()])
            segments.extend(_split_action_segments(plain))
        if match.group("bolditalic") is not None:
            inner, bold, italic = match.group("bolditalic_text"), True, True
        elif match.group("bold") is not None:
            inner, bold, italic = match.group("bold_text"), True, False
        elif match.group("italic") is not None:
            inner, bold, italic = match.group("italic_text"), False, True
        elif match.group("bolditalic_u") is not None:
            inner, bold, italic = match.group("bolditalic_u_text"), True, True
        elif match.group("bold_u") is not None:
            inner, bold, italic = match.group("bold_u_text"), True, False
        else:
            inner, bold, italic = match.group("italic_u_text"), False, True
        segments.extend(_split_action_segments(inner or "", bold=bold, italic=italic))
        position = match.end()
    if position < len(value):
        segments.extend(_split_action_segments(_strip_stray_inline_markers(value[position:])))
    if not segments and value:
        segments.extend(_split_action_segments(_strip_stray_inline_markers(value)))
    return segments


def _plain_inline_text(text: str) -> str:
    return "".join(segment[0] for segment in _parse_inline_segments(text))


def _add_inline_runs(paragraph, text: str, force_bold: bool = False) -> None:
    """Render Markdown emphasis and colour author-action text red."""
    from docx.shared import RGBColor

    for visible, bold, italic, action_required in _parse_inline_segments(text):
        run = paragraph.add_run(visible)
        run.bold = bool(bold or force_bold)
        run.italic = bool(italic)
        if action_required:
            run.font.color.rgb = RGBColor(*_ACTION_RED)

def _add_markdown_table(doc, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells and not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row_idx, cells in enumerate(rows):
        row = table.add_row().cells
        for i in range(width):
            value = cells[i] if i < len(cells) else ""
            paragraph = row[i].paragraphs[0]
            paragraph.clear()
            _add_inline_runs(paragraph, value, force_bold=row_idx == 0)

def _add_equation_paragraph(doc, equation: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(equation.strip())
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)


def _add_code_block(doc, code_lines: list[str]) -> None:
    from docx.shared import Pt
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def export_article_docx(article_text: str, title: str = "Journal Article Draft") -> tuple[io.BytesIO, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", (title or "journal_article")[:80]).strip("_") or "journal_article"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False
    equation_buffer: list[str] = []
    in_equation = False

    for raw_line in _finalise_article_text(article_text).splitlines():
        line = raw_line.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                if table_buffer:
                    _add_markdown_table(doc, table_buffer)
                    table_buffer = []
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        stripped = line.strip()
        if stripped == "$$":
            if in_equation:
                _add_equation_paragraph(doc, " ".join(equation_buffer))
                equation_buffer = []
                in_equation = False
            else:
                if table_buffer:
                    _add_markdown_table(doc, table_buffer)
                    table_buffer = []
                in_equation = True
            continue
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            if table_buffer:
                _add_markdown_table(doc, table_buffer)
                table_buffer = []
            _add_equation_paragraph(doc, stripped.strip("$").strip())
            continue
        if in_equation:
            equation_buffer.append(stripped)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            _add_markdown_table(doc, table_buffer)
            table_buffer = []
        if not stripped:
            continue
        if line.startswith("# "):
            p = doc.add_heading("", level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline_runs(p, line[2:].strip())
        elif line.startswith("## "):
            p = doc.add_heading("", level=1)
            _add_inline_runs(p, line[3:].strip())
        elif line.startswith("### "):
            p = doc.add_heading("", level=2)
            _add_inline_runs(p, line[4:].strip())
        elif re.match(r"^[-*•]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, re.sub(r"^[-*•]\s+", "", line).strip())
        elif re.match(r"^\d+[.)]\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+[.)]\s+", "", line).strip())
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)
    if table_buffer:
        _add_markdown_table(doc, table_buffer)
    if code_buffer:
        _add_code_block(doc, code_buffer)
    if equation_buffer:
        _add_equation_paragraph(doc, " ".join(equation_buffer))

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream, f"{safe_title}_journal_article_draft.docx"
