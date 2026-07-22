from __future__ import annotations

import io
import re
from typing import Any


def _safe_filename(value: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9_-]+", "_", str(value or "article_topic_ideas")[:90]).strip("_")
    return clean or "article_topic_ideas"


def _text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _add_labelled_paragraph(doc, label: str, value: Any) -> None:
    text = _text(value)
    if not text:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(text)


def _add_list(doc, label: str, values: Any) -> None:
    items = values if isinstance(values, list) else []
    items = [_text(item) for item in items if _text(item)]
    if not items:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{label}:")
    run.bold = True
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_resource_table(doc, heading: str, items: Any) -> None:
    rows = [item for item in (items or []) if isinstance(item, dict)]
    if not rows:
        return
    doc.add_heading(heading, level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Resource", "Provider", "Coverage or purpose", "Why it may fit", "Access or permission check"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.bold = True
    for item in rows:
        cells = table.add_row().cells
        values = [
            item.get("name"),
            item.get("provider"),
            item.get("coverage") or item.get("purpose"),
            item.get("suitability"),
            item.get("access_note") or item.get("permission_note"),
        ]
        for index, value in enumerate(values):
            cells[index].text = _text(value)
        url = _text(item.get("url"))
        if url:
            cells[0].paragraphs[0].add_run(f"\n{url}")


def export_article_ideas_docx(payload: dict[str, Any]) -> tuple[io.BytesIO, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt

    research_area = _text(payload.get("research_area")) or "Article Topic Ideas"
    ideas = [item for item in (payload.get("ideas") or []) if isinstance(item, dict)][:20]
    if not ideas:
        raise ValueError("Generate at least one article idea before exporting.")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_heading("Article Topic Ideas", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(research_area)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    doc.add_heading("Research and publication context", level=1)
    _add_labelled_paragraph(doc, "Source mode", payload.get("source_mode"))
    _add_labelled_paragraph(doc, "Preferred article type", payload.get("article_type"))
    _add_labelled_paragraph(doc, "Study context", payload.get("context"))
    _add_labelled_paragraph(doc, "Target journal", payload.get("target_journal"))
    _add_labelled_paragraph(doc, "Portfolio guidance", payload.get("portfolio_note"))

    doc.add_heading("Proposed article portfolio", level=1)
    for position, idea in enumerate(ideas, start=1):
        number = idea.get("idea_number") or position
        heading = f"{number}. {_text(idea.get('title')) or 'Untitled article idea'}"
        doc.add_heading(heading, level=2)
        meta = doc.add_paragraph()
        meta.add_run("Article type: ").bold = True
        meta.add_run(_text(idea.get("article_type")) or "Not specified")
        meta.add_run("    Research route: ").bold = True
        meta.add_run(_text(idea.get("research_route")).replace("_", " ") or "Not determined")
        meta.add_run("    Readiness: ").bold = True
        meta.add_run(f"{idea.get('readiness_score', '')}%" if idea.get("readiness_score") not in (None, "") else "Not scored")

        _add_labelled_paragraph(doc, "Article angle", idea.get("angle"))
        _add_labelled_paragraph(doc, "Publishable gap", idea.get("gap"))
        _add_labelled_paragraph(doc, "Overall objective", idea.get("objective"))
        _add_list(doc, "Questions or hypotheses", idea.get("questions_or_hypotheses"))
        _add_labelled_paragraph(doc, "Contribution", idea.get("contribution"))
        _add_labelled_paragraph(doc, "Method and data route", idea.get("method_and_data_route"))
        _add_labelled_paragraph(doc, "Journal fit", idea.get("journal_fit"))
        _add_list(doc, "Suggested sections", idea.get("suggested_sections"))
        _add_list(doc, "Keywords", idea.get("keywords"))
        _add_list(doc, "Evidence still needed", idea.get("evidence_needed"))
        _add_labelled_paragraph(doc, "Scope warning", idea.get("scope_warning"))

        guidance = idea.get("resource_guidance") or {}
        if isinstance(guidance, dict):
            _add_labelled_paragraph(doc, "Resource route", guidance.get("research_route_label") or guidance.get("research_route"))
            _add_resource_table(doc, "Possible secondary data sources", guidance.get("possible_data_sources"))
            _add_resource_table(doc, "Possible questionnaire or instrument sources", guidance.get("possible_instruments"))
            _add_labelled_paragraph(doc, "Resource verification note", guidance.get("guidance_note"))

    resources = payload.get("research_resources") or {}
    if isinstance(resources, dict):
        data_sources = resources.get("data_sources") or []
        instrument_sources = resources.get("instrument_sources") or []
        if data_sources or instrument_sources:
            doc.add_page_break()
            doc.add_heading("Research resource guidance", level=1)
            _add_labelled_paragraph(doc, "Research route", resources.get("research_route_label") or resources.get("research_route"))
            _add_labelled_paragraph(doc, "Search note", resources.get("search_note"))
            _add_resource_table(doc, "Candidate secondary data sources", data_sources)
            _add_resource_table(doc, "Candidate questionnaire or instrument sources", instrument_sources)

    source_records = [item for item in (payload.get("source_records_used") or []) if isinstance(item, dict)][:100]
    if source_records:
        doc.add_page_break()
        doc.add_heading("Relevant scholarly source records retained", level=1)
        for index, source in enumerate(source_records, start=1):
            paragraph = doc.add_paragraph(style="List Number")
            authors = source.get("authors") or []
            if isinstance(authors, list):
                authors = ", ".join(_text(author) for author in authors if _text(author))
            citation = " ".join(part for part in [
                _text(authors),
                f"({_text(source.get('year'))})" if _text(source.get("year")) else "",
                _text(source.get("title")),
                _text(source.get("source") or source.get("database")),
            ] if part)
            paragraph.add_run(citation or f"Source record {index}")
            doi_or_url = _text(source.get("doi") or source.get("url"))
            if doi_or_url:
                paragraph.add_run(f"\n{doi_or_url}")

    quality_filters = [_text(item) for item in (payload.get("quality_filters") or []) if _text(item)]
    if quality_filters:
        doc.add_heading("Quality and verification notes", level=1)
        for item in quality_filters:
            doc.add_paragraph(item, style="List Bullet")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream, f"{_safe_filename(research_area)}_article_topic_ideas.docx"
